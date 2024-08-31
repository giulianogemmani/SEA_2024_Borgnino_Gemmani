# Encoding: utf-8
'''
Created on 2 Oct 2018

@author: AMB 
'''

import wx

from enum import Enum
from enum import IntEnum
from collections import namedtuple
import os

from pathlib import Path

from shutil import copy2

from docx import Document
from docx.shared import Inches
from PIL import Image, ImageFile
from PIL import WmfImagePlugin
from PIL.WmfImagePlugin import WmfHandler

from win32com import client

from tracer.grid import *

from itertools import repeat

from xlsxwriter.workbook import Workbook
from openpyxl import load_workbook
#from openpyxl import Workbook

from  calc_sw_interface.powerfactory_relay_interface import *

import webbrowser
import itertools
import tracer.branch

from calc_sw_interface.db_interface import *
from calc_sw_interface.powerfactory_relay_interface import RelaySetting

import copy
from pydoc import doc

import win32com.client as win32

import datetime
import sys

class FaultType(Enum):
    THREEPHASE = '3psc'
    PHASEPHASE = '2psc'
    SINGLEPHASE = 'spgf'
    PHASEPHASEGROUND = '2pgf'
    THREEPHASEUNBALANCE = '3rst'
    SINGLEPHASENEUTRAL = 'spnf'
    PHASEPHASENEUTRAL = '2pnf'
    THREEPHASENEUTRAL = '3pnf'


class OutputDetail(IntEnum):
    DISABLED = 0
    NORMAL = 1
    DEBUG = 2
    VERBOSEDEBUG = 3


class BreakerID(IntEnum):
    BREAKER1 = 0
    BREAKER2 = 1


class CoordRule:
    '''
    class used to define a relay type, relay location (line, trafo, etc) and relay
    model to create then a link to a given coordination rule
    '''

    def __init__(self, relay_type, location, relay_model='any'):
        self.relay_type = relay_type
        self.location = location
        self.relay_model = relay_model

    def __hash__(self):
        return hash((self.relay_type, self.location, self.relay_model))

    def __eq__(self, other):
        return (self.relay_type, self.location, self.relay_model) == \
            (other.relay_type, other.location, other.relay_model) or \
            ((self.relay_model == 'any' or other.relay_model == 'any') and \
             (self.relay_type, self.location) == (other.relay_type, other.location))

    def __ne__(self, other):
        # Not strictly necessary, but to avoid having both x==y and x!=y
        # True at the same time
        return not(self == other)


class RelayDB:
    '''
    class used to define a relay type, relay location (line, trafo, etc) and relay
    model to create then a link to a given table in a DB
    '''

    def __init__(self, relay_category, protected_item_type, relay_model='any'):
        self.relay_category = relay_category
        self.protected_item_type = protected_item_type
        self.relay_model = relay_model

    def __hash__(self):
        return hash((self.relay_category, self.protected_item_type, self.relay_model))

    def __eq__(self, other):
        return (self.relay_category, self.protected_item_type, self.relay_model) == \
           (other.relay_category, other.protected_item_type, other.relay_model) or\
        ((self.relay_model == 'any' or other.relay_model == 'any') and \
             (self.relay_category, self.protected_item_type) == (other.relay_category, other.protected_item_type))

    def __ne__(self, other):
        # Not strictly necessary, but to avoid having both x==y and x!=y
        # True at the same time
        return not(self == other)


class PSET():
    '''
    Protective device setting calculation logic class
    '''

    def __init__(self, interface):
        '''
        Constructor
        '''
        self.interface = interface
        self.output_detail = OutputDetail.VERBOSEDEBUG
        self.Relay = namedtuple(
            "Relay", "name manufacturer model substation busbar cubicle \
                      protected_item voltage cbr_optime pf_relay is_backup_relay\
                      relay_branch from_station to_station \
                      phase_minimum_threshold ground_minimum_threshold measurement")

        self.relay_list = []
        self.Fault = namedtuple(
            "Fault", "faulted_line network_status type fault_position reference_breaker\
                     fault_resistance disconnected_elements_names voltage area \
                     zone from_station to_station fault_clearance_time \
                     error")
        self.LDF = namedtuple(
            "LDF", "network_status disconnected_elements_names \
            ldf_trip_time error")
        # data structure used to get the info froma sheet to the DB
        self.RelaySheetRecord = namedtuple("RelaySheetRecord", \
                               "file sheet_name setting_table multiple")
        self.Currents = namedtuple("Currents", "Ia Ib Ic In")
        self.Voltages = namedtuple("Voltages", "Va Vb Vc Vn")
        self.Tripping_data = namedtuple ("Tripping_data", "trip_time \
                                            tripping_element_string currents \
                                            at_load_bus breaker_failure")
        self.grid = None  # the object containing the whole one line diagram info
        self.calculation_list = []  # the list with all fault objects
        self.trip_time_list = [[]]  # matrix with all trip time (one row for each fault)
        self.relay_matrix = []  # the logic time interconnection between relays

        # definition of the coordination rules
        self.coordination_rules = {}

        self.test_running = False
        self.workbook = None
        # repository for all already calculated shc values
        self.lines_iscc_data = {}
        self.shunts_iscc_data = {}
        self.generators_zscc_data = {}
        self.z3s_zscc_data = {}
        # constant values
        self.enabled = 0
        self.disabled = 1
        # generator critical times repository
        self.limit_times = {}

    def initialize(self, window, input_settings=None):
        '''
        function initializing all dictionaries to configure the network, 
        the fault type, the fault location and the other selection criteria
        '''
        self.relay_list.clear()
        self.calculation_list.clear()
        self.trip_time_list.clear()
        self.relay_matrix.clear()

        settings = input_settings if input_settings != None \
        else window.GetSettings()
        netstatus = namedtuple("netstatus", "active function")

        FaultConfig = namedtuple("ftconfig", "active R type")

        self.fault_types = {
            "1PH": FaultConfig(active=settings['Fslg'], R=0, type=FaultType.SINGLEPHASE),
            "2PHG": FaultConfig(active=settings['Fdlg'], R=0, type=FaultType.PHASEPHASEGROUND),
            "2PH": FaultConfig(active=settings['Fltl'], R=0, type=FaultType.PHASEPHASE),
            "3PH": FaultConfig(active=settings['Ftph'], R=0, type=FaultType.THREEPHASE),
            "1PHR": FaultConfig(active=settings['Fslgr'], R=settings['FslgrValue'], type=FaultType.SINGLEPHASE),
            "2PHR": FaultConfig(active=settings['Fltlr'], R=settings['FltlrValue'], type=FaultType.PHASEPHASE),
            "2PHGR":  FaultConfig(active=settings['Fdlgr'], R=settings['FdlgrValue'], type=FaultType.PHASEPHASEGROUND),
        }

        # create the selection criteria dictionary
        Criteria = namedtuple("Criteria", "itemlist function")

        grid_lists = [self.interface.get_content(window.grid_list[gridindex]) for gridindex in settings["GridList"]]
        grid_items = list(itertools.chain(*grid_lists))

        self.line_selection_criteria = {
            "cpArea":   Criteria(itemlist=[window.area_list[areaindex] for areaindex in settings["AreaList"]], function=self.interface.get_attribute),
            "cpZone":   Criteria(itemlist=[window.zone_list[zoneindex] for zoneindex in settings["ZoneList"]], function=self.interface.get_attribute),
            "uknom":    Criteria(itemlist=[window.voltage_list[voltageindex] for voltageindex in settings["VoltageList"]], function=self.interface.get_area_bus_voltage),
            "grid": Criteria(itemlist=grid_items, function=self.get_element),
            "path": Criteria(itemlist=[self.interface.get_content(window.path_list[pathindex]) for pathindex in settings["PathList"]], function=self.void_function_2_params)
        }

        self.generator_selection_criteria = {
            "cpArea":   Criteria(itemlist=[window.area_list[areaindex] for areaindex in settings["AreaList"]], function=self.interface.get_attribute),
            "cpZone":   Criteria(itemlist=[window.zone_list[zoneindex] for zoneindex in settings["ZoneList"]], function=self.interface.get_attribute),
            "ugn":    Criteria(itemlist=[window.voltage_list[voltageindex] for voltageindex in settings["VoltageList"]], function=self.interface.get_generator_voltage),
            "grid": Criteria(itemlist=grid_items, function=self.get_element),
            "path": Criteria(itemlist=[self.interface.get_content(window.path_list[pathindex]) for pathindex in settings["PathList"]], function=self.void_function_2_params)
        }

        self.trafo_selection_criteria = {
            "cpArea":   Criteria(itemlist=[window.area_list[areaindex] for areaindex in settings["AreaList"]], function=self.interface.get_attribute),
            "cpZone":   Criteria(itemlist=[window.zone_list[zoneindex] for zoneindex in settings["ZoneList"]], function=self.interface.get_attribute),
            "bushv":    Criteria(itemlist=[window.voltage_list[voltageindex] for voltageindex in settings["VoltageList"]], function=self.interface.get_transformer_winding_voltage_of),
            "grid": Criteria(itemlist=grid_items, function=self.get_element),
            "path": Criteria(itemlist=[self.interface.get_content(window.path_list[pathindex]) for pathindex in settings["PathList"]], function=self.void_function_2_params)
        }

        # definition of the coordination rules
#         self.coordination_rules.update({CoordRule(relay_type = self.interface.RelayType.OVERCURRENT.value,
#                                                    location = 'ElmTr') : \
#                                                   self.set_trafo_overcurrent})
        self.coordination_rules.update({CoordRule(relay_type=self.interface.RelayType.DISTANCE.value, \
                                                   location='ElmLne', \
                                                   relay_model='F21 Distance Polygonal') : \
                                                  self.set_line_distance_polygonal})
        self.coordination_rules.update({CoordRule(relay_type=self.interface.RelayType.DISTANCE.value, \
                                                   location='ElmLne', \
                                                   relay_model='F21 Distance Mho') : \
                                                  self.set_line_distance_mho})
#         self.coordination_rules.update({CoordRule(relay_type = self.interface.RelayType.OVERCURRENT.value,
#                                                    location = 'ElmLne') : \
#                                                   self.set_line_overcurrent})
#         self.coordination_rules.update({CoordRule(relay_type = self.interface.RelayType.DISTANCE.value,\
#                                     location = 'ElmSym',\
#                                     relay_model = 'F21 Under Z Polygonal') : \
#                                     self.set_generator_underz_polygonal})
        self.coordination_rules.update({CoordRule(relay_type=self.interface.RelayType.OVERCURRENT.value, \
                                    location='ElmShnt', \
                                    relay_model='F50_F51  Phase overcurrent') : \
                                    self.set_shunt_phase_overcurrent})
        self.coordination_rules.update({CoordRule(relay_type=self.interface.RelayType.OVERCURRENT.value, \
                                    location='ElmShnt', \
                                    relay_model='F50N_F51N Neutral overcurrent') : \
                                    self.set_shunt_neutral_overcurrent})

        # ldf status dictionary
        self.Ldfstatus = namedtuple("Ldfstatus", "ldf_already_calculated ldf_failed")

        self.ldf_status_list = { "Intact network" : self.Ldfstatus(ldf_already_calculated=False, ldf_failed=False)}
        # the actual ldf status
        self.ldf_status = self.ldf_status_list["Intact network"]
        return settings

    def create_system_layout(self, window, input_settings, grid_only=False):
        '''
        function creating the self.grid, and the self.relay_list variable
        it returns the lines and the transformers which have been selected
        accordingly to the criteria
        '''
        try:
            # init return variables
            lines , transformers = [] , []
            # basic initialization
            settings = self.initialize(window, input_settings)
            if self.is_dialog_setting_ok(window, settings) == False:
                return []
            if self.output_detail >= OutputDetail.NORMAL:
                self.interface.print("\nLoading grid elements")
            if self.grid == None:
                self.grid = Grid(self.interface)  # @UndefinedVariable
            # exit point if we need only the grid
            if grid_only == True:
                return lines , transformers
            if self.output_detail >= OutputDetail.NORMAL:
                self.interface.print("\nLoading protective elements")
            self.fill_relay_list(self.grid)

            # the "stay around bus bar" is not defined , use the other criteria
            if len(settings['StudySelectedBus']) == 0:
                lines = self.interface.get_lines(self.line_selection_criteria)
                transformers = self.interface.get_transformers(self.trafo_selection_criteria)
            else:
                lines = self.grid.get_bus_lines_from(busbar_name=settings['StudySelectedBus'],
                                        number_of_steps=settings['StudySelectedBusExtent'])
                transformers = self.grid.get_bus_transformers_from(busbar_name=settings['StudySelectedBus'],
                                        number_of_steps=settings['StudySelectedBusExtent'])
            if self.output_detail >= OutputDetail.DEBUG:
                self.interface.print("\nRelay list")
#                     for relay in self.relay_list:
#                         self.interface.print("%r" % (relay,))
            if self.output_detail >= OutputDetail.NORMAL:
                self.interface.print("\nCreating protection layout")
            self.create_protecton_layout_logic(self.grid, lines,
                        number_of_levels=2)
        except KeyboardInterrupt:
                if self.output_detail >= OutputDetail.NORMAL:
                    self.interface.print("\nInterrupted by the user!")

        self.interface.print("\nData collection completed!")
        return lines, transformers

################################################################################
################################################################################
################################################################################

    def set_relays(self, window, input_settings=None):
        '''
        the main procedures of the relay automatic calculation
        '''
        settings = self.initialize(window, input_settings)
        if self.is_dialog_setting_ok(window, settings) == False:
            return 1

        # Print some info in the PF output window
        if self.output_detail.value >= OutputDetail.NORMAL.value:
            self.interface.print("*********************************************************")
            self.interface.print("    Time Distance Creator Tool (Beta) 0.1")
            self.interface.print("                 Setting relays           ")
            self.interface.print("*********************************************************\n")

        try:
            # create the header files for the Z3 calculation results
            z3_line_workbook = self.create_z3_output_file_header(window, \
                                                                       'line')
            z3_trafo_workbook = self.create_z3_output_file_header(window, \
                                                                 'transformer')
            lines = []
            trafos = []
            generators = []
            lines, trafos = self.create_system_layout(window, input_settings)
            generators = self.interface.get_available_generators()
            shunts = self.interface.get_available_shunts()
            elements = lines  # + trafos + generators + shunts

            # set initially the SHC to complete
            self.interface.set_shc_basic_configuration(\
                                            self.interface.SHC_Mode.COMPLETE)
#             lines = self.interface.get_lines(self.line_selection_criteria)
            for element_index, element in enumerate(elements):
#                 cubicles = [self.interface.get_line_cubicle_i_of(element)]
#                 if self.interface.get_class_name_of(element) != "ElmSym":
#                     cubicles.append(self.interface.get_line_cubicle_j_of(element))
                cubicles = self.interface.get_branch_cubicles_of(element)
                found_relays = []
                for cubicle in cubicles:
                    if cubicle:
                        found_relays += self.interface.get_cubicle_relay_of(cubicle)
                if self.output_detail >= OutputDetail.NORMAL:
                    self.interface.print("\n\nSetting element " + self.\
                    interface.get_name_of(element) + '(' + \
                    str(element_index) + '\\' + str(len(elements)) + ')')
                for relay in found_relays:
                    # temporary code to add the shunt neutral relays if missing
                    relay_type_strings = [self.interface.get_relay_model_name_of(relay)\
                                      for relay in found_relays]
                    if "F50N_F51N Neutral overcurrent" not in relay_type_strings and\
                    self.interface.get_class_name_of(element) == "ElmShnt":
                        phase_relay_type = self.interface.get_relay_type_of(relay)
                        type_folder = self.interface.get_parent(phase_relay_type)
                        neutral_relay_types = self.interface.get_content(\
                                type_folder,
                                search_string='F50N_F51N Neutral overcurrent.*')
                        if neutral_relay_types:
                            neutral_relay_type = neutral_relay_types[0]
                            phase_relay_name = self.interface.get_name_of(relay)
                            neutral_relay_name = phase_relay_name.replace("50", "50N")
                            neutral_relay_name = neutral_relay_name.replace("51", "51N")
                            new_neutral_relay = self.interface.create_relay(\
                                        neutral_relay_type, \
                                        self.interface.get_relay_cubicle_of(relay), \
                                        neutral_relay_name)
                            found_relays.append(new_neutral_relay)
                            self.interface.print("\t The " + self.interface.get_name_of(\
                                new_neutral_relay) + " relay has been created")
                        else:
                            self.interface.print("")
                    # END temporary code
                    try:
                        self.coordination_rules[CoordRule(
                                self.interface.get_relay_category_of(relay), \
                                self.interface.get_class_name_of(element),
                                self.interface.get_relay_model_name_of(relay))]\
                                (relay, element)
                    except KeyError:
                        pass
                # write in the excel result file the Z3 calculation results
                if self.output_detail >= OutputDetail.NORMAL:
                    self.interface.print("\t\t Saving relays measured impedances")
                self.fill_z3_output_file(z3_line_workbook, z3_trafo_workbook)
                if self.output_detail >= OutputDetail.NORMAL:
                    self.interface.print("\t\t Saved relay measured impedances")
#                     if self.interface.is_overcurrent_relay(relay):
#                         if self.output_detail >= OutputDetail.NORMAL:
#                             self.interface.print("Setting line relay " +\
#                                                 self.interface.get_name_of(relay))
#                         self.coordination_rules[CoordRule(self.interface.RelayType.\
#                                         DIRECTIONAL, 'ElmLne')](relay, line)

#             for trafo in trafos:
#                 cubicles = self.interface.get_transformer_cubicle_of(trafo)
#                 found_relays = self.interface.get_cubicle_relay_of(cubicles[0])
#                 for relay in found_relays:
#                     if self.interface.is_overcurrent_relay(relay):
#                         self.coordination_rules[CoordRule(self.interface.RelayType.\
#                                         OVERCURRENT, 'ElmTr')](relay, trafo)

        except KeyboardInterrupt:
            if self.output_detail >= OutputDetail.NORMAL:
                self.interface.print("\nInterrupted by the user!")

        self.interface.print("\nSetting calculation completed!")
        return

################################################################################
################################################################################
################################################################################

    def save_relay_settings_to_DB(self):
        '''
        function savening all relay settings in the DB
        '''
        self.interface.print("\n\t\t *** Saving all relay settings into the DB ***")
        # create the db connection
        try:
            db = MYSQLInterface(self.interface)
        except:
            self.interface.print("ERROR: Cannot connect to the MySQL database")
            return
        relays = self.interface.get_relays([self.interface.RelayType.OVERCURRENT,
                            self.interface.RelayType.DISTANCE])
        # here the relationship between the relay type and its interface
        relay_pf_interface_dict = {
            "F21 Distance Mho" : PowerFactoryMhoDistanceRelayInterface,
            "F21 Distance Polygonal" : PowerFactoryPolygonalDistanceRelayInterface,
            "F21 Under Z Polygonal" : PowerFactoryPolygonalDistanceRelayInterface,
            "F50N_F51N Neutral overcurrent" : PowerFactoryNeutralOvercurrentRelayInterface ,
            "F50_F51  Phase overcurrent" : PowerFactoryOvercurrentRelayInterface,
            "F67N_F50N_F51N Neutral directional overc" : PowerFactoryNeutralOvercurrentRelayInterface,
            "F67_F50_F51 Phase directional overc" : PowerFactoryOvercurrentRelayInterface}
        # here the relationship between the type of the protected item and the
        # relay settings table in the DB
        relay_db_interface_dict = { RelayDB(relay_category="Overcurrent", \
                                          protected_item_type="Elmlne"): \
                                          self.save_line_overcurrent, \
                                    RelayDB(relay_category="Distance", \
                                          protected_item_type="ElmLne", \
                                          relay_model="F21 Distance Polygonal"):\
                                          self.save_line_distance_poly, \
                                    RelayDB(relay_category="Distance", \
                                          protected_item_type="ElmLne", \
                                          relay_model="F21 Distance Mho"):\
                                          self.save_line_distance_mho, \
                                    RelayDB(relay_category="Overcurrent", \
                                          protected_item_type="ElmTr2", \
                                          relay_model="F50_F51  Phase overcurrent"):\
                                          self.save_trafo_phase_overcurrent, \
                                    RelayDB(relay_category="Overcurrent", \
                                          protected_item_type="ElmTr2", \
                                          relay_model="F50N_F51N Neutral overcurrent"):\
                                          self.save_trafo_neutral_overcurrent, \
                                    RelayDB(relay_category="Overcurrent", \
                                          protected_item_type="ElmTr3", \
                                          relay_model="F50_F51  Phase overcurrent"):\
                                          self.save_trafo_phase_overcurrent, \
                                    RelayDB(relay_category="Overcurrent", \
                                          protected_item_type="ElmTr3", \
                                          relay_model="F50N_F51N Neutral overcurrent"):\
                                          self.save_trafo_neutral_overcurrent,
                                    RelayDB(relay_category="Overcurrent", \
                                          protected_item_type="ElmShnt", \
                                          relay_model="F50_F51  Phase overcurrent"):\
                                          self.save_shunt_phase_overcurrent,
                                    RelayDB(relay_category="Overcurrent", \
                                          protected_item_type="ElmShnt", \
                                          relay_model="F50N_F51N Neutral overcurrent"):\
                                          self.save_shunt_neutral_overcurrent,
                                    RelayDB(relay_category="Distance", \
                                          protected_item_type="ElmSym", \
                                          relay_model="F21 Under Z Polygonal"):\
                                          self.save_generator_underimpdance_poly
                                          }

        for relay_index, relay in enumerate(relays):
            protected_item = self.interface.get_relay_protected_item_of(relay)
            protected_item_type = self.interface.get_class_name_of(protected_item)
            relay_name = self.interface.get_name_of(relay)
            relay_type_name = self.interface.get_relay_model_name_of(relay)
            relay_category = self.interface.get_relay_category_of(relay)
            relay_interface = relay_pf_interface_dict[relay_type_name](self.interface, relay)
            relay_settings = relay_interface.read_settings()
            if self.output_detail >= OutputDetail.NORMAL:
                self.interface.print("Evaluating relay " + relay_name + "(" + \
                                     str(relay_index) + "\\" + str(len(relays))\
                                     +")")
            try:
                transfer_to_db = relay_db_interface_dict[RelayDB(relay_category=relay_category, \
                                          protected_item_type=protected_item_type, \
                                          relay_model=relay_type_name)]
                ct_ratio, vt_ratio = self.get_relay_ratios_of(relay)
                transfer_to_db(relay_name, relay_settings, db, ct_ratio, vt_ratio)
            except Exception as e:
                self.interface.print("Query/function call error: " + str(e))
        self.interface.print("\n\t\t *** All relay settings have been saved into the DB *** ")

    def save_trafo_phase_overcurrent(self, relay_name, settings, db, ct_ratio, vt_ratio):
        '''
        function saving in the DB the phase overcurrent settings protecting a trafo
        '''
        self._save_into_DB(table_name="pt_tr_phase_overcurrent", \
                     setting_table=\
                    {'i_stg1_inservice' :  0 if settings['Enable #1'] else 1,
                   'i_stg1_trippingdirection' : '1',
                   'i_stg1_characteristictype' : self.interface.get_name_of(\
                                                settings['Characteristic #1']),
                   'i_stg1_pickupprimary' : settings['Current #1'] * ct_ratio,
                   'i_stg1_pickupsecondary' : settings['Current #1'],
                   'timedial' : settings['Time #1'],
                   'i_stg2_inservice' :  0 if settings['Enable #2'] else 1,
                   'i_stg2_trippingdirection' : '1',
                   'i_stg2_characteristictype' : 'Definite time',
                   'i_stg2_pickupprimary' : settings['Current #2'] * ct_ratio,
                   'i_stg2_pickupsecondary' : settings['Current #2'],
                   'trippingtimedelay' : settings['Time #2']}, \
                   relay_name=relay_name, settings=settings, db=db)

    def save_trafo_neutral_overcurrent(self, relay_name, settings, db, ct_ratio, vt_ratio):
        '''
        function saving in the DB the neutral overcurrent settings protecting a trafo
        '''
        self._save_into_DB(table_name="pt_tr_neutral_overcurrent", \
                     setting_table=\
                    {'ig_stg1_inservice' :  0 if settings['Enable #1'] else 1,
                   'ig_stg1_trippingdirection' : '1',
                   'ig_stg1_characteristictype' : self.interface.get_name_of(\
                                                settings['Characteristic #1']),
                   'ig_stg1_pickupprimary' : settings['Current #1'] * ct_ratio,
                   'ig_stg1_pickupsecondary' : settings['Current #1'],
                   'timedial' : settings['Time #1'],
                   'ig_stg2_inservice' :  0 if settings['Enable #2'] else 1,
                   'ig_stg2_trippingdirection' : '1',
                   'ig_stg2_characteristictype' : 'Definite time',
                   'ig_stg2_pickupprimary' : settings['Current #2'] * ct_ratio,
                   'ig_stg2_pickupsecondary' : settings['Current #2'],
                   'trippingtimedelay' : settings['Time #2']}, \
                   relay_name=relay_name, settings=settings, db=db)

    def save_line_distance_poly(self, relay_name, settings, db, ct_ratio, vt_ratio):
        '''
        function saving in the DB the polygonal distance settings protecting a line
        '''
        z_ratio = vt_ratio / ct_ratio
        self._save_into_DB(table_name="pt_line_distance", \
                     setting_table={\
                   'zone1_inservice' : 0 if settings['Phase Phase Polygonal 1 Out service'] else 1,
                   'zone1_trippingdirection' : 1,
                   'zone1_ph_ph_prireach_x' : settings['Phase Phase Polygonal 1 X'],
                   'zone1_ph_ph_secreach_x' : settings['Phase Phase Polygonal 1 X'] / z_ratio,
                   'zone1_ph_ph_prireach_r' : settings['Phase Phase Polygonal 1 R'],
                   'zone1_ph_ph_secreach_r' : settings['Phase Phase Polygonal 1 R'] / z_ratio,
                   'zone1_ph_e_prireach_x' : settings['Phase Earth Polygonal 1 X'],
                   'zone1_ph_e_secreach_x' : settings['Phase Earth Polygonal 1 X'] / z_ratio,
                   'zone1_ph_e_prireach_r' : settings['Phase Earth Polygonal 1 R'],
                   'zone1_ph_e_secreach_r' : settings['Phase Earth Polygonal 1 R'] / z_ratio,
                   'zone1_relayangle' : settings['Phase Phase Polygonal 1 Relay Angle'],
                   'zone1_tdelay' : settings['Phase Phase Polygonal 1 delay'],
                   'zone2_inservice' : 0 if settings['Phase Phase Polygonal 2 Out service'] else 1,
                   'zone2_trippingdirection' : '1',
                   'zone2_ph_ph_prireach_x' : settings['Phase Phase Polygonal 2 X'],
                   'zone2_ph_ph_secreach_x' : settings['Phase Phase Polygonal 2 X'] / z_ratio,
                   'zone2_ph_ph_prireach_r' : settings['Phase Phase Polygonal 2 R'],
                   'zone2_ph_ph_secreach_r' : settings['Phase Phase Polygonal 2 R'] / z_ratio,
                   'zone2_ph_e_prireach_x' : settings['Phase Earth Polygonal 2 X'],
                   'zone2_ph_e_secreach_x' : settings['Phase Earth Polygonal 2 X'] / z_ratio,
                   'zone2_ph_e_prireach_r' : settings['Phase Earth Polygonal 2 R'],
                   'zone2_ph_e_secreach_r' : settings['Phase Earth Polygonal 2 R'] / z_ratio,
                   'zone2_relayangle' : settings['Phase Phase Polygonal 2 Relay Angle'],
                   'zone3_inservice' : 0 if settings['Phase Phase Polygonal 3 Out service'] else 1,
                   'zone3_trippingdirection' : '1',
                   'zone3_ph_ph_prireach_x' : settings['Phase Phase Polygonal 3 X'],
                   'zone3_ph_ph_secreach_x' : settings['Phase Phase Polygonal 3 X'] / z_ratio,
                   'zone3_ph_ph_prireach_r' : settings['Phase Phase Polygonal 3 R'],
                   'zone3_ph_ph_secreach_r' : settings['Phase Phase Polygonal 3 R'] / z_ratio,
                   'zone3_ph_e_prireach_x' : settings['Phase Earth Polygonal 3 X'],
                   'zone3_ph_e_secreach_x' : settings['Phase Earth Polygonal 3 X'] / z_ratio,
                   'zone3_ph_e_prireach_r' : settings['Phase Earth Polygonal 3 R'],
                   'zone3_ph_e_secreach_r' : settings['Phase Earth Polygonal 3 R'] / z_ratio,
                   'zone3_relayangle' : settings['Phase Phase Polygonal 3 Relay Angle'],
                   'zone3_tdelay' : settings['Phase Phase Polygonal 3 delay'],
                   'zone4_inservice' : 0 if settings['Phase Phase Polygonal 4 Out service'] else 1,
                   'zone4_trippingdirection' : '1',
                   'zone4_ph_ph_prireach_x' : settings['Phase Phase Polygonal 4 X'],
                   'zone4_ph_ph_secreach_x' : settings['Phase Phase Polygonal 4 X'] / z_ratio,
                   'zone4_ph_ph_prireach_r' : settings['Phase Phase Polygonal 4 R'],
                   'zone4_ph_ph_secreach_r' : settings['Phase Phase Polygonal 4 R'] / z_ratio,
                   'zone4_ph_e_prireach_x' : settings['Phase Earth Polygonal 4 X'],
                   'zone4_ph_e_secreach_x' : settings['Phase Earth Polygonal 4 X'] / z_ratio,
                   'zone4_ph_e_prireach_r' : settings['Phase Earth Polygonal 4 R'],
                   'zone4_ph_e_secreach_r' : settings['Phase Earth Polygonal 4 R'] / z_ratio,
                   'zone4_relayangle' : settings['Phase Phase Polygonal 4 Relay Angle'],
                   'zone4_tdelay' : settings['Phase Phase Polygonal 4 delay'],
                   'k0' : settings['k0'],
                   'phik0' : settings['k0 Angle']}, \
                   relay_name=relay_name, settings=settings, db=db)

    def save_line_distance_mho(self, relay_name, settings, db, ct_ratio, vt_ratio):
        '''
        function saving in the DB the mho distance settings protecting a line
        '''
        z_ratio = vt_ratio / ct_ratio
        self._save_into_DB(table_name="pt_line_distance", \
                     setting_table={\
                   'zone1_inservice' : 0 if settings['Phase Phase Mho 1 Out service'] else 1,
                   'zone1_trippingdirection' : 1, \
                   'zone1_ph_ph_prireach_z' : settings['Phase Phase Mho 1 Replica Impedance'] * z_ratio,
                   'zone1_ph_ph_secreach_z' : settings['Phase Phase Mho 1 Replica Impedance'],
                   'zone1_ph_e_prireach_z' : settings['Phase Earth Mho 1 Replica Impedance'] * z_ratio,
                   'zone1_ph_e_secreach_z' : settings['Phase Earth Mho 1 Replica Impedance'],
                   'zone1_relayangle' : settings['Phase Phase Mho 1 Relay Angle'],
                   'zone1_tdelay' : settings['Phase Phase Mho 1 Delay'],
                   'zone2_inservice' : 0 if settings['Phase Phase Mho 2 Out service'] else 1,
                   'zone2_trippingdirection' : '1',
                   'zone2_ph_ph_prireach_z' : settings['Phase Phase Mho 2 Replica Impedance'] * z_ratio,
                   'zone2_ph_ph_secreach_z' : settings['Phase Phase Mho 2 Replica Impedance'],
                   'zone2_ph_e_prireach_z' : settings['Phase Earth Mho 2 Replica Impedance'] * z_ratio,
                   'zone2_ph_e_secreach_z' : settings['Phase Earth Mho 2 Replica Impedance'],
                   'zone2_relayangle' : settings['Phase Phase Mho 2 Relay Angle'],
                   'zone2_tdelay' : settings['Phase Phase Mho 2 Delay'],
                   'zone3_inservice' : 0 if settings['Phase Phase Mho 3 Out service'] else 1,
                   'zone3_trippingdirection' : '1',
                   'zone3_ph_ph_prireach_z' : settings['Phase Phase Mho 3 Replica Impedance'] * z_ratio,
                   'zone3_ph_ph_secreach_z' : settings['Phase Phase Mho 3 Replica Impedance'],
                   'zone3_ph_e_prireach_z' : settings['Phase Earth Mho 3 Replica Impedance'] * z_ratio,
                   'zone3_ph_e_secreach_z' : settings['Phase Earth Mho 3 Replica Impedance'],
                   'zone3_relayangle' : settings['Phase Phase Mho 3 Relay Angle'],
                   'zone3_tdelay' : settings['Phase Phase Mho 3 Delay'],
                   'zone4_inservice' : 0 if settings['Phase Phase Mho 4 Out service'] else 1,
                   'zone4_trippingdirection' : '1',
                   'zone4_ph_ph_prireach_z' : settings['Phase Phase Mho 4 Replica Impedance'] * z_ratio,
                   'zone4_ph_ph_secreach_z' : settings['Phase Phase Mho 4 Replica Impedance'],
                   'zone4_ph_e_prireach_z' : settings['Phase Earth Mho 4 Replica Impedance'] * z_ratio,
                   'zone4_ph_e_secreach_z' : settings['Phase Earth Mho 4 Replica Impedance'],
                   'zone4_relayangle' : settings['Phase Phase Mho 4 Relay Angle'],
                   'zone4_tdelay' : settings['Phase Phase Mho 4 Delay'],
                   'k0' : settings['k0'],
                   'phik0' : settings['k0 Angle']}, \
                   relay_name=relay_name, settings=settings, db=db)

    def save_line_overcurrent(self, relay_name, settings, db, ct_ratio, vt_ratio):
        '''
        function saving in the DB the overcurrent settings protecting a line
        '''
        self._save_into_DB(table_name="pt_line_phase_overcurrent", \
                     setting_table=\
                    {'i_stg1_inservice' :  0 if settings['Enable #1'] else 1,
                   'i_stg1_trippingdirection' : '1',
                   'i_stg1_characteristictype' : self.interface.get_name_of(\
                                                settings['Characteristic #1']),
                   'i_stg1_pickupprimary' : settings['Current #1'] * ct_ratio,
                   'i_stg1_pickupsecondary' : settings['Current #1'],
                   'timedial' : settings['Time #1'],
                   'i_stg2_inservice' :  0 if settings['Enable #2'] else 1,
                   'i_stg2_trippingdirection' : '1',
                   'i_stg2_characteristictype' : 'Definite time',
                   'i_stg2_pickupprimary' : settings['Current #2'] * ct_ratio,
                   'i_stg2_pickupsecondary' : settings['Current #2'],
                   'trippingtimedelay' : settings['Time #2']}, \
                   relay_name=relay_name, settings=settings, db=db)

    def save_shunt_phase_overcurrent(self, relay_name, settings, db, ct_ratio, vt_ratio):
        '''
        function saving in the DB the phase overcurrent settings protecting a shunt
        '''
        self._save_into_DB(table_name="pt_sh_phase_overcurrent", \
                     setting_table=\
                    {'i_stg1_inservice' :  0 if settings['Enable #1'] else 1,
                   'i_stg1_trippingdirection' : '1',
                   'i_stg1_characteristictype' : self.interface.get_name_of(\
                                                settings['Characteristic #1']),
                   'i_stg1_pickupprimary' : settings['Current #1'] * ct_ratio,
                   'i_stg1_pickupsecondary' : settings['Current #1'],
                   'timedial' : settings['Time #1'],
                   'i_stg2_inservice' :  0 if settings['Enable #2'] else 1,
                   'i_stg2_trippingdirection' : '1',
                   'i_stg2_characteristictype' : settings['Characteristic #1'],
                   'i_stg2_pickupprimary' : settings['Current #2'] * ct_ratio,
                   'i_stg2_pickupsecondary' : settings['Current #2'],
                   'trippingtimedelay' : settings['Time #2']}, \
                   relay_name=relay_name, settings=settings, db=db, \
                   from_bus_number_available=False)

    def save_shunt_neutral_overcurrent(self, relay_name, settings, db, ct_ratio, vt_ratio):
        '''
        function saving in the DB the neutral overcurrent settings protecting
        a shunt
        '''
        self._save_into_DB(table_name="pt_sh_neutral_Overcurrent", \
                     setting_table=\
                    {'ig_stg1_inservice' :  0 if settings['Enable #1'] else 1,
                   'ig_stg1_trippingdirection' : '1',
                   'ig_stg1_characteristictype' : self.interface.get_name_of(\
                                                settings['Characteristic #1']),
                   'ig_stg1_pickupprimary' : settings['Current #1'] * ct_ratio,
                   'ig_stg1_pickupsecondary' : settings['Current #1'],
                   'timedial' : settings['Time #1'],
                   'ig_stg2_inservice' :  0 if settings['Enable #2'] else 1,
                   'ig_stg2_trippingdirection' : '1',
                   'ig_stg2_characteristictype' : 'Definite time',
                   'ig_stg2_pickupprimary' : settings['Current #2'] * ct_ratio,
                   'ig_stg2_pickupsecondary' : settings['Current #2'],
                   'trippingtimedelay' : settings['Time #2']}, \
                   relay_name=relay_name, settings=settings, db=db, \
                   from_bus_number_available=False)

    def save_generator_underimpdance_poly(self, relay_name, settings, db, ct_ratio, vt_ratio):
        '''
        function saving in the DB the under Z distance settings protecting a 
        generator
        '''
        z_ratio = vt_ratio / ct_ratio
        self._save_into_DB(table_name="pt_gen_underimpedance", \
                     setting_table={\
                   'zone1_inservice' : 0 if settings['Phase Phase Polygonal 1 Out service'] else 1,
                   'zone_stg1_prireach_x' : settings['Phase Phase Polygonal 1 X'],
                   'zone_stg1_secreach_x' : settings['Phase Phase Polygonal 1 X'] / z_ratio,
                   'zone_stg1_prireach_r' : settings['Phase Phase Polygonal 1 R'],
                   'zone_stg1_secreach_r' : settings['Phase Phase Polygonal 1 R'] / z_ratio,
                   'zone_stg2_prireach_x' : settings['Phase Earth Polygonal 1 X'],
                   'zone_stg2_secreach_x' : settings['Phase Earth Polygonal 1 X'] / z_ratio,
                   'zone_stg2_prireach_r' : settings['Phase Earth Polygonal 1 R'],
                   'zone_stg2_secreach_r' : settings['Phase Earth Polygonal 1 R'] / z_ratio,
                   'zone_stg1_tdelay' : settings['Phase Phase Polygonal 1 delay'],
                   'zone_stg2_tdelay' : settings['Phase Phase Polygonal 2 delay'],
                   'zone_angle' : settings['Phase Phase Polygonal 1 Relay Angle']},
                   relay_name=relay_name, settings=settings, db=db, \
                   from_bus_number_available=False)

    def _save_into_DB(self, table_name, setting_table, relay_name, \
                     settings, db, from_bus_number_available=True):
        '''
        base function saving the data for a relay in the DB
        '''
        if self.output_detail >= OutputDetail.NORMAL:
            self.interface.print("Saving relay: " + relay_name)
        table = db.get_table(table_name)
        name_items = relay_name.split('_')
        if from_bus_number_available:
            update_query = table.update().where(\
                        table.c.from_bus_number == name_items[1] and\
                        table.c.to_bus_number == name_items[2] and\
                        table.c.id == name_items[3] and\
                        table.c.protectionbus == name_items[4])\
                        .values(setting_table)
        else:
            update_query = table.update().where(\
                        table.c.busnumber == name_items[2] and\
                        table.c.id == name_items[3] and\
                        table.c.protectionbus == name_items[4])\
                        .values(setting_table)
        db.execute(update_query)

################################################################################
################################################################################
################################################################################

    def save_differential_relay_settings_from_sheet_to_DB(self):
        '''
        function saving differential settings from an excell sheet page to the DB
        '''
        self.interface.print("\n\t\t *** Loading all differential relay settings into the DB ***")
        # create the db connection
        db = None
        try:
            db = MYSQLInterface(self.interface)
        except:
            self.interface.print("ERROR: Cannot connect to the MySQL database")
            return
        # each record is a set of cells in a file and in a sheeet
        records = []
        records.append(self.RelaySheetRecord(file=\
            "D:\\Setting_Line_differential_protection.xlsx",
                         sheet_name='Setting Line Dif. relay',
                         setting_table={\
                                'relay_name'   : ('A', 5), \
                                'idmin'        : ('C', 5), \
                                'point1'       : ('D', 5), \
                                'point2'       : ('E', 5), \
                                'slope1'       : ('F', 5), \
                                'slope2'       : ('G', 5), \
                                'idnorestrain' : ('H', 5)}, \
                                 multiple=True))

        # here some other records are added in the list if needed...
        # ....
        # call the function which processes all records
        self._transfer_sheets_to_DB(records, db)
        self.interface.print("\n\t\t *** All differential relay settings have been saved into the DB *** ")

    def _transfer_sheets_to_DB(self, records, db):
        '''
        functions reading the sheet data and transferring then in to the BD
        '''
        table_name = 'pt_line_differential_protection'
        exit_loops = False
        settings = None
        for record in records:
            wb = load_workbook(record.file, data_only=True)
            # check if a file has been found
            if wb:
                sheet = wb[record.sheet_name]
                for row in range(10000):
                    setting_table = {}
                    # exit after the first record if the record set is not multiple
                    if row > 0 and record.multiple == False:
                        break
                    for index, (db_setting_name, position) in\
                         enumerate(record.setting_table.items()):
                        if index == 0:
                            relay_name = str(sheet[position[0] + \
                                             str(position[1] + row)].value)
                            # there is no relay name...leave the loops
                            if relay_name == 'None':
                                exit_loops = True
                                break
                        else:
                            value = str(sheet[position[0] + str(position[1] + row)].value)
                            setting_table.update({db_setting_name : value})
                    # exit if no relay has been found
                    if exit_loops:
                        break
                    self._save_into_DB(table_name, setting_table, relay_name, \
                         settings, db, from_bus_number_available=True)
            else:
                self.interface.print("ERROR: Cannot load " + record.file)

################################################################################
################################################################################
################################################################################
    def run_simulation_for_multiple_study_cases(self):
        '''
        run a simulation along all available study cases
        '''
        self.run_multiple_study_cases(\
                    self.run_rms_simulation_with_different_fault_removal_times)

################################################################################
################################################################################
################################################################################

    def run_multiple_study_cases(self, payload_function, window=None):
        '''
        run a process along all available study cases running for each of them
        the given payload function which does something
        '''
        # Print some info in the PF output window
        if self.output_detail.value >= OutputDetail.NORMAL.value:
            self.interface.print("*********************************************************")
            self.interface.print("    Time Distance Creator Tool (Beta) 0.1")
            self.interface.print("            Run multiple study cases           ")
            self.interface.print("*********************************************************\n")
        # self.interface.set_echo_off()
        study_cases = self.interface.get_study_cases()
        for study_case in study_cases:
            self.interface.activate_study_case(study_case)
            if self.output_detail >= OutputDetail.NORMAL:
                self.interface.set_echo_on()
                self.interface.print(self.interface.get_name_of(study_case) + \
                                  " study case has been activated")
                self.interface.set_echo_off()
            # defintion of the function which performs the required operations
            # like running a simulation, a LDF etc
            payload_function(study_case, window)
        self.interface.set_echo_on()
        if self.output_detail.value >= OutputDetail.NORMAL.value:
            self.interface.print("***     Run multiple study cases  completed     ***")

################################################################################

    def run_rms_simulation_with_different_fault_removal_times(self, study_case, window):
        '''
        function performing operations for generator critical time calculation
         YudTet Eng projects
        '''
        # no operational messages will be printed....
        self.interface.set_echo_off()
        study_case_name = self.interface.get_name_of(study_case)
        # dictionary with the result files and the relevant removal times
        if len(self.limit_times) > 0:
            operation_parameters = {"150 ms": 0.150,
                      "max stable": self.limit_times[study_case_name]["max stable"],
                      "min unstable": self.limit_times[study_case_name]["min unstable"]}
        else:
            operation_parameters = {"150 ms": 0.150,
                      "max stable": 0.165,
                      "min unstable": 0.175}  if '3B' in study_case_name else\
                        {"Results": 0.000} if '5000' in study_case_name else\
                        {"Results": 0.000} if '123' in study_case_name or\
                                            '23' in study_case_name else\
                            {"150 ms": 0.150,
                             "max stable": 0.160,
                             "min unstable": 0.170}
        # event name (to customize)
        if '123' in study_case_name or\
           '5000' in study_case_name or\
            '23' in study_case_name:
            self.interface.print("    ZEELIM Projet: ")
            event_name = 'Short-Circuit Event'
        else:
            event_name = 'Short-Circuit Event(1)'
        events = self.interface.get_study_case_events(study_case, event_name)
        simulation_init_objects = self.interface.get_simulation_inits_of(study_case)
        # check that the simulation init object and the events are availablle
        if simulation_init_objects and events:
            for operation_key, operation_value in operation_parameters.items():
                results = self.interface.get_study_case_results(study_case, operation_key)
                if results:
                    if self.output_detail.value >= OutputDetail.NORMAL.value:
                        self.interface.set_echo_on()
                        self.interface.print("    Result file: " + \
                        str(operation_key) + "  Value: " + str(operation_value))
                        self.interface.set_echo_off()
                    # set the right result object in the simulation init
                    self.interface.set_attribute(\
                        simulation_init_objects[0], 'p_resvar', results[0])
                    # set the fault clearing time in the events
                    self.interface.set_attribute(events[0], 'time', operation_value)
                    # here init and run the simulation
                    self.interface.set_echo_off()
                    simulation_init_objects[0].Execute()
                    simulation_objects = self.interface.\
                                                get_simulation_objects_of(study_case)
                    if simulation_objects:
                        if self.output_detail.value >= OutputDetail.NORMAL.value:
                            self.interface.set_echo_on()
                            self.interface.print("    running simulation")
                            self.interface.set_echo_off()
                        # run the simulation
                        simulation_objects[0].Execute()
                        # get the results data for the rotor angle
                        variable_name = "c:firel"
                        max_value = self.interface.get_max_of_element_variable_results(\
                                                    results[0], variable_name)
                        if '123' not in study_case_name and\
                            '23' not in study_case_name and\
                             '5000' not in study_case_name:
                            self.interface.print("\t   The max rotor angle is: " + \
                                        str(round(max_value, 2)))
                    else:
                        self.interface.set_echo_on()
                        self.interface.print("ERROR: no sim object available for " + \
                        self.interface.get_name_of(study_case) + " study case")
                else:
                    self.interface.set_echo_on()
                    self.interface.print("ERROR: no " + str(operation_key) \
                                         +" result available for" + \
                        self.interface.get_name_of(study_case) + " study case")
                    self.interface.set_echo_off()
        else:
            self.interface.set_echo_on()
            if simulation_init_objects:
                self.interface.print("ERROR: no events available for " + \
                                            study_case_name + " study case")
            else:
                self.interface.print("ERROR: no simulation init available for " + \
                                            study_case_name + " study case")
            self.interface.set_echo_off()

################################################################################
################################################################################
################################################################################
    def calculate_all_critical_times(self, window):
        '''
        calculate the generator critical time for all available study cases
        '''
        self.limit_times = {}
        self.run_multiple_study_cases(self.calculate_generator_critical_time, window)

    def calculate_generator_critical_time(self, study_case, window):
        '''
        calculate the generator critical time running many simulations with 
        different fault removal times
        '''
        first_attempt_removal_time = 0.15
        generator_name = window.GetSettings()['TargetDevice'].rstrip()
        # the search step
        time_step = 0.005
        # event name (to customize)
        event_name = 'Short-Circuit Event(1)'
        events = self.interface.get_study_case_events(study_case, event_name)
        simulation_init_objects = self.interface.get_simulation_inits_of(study_case)
        removal_time = first_attempt_removal_time
        generator = self.interface.get_element_by_name(generator_name)
        # here the rotro angle limit value to declare out of step
        max_safety_angle = 150
        # flag to kn ow if we are already over the safety angle
        safety_angle_reached = False
        while generator and events:
            # set the fault clearing time in the events
            self.interface.set_attribute(events[0], 'time', removal_time)
            # here init and run the simulation
            self.interface.set_echo_off()
            simulation_init_objects[0].Execute()
            simulation_objects = self.interface.\
                                        get_simulation_objects_of(study_case)
            if simulation_objects:
                if self.output_detail.value >= OutputDetail.VERBOSEDEBUG.value:
                    self.interface.set_echo_on()
                    self.interface.print("    running simulation with " + \
                                         str(round(removal_time, 3)) + " s removal time")
                    self.interface.set_echo_off()
                # run the simulation
                simulation_objects[0].Execute()
                is_out_of_step = self.interface.get_attribute(generator[0], "s:outofstep")
                if self.output_detail.value >= OutputDetail.VERBOSEDEBUG.value:
                    self.interface.set_echo_on()
                    if is_out_of_step == True:
                        self.interface.print("    Out of step condition detected! ")
                    else:
                        self.interface.print("    " + generator_name + \
                                                            " remains stable")
                    self.interface.set_echo_off()

                # get the results data for the rotor angle
                results = self.interface.get_attribute(\
                                        simulation_init_objects[0], 'p_resvar')
                if results:
                    variable_name = "c:firel"
                    max_angle_value = self.interface.get_max_of_element_variable_results(\
                                            results, variable_name)
                    if safety_angle_reached == False:
                        if max_angle_value > max_safety_angle:
                            if removal_time > first_attempt_removal_time:
                                self.interface.print("   Critical time has been found\
                                 with safety angle = " + str(max_safety_angle))
                                self.interface.print("   Min unstable time " + \
                                                     str(round(removal_time, 3)) + " s")
                                self.interface.print("   Max stable time " + \
                                        str(round(removal_time - time_step, 3)) + " s")
                            else:
                                if removal_time < first_attempt_removal_time:
                                    self.interface.print("   Critical time has been found\
                                 with safety angle = " + str(max_safety_angle))
                                    self.interface.print("   Min unstable time " + \
                                            str(round(removal_time + time_step, 3)) + " s")
                                    self.interface.print("   Max stable time " + \
                                                        str(round(removal_time, 3)) + " s")
                            self.interface.print("\t   The max rotor angle is: " + \
                                str(round(max_angle_value, 2)))
                            safety_angle_reached = True

                if is_out_of_step == True:
                    if removal_time > first_attempt_removal_time:
                        self.interface.print("   Critical time has been found")
                        self.interface.print("   Min unstable time " + \
                                             str(round(removal_time, 3)) + " s")
                        self.interface.print("   Max stable time " + \
                                str(round(removal_time - time_step, 3)) + " s")

                        self.limit_times.update({self.interface.\
                                                 get_name_of(study_case):
                                {"max stable": round(removal_time - time_step, 3),
                                        "min unstable": round(removal_time, 3)}})
                        break
                else:
                    if removal_time < first_attempt_removal_time:
                        self.interface.print("   Critical time has been found")
                        self.interface.print("   Min unstable time " + \
                                str(round(removal_time + time_step, 3)) + " s")
                        self.interface.print("   Max stable time " + \
                                            str(round(removal_time, 3)) + " s")
                        self.limit_times.update({self.interface.\
                                                 get_name_of(study_case):
                                {"max stable": round(removal_time, 3),
                                "min unstable": round(removal_time + time_step, 3)}})
                        break

                removal_time += -time_step if is_out_of_step == True else time_step
            else:
                self.interface.set_echo_on()
                self.interface.print("ERROR: no sim object available for " + \
                self.interface.get_name_of(study_case) + " study case")

        self.interface.set_echo_on()
        if not generator:
            self.interface.print("ERROR: the " + generator_name + \
                                 " generator is not available")
        if not events:
            self.interface.print("ERROR: no events available for " + \
                        self.interface.get_name_of(study_case) + " study case")
        self.interface.set_echo_off()

################################################################################
################################################################################
################################################################################

    def move_all_pictures_to_word(self, window, input_settings=None):
        '''
        function getting all pictures available in all study cases, saving them
        in the Pictures
        '''
        # just for test
        # self.interface.rebuild_pf()
        # return
        # just check that the settings are ok
        settings = self.initialize(window, input_settings)
        if self.is_dialog_setting_ok(window, settings) == False:
            return 1
        if self.output_detail >= OutputDetail.NORMAL:
                self.interface.print("Saving all pictures ...")

        # create the Pictures directory if not present in the path where the
        # result file is
        picture_dir_path = os.path.split(window.results_file_name.GetValue())[0] + \
                                                 '\\Pictures'
        if not os.path.exists(picture_dir_path):
            os.mkdir(picture_dir_path)

        # just for  test
#         self.add_wmfs_in_word(word_file_name = window.results_file_name.GetValue(),\
#                              wmf_path = picture_dir_path)

        # go throw all study cases
        study_cases = self.interface.get_study_cases()
        for study_case in study_cases:
            # activate the study case
            self.interface.set_echo_off()
            self.interface.activate_study_case(study_case)
            if self.output_detail >= OutputDetail.NORMAL:
                self.interface.set_echo_on()
                self.interface.print(self.interface.get_name_of(study_case) + \
                                  " study case has been activated")
                self.interface.set_echo_off()
            diagram_path = picture_dir_path + '\\' + \
                                        self.interface.get_name_of(study_case)
            # save the diagram
            self.save_diagrams_as_wmf(window, path=diagram_path, create_copy=True)
        self.interface.set_echo_on()
        if self.output_detail >= OutputDetail.NORMAL:
                self.interface.print("All pictures have been saved!")
                self.interface.print("Trasfering pictures to word....")
        # here process the word file to replace the pictures
        self.add_wmfs_in_word(word_file_name=window.results_file_name.GetValue(), \
                             wmf_path=picture_dir_path)
        if self.output_detail >= OutputDetail.NORMAL:
            self.interface.print("All pictures have been Transfered!")
        self.interface.print(" ***   Task completed    ***")

################################################################################
################################################################################
################################################################################

    def set_line_overcurrent(self, relay , trafo):
        '''
        set the given overcurrent relay to protect the given trafo
        '''
        pf_relay = PowerFactoryOvercurrentRelayInterface(self.interface, relay)
        relay_settings = pf_relay.read_settings()
        si_curve = pf_relay.get_curve_object('IEC Class A (Standard Inverse)')
        if si_curve:
            relay_settings['Characteristic #1'] = si_curve
        pf_relay.write_settings(relay_settings)

    def set_trafo_overcurrent(self, relay , trafo):
        '''
        set the given overcurrent relay to protect the given trafo
        '''
        pf_relay = PowerFactoryOvercurrentRelayInterface(self.interface, relay)
        relay_settings = pf_relay.read_settings()
        trafo_in = self.interface.get_transformer_rated_i()
        pf_relay.write_settings(relay_settings)

    def set_shunt_phase_overcurrent(self, relay, shunt):
        '''
        set the given phase overcurrent relay to protect the given shunt
        '''
        if self.output_detail >= OutputDetail.NORMAL:
            self.interface.print("     Setting " + self.interface.get_name_of(relay)\
                                 +" phase overcurrent shunt relay (shunt: " + \
                                 self.interface.get_name_of(shunt) + ")")
        ct_list = self.interface.get_cubicle_CT_of(\
                                    self.interface.get_relay_cubicle_of(relay))
        ct_ratio = self.interface.get_ct_ratio(ct_list[0]) if ct_list else 1
        pf_relay = PowerFactoryOvercurrentRelayInterface(self.interface, relay)
        relay_settings = pf_relay.read_settings()
        shunt_in = self.interface.get_shunt_rated_i(shunt)
        relay_settings['Enable #1'] = self.enabled
        # trip thershold at 130% shunt In
        relay_settings['Current #1'] = round(shunt_in * 1.3 / ct_ratio, 2)
        # time dial calculated at 2.5/ trip threshold 2 seconds
        relay_settings['Time #1'] = round((((2.5 / 1.3) ** 0.02 - 1) * 2) / 0.14, 2)
        si_curve = pf_relay.get_curve_object('IEC Class A (Standard Inverse)')
        if si_curve:
            relay_settings['Characteristic #1'] = si_curve
        relay_settings['Enable #2'] = self.enabled
        relay_settings['Current #2'] = round(shunt_in * 4 / ct_ratio, 2)
        relay_settings['Time #2'] = 0.1
        relay_settings['Enable #3'] = self.disabled
        relay_settings['Enable #4'] = self.disabled
        pf_relay.write_settings(relay_settings)

    def set_shunt_neutral_overcurrent(self, relay, shunt):
        '''
        set the given neutral overcurrent relay to protect the given shunt
        '''
        if self.output_detail >= OutputDetail.NORMAL:
            self.interface.print("     Setting " + self.interface.get_name_of(relay)\
                                 +" neutral overcurrent shunt relay (shunt: " + \
                                 self.interface.get_name_of(shunt) + ")")
        ct_list = self.interface.get_cubicle_CT_of(\
                                    self.interface.get_relay_cubicle_of(relay))
        ct_ratio = self.interface.get_ct_ratio(ct_list[0]) if ct_list else 1
        pf_relay = PowerFactoryNeutralOvercurrentRelayInterface(self.interface, relay)
        relay_settings = pf_relay.read_settings()
        shunt_in = self.interface.get_shunt_rated_i(shunt)
        relay_settings['Enable #1'] = self.enabled
        # trip thershold at 10% shunt In
        relay_settings['Current #1'] = round(shunt_in * 0.1 / ct_ratio, 2)
        # time dial calculated at 1.2/ trip thershold 1.2 seconds
        relay_settings['Time #1'] = round((((1.2 / 0.1) ** 0.02 - 1) * 1.2) / 0.14, 2)
        si_curve = pf_relay.get_curve_object('IEC Class A (Standard Inverse)')
        if si_curve:
            relay_settings['Characteristic #1'] = si_curve
        relay_settings['Enable #2'] = self.enabled
        relay_settings['Current #2'] = round(shunt_in * 4 / ct_ratio, 2)
        relay_settings['Time #2'] = 0.1
        relay_settings['Enable #3'] = self.disabled
        pf_relay.write_settings(relay_settings)

    def set_generator_underz_polygonal(self, relay, generator):
        '''
        set the given under impedance relay to protect the given generator 
        '''
        z1_reach = 0.7
        z2_reach = 1.2
        if self.output_detail >= OutputDetail.NORMAL:
            self.interface.print("\nCalculating setting for " + \
                                self.interface.get_name_of(relay))

        # create relay interface
        pf_relay = PowerFactoryPolygonalDistanceRelayInterface(self.interface, relay)
        relay_settings = pf_relay.read_settings()
        relay_settings['Phase Phase Polygonal 1 Out service'] = 1

        branch = self.grid.get_branch_of(generator)
        z = 0
        # if the branch is a transformer branch get the other side
        if branch:
            if branch.is_transformer_branch():
                z = self.interface.get_transformer_z(branch.transformer)
            else:  # otherwise multiple generators connected to the same busbar
                bus = self.interface.get_branch_busses_of(generator, \
                                                    self.interface.Side.Side_1)
                branches = self.grid.get_branch_of(bus)
                if branches:
                    for branch in branches:
                        if branch.is_transformer_branch():
                            z = self.interface.get_transformer_z(\
                                                        branch.transformer)
                            break
        z1_z_reach = z1_reach * z
        z1_z_reach_real = z1_z_reach.real if z1_z_reach.real > 0 else z1_z_reach.imag
        relay_settings['Phase Phase Polygonal 1 X'] = round(z1_z_reach.imag, 2)
        relay_settings['Phase Phase Polygonal 1 R'] = round(z1_z_reach_real, 2)
        pf_relay.write_settings(relay_settings)
        if self.output_detail >= OutputDetail.NORMAL:
            self.interface.print("     Set at " + str(round(z1_z_reach_real, 2)) + \
                                 " , " + str(round(z1_z_reach.imag, 2)))

    def set_line_distance_polygonal(self, relay, line):
        '''
        procedure to calculate the polygonal distance relay settings
        '''
        from math import sin, cos, pi
        from cmath import phase
        mult1 = 0.8

        if self.output_detail >= OutputDetail.NORMAL:
            self.interface.print("\nCalculating setting for " + \
                                self.interface.get_name_of(relay))
        pf_relay = PowerFactoryPolygonalDistanceRelayInterface(self.interface, relay)
        relay_settings = pf_relay.read_settings()

        # get the impedances for z1 and z2
        z1_z_reach, z2_z_reach = self.get_z1z2_impedance(relay, line)

        branch = self.grid.get_branch_of(line)
        if branch:
            branch_z = branch.get_branch_z()
            branch_angle_rad = phase(branch_z)
        else:
            branch_angle_rad = 0
        branch_angle = branch_angle_rad * 180 / pi

        relay_settings['Phase Phase Polygonal 1 X'] = z1_z_reach.imag
        relay_settings['Phase Phase Polygonal 1 R'] = z1_z_reach.real
        relay_settings['Phase Phase Polygonal 1 Relay Angle'] = branch_angle
        relay_settings['Phase Phase Polygonal 1 delay'] = 0
        relay_settings['Phase Phase Polygonal 2 X'] = z2_z_reach.imag
        relay_settings['Phase Phase Polygonal 2 R'] = z2_z_reach.real
        relay_settings['Phase Phase Polygonal 2 Relay Angle'] = branch_angle
        relay_settings['Phase Phase Polygonal 2 delay'] = 0.4

        if branch:
            z3_z_reach = self.get_z3_impedance(relay, line)
            if z3_z_reach > 0.0001:
                relay_settings['Phase Phase Polygonal 3 X'] = z3_z_reach * sin(branch_angle_rad)
                relay_settings['Phase Phase Polygonal 3 R'] = z3_z_reach * cos(branch_angle_rad)
                relay_settings['Phase Phase Polygonal 3 Relay Angle'] = branch_angle
            relay_settings['Phase Phase Polygonal 3 delay'] = 0.8
            self.interface.set_echo_off()
            pf_relay.write_settings(relay_settings)
            self.interface.set_echo_on()
            if self.output_detail >= OutputDetail.NORMAL:
                self.interface.print("   '" + self.interface.get_name_of(line) + "' Z = " + \
                            str(round(abs(z1_z_reach / mult1), 2)) + " ohm  , angle = " + str(round(branch_angle, 2)) + " deg")
                self.interface.print("   Z1 reach = " + str(round(abs(z1_z_reach), 2)) + "(" + \
                str(round(z1_z_reach.real, 3)) + "," + str(round(z1_z_reach.imag, 3))
                                      +") pri ohm")
                self.interface.print("   Z2 reach = " + str(round(abs(z2_z_reach), 2)) + "(" + \
                str(round(z2_z_reach.real, 3)) + "," + str(round(z2_z_reach.imag, 3)) + ") pri ohm")

                self.interface.print("   Z3 reach = " + str(round(abs(z3_z_reach), 2)) + "(" + \
                str(round(z3_z_reach.real, 3)) + "," + str(round(z3_z_reach.imag, 3)) + ") pri ohm")

    def set_line_distance_mho(self, relay, line):
        '''
        procedure to calculate the mho distance relay settings
        '''
        from math import pi
        from cmath import phase
        if self.output_detail >= OutputDetail.NORMAL:
            self.interface.print("\nCalculating setting for " + \
                                self.interface.get_name_of(relay))
        pf_relay = PowerFactoryMhoDistanceRelayInterface(self.interface, relay)
        relay_settings = pf_relay.read_settings()

        # get the impedances for z1 and z2
        z1_z_reach, z2_z_reach = self.get_z1z2_impedance(relay, line)

        branch = self.grid.get_branch_of(line)
        if branch != None:
            branch_z = branch.get_branch_z()
            branch_angle_rad = phase(branch_z)
        else:
            branch_angle_rad = 0
        branch_angle = branch_angle_rad * 180 / pi

        relay_settings['Phase Phase Mho 1 Replica Impedance'] = abs(z1_z_reach)
        relay_settings['Phase Phase Mho 1 Relay Angle'] = branch_angle
        relay_settings['Phase Phase Mho 1 Delay'] = 0

        relay_settings['Phase Phase Mho 2 Replica Impedance'] = abs(z2_z_reach)
        relay_settings['Phase Phase Mho 2 Relay Angle'] = branch_angle
        relay_settings['Phase Phase Mho 2 Delay'] = 0.4

        if branch:
            z3_z_reach = self.get_z3_impedance(relay, line)
            if z3_z_reach > 0.0001:
                relay_settings['Phase Phase Mho 3 Replica Impedance'] = abs(z3_z_reach)
                relay_settings['Phase Phase Mho 3 Relay Angle'] = branch_angle
            relay_settings['Phase Phase Mho 2 Delay'] = 0.8
            if self.output_detail >= OutputDetail.NORMAL:
                self.interface.print("          Writing settings in the relay...")
            self.interface.set_echo_off()
            pf_relay.write_settings(relay_settings)
            self.interface.set_echo_on()
            if self.output_detail >= OutputDetail.NORMAL:
                self.interface.print("   '" + self.interface.get_name_of(line) + "' Z = " + \
                            str(round(abs(branch_z), 2)) + " ohm  , angle = " + str(round(branch_angle, 2)) + " deg")
                self.interface.print("   Z1 reach = " + str(round(abs(z1_z_reach), 2)) + " pri ohm")
                self.interface.print("   Z2 reach = " + str(round(abs(z2_z_reach), 2)) + " pri ohm")
                self.interface.print("   Z3 reach = " + str(round(abs(z3_z_reach), 2)) + " pri ohm")

    def is_tap_bus(self, busbar):
        '''
        function checking if at least one relayis present inside the cubicles of the
        given busbar. If there is no relay the given busbar is declared as a tap
        '''
        cubicles = self.interface.get_bus_cubicles(busbar)
        if cubicles:
            for cubicle in cubicles:
                if len(self.interface.get_cubicle_relay_of(cubicle)) > 0:
                    return False
        return True

    def _get_remote_busses_of(self, bus_bar, branch, remote_busses, remote_elements):
        '''
        ricorsive function returning all remote bus bar connected and not stopping on
        the tap busbars
        bus_bar = the bs bar from which we start the search
        branch = the line from which are coming from (used to avoid to come back)
        remote_busses = a list of all remote busses which have been found
        remote_elements  = a list of all elements to go to remote busses which have been found
        '''
        connected_element_list = self.interface.get_bus_connections_of(bus_bar)
        if connected_element_list:
            for element in connected_element_list:
                # skip the line from which we are coming from
                element_branch = self.grid.get_branch_of(element)
                if element_branch == branch or element_branch == None:
                    continue
                remote_bus_bar_2nd_step = element_branch.get_other_busbar_of(bus_bar)
                # skip the element if it doesn't have a remote bus (so it's i.e.a load)
                # of if the remote bus has already been collected
                if remote_bus_bar_2nd_step == None or \
                remote_bus_bar_2nd_step in remote_busses:
                    continue
                remote_elements.append(element)
                remote_busses.append(remote_bus_bar_2nd_step)
                if self.is_tap_bus(remote_bus_bar_2nd_step):
                    new_remote_busses = self._get_remote_busses_of(\
                    remote_bus_bar_2nd_step, element_branch, remote_busses, remote_elements)
                    remote_busses += new_remote_busses
                    remote_busses = list(set(remote_busses))
        else:
            pass
        return remote_busses

    def get_z1z2_impedance(self, relay, line):
        '''
        function which calculates the Z1 and Z2 impedances taking care of tap busbars
        '''
        mult1 = 0.8
        mult2 = 1.2
        if self.output_detail >= OutputDetail.NORMAL:
            self.interface.print("    Calculating setting for Zone 1 and Zone 2 ")
        # get the branch at which the given line belongs
        branch = self.grid.get_branch_of(line)
        if branch == None:
            if self.output_detail >= OutputDetail.NORMAL:
                self.interface.print("    Error: no branch available!")
            return 0, 0
        # get just the branch z
        branch_z = branch.get_branch_z()
        # get the cubicle hosting the given relay
        cubicle = self.interface.get_relay_cubicle_of(relay)
        # get the relay polarizing element
        polarizing_element = self.interface.get_relay_polarizing_element(relay)
        if cubicle and polarizing_element:
            # set fault type and resistance
            self.interface.set_fault(self.fault_types['3PH'].type, 0)
            # get the relay busbar
            bus_bar = self.interface.get_cubicle_busbar_of(cubicle)
            z_max = 0
            remote_bus_bar = branch.get_last_busbar_from(bus_bar)
            # use a list for the remote busbar to take care of the taps
            remote_bus_bars = []
            remote_paths = []
            if remote_bus_bar and self.is_tap_bus(remote_bus_bar):
                remote_bus_bars = self._get_remote_busses_of(remote_bus_bar, branch, \
                                                            remote_bus_bars, remote_paths)
                if len(remote_bus_bars):
                    for remote_bus_bar in remote_bus_bars:
                        if remote_bus_bar:
                            # fault at the remote bus bar locations
                            self.interface.set_fault_position(remote_bus_bar, \
                                                              position=0.01, \
                                                              single_shc=True)
                            if self.output_detail >= OutputDetail.NORMAL:
                                self.interface.print("       SHC at " + self.interface.\
                                                get_name_of(remote_bus_bar))
                            # run shc
                            try:
                                self.interface.set_echo_off()
                                error = self.interface.run_shc()
                                self.interface.set_echo_on()
                                if error == 0:
                                    zline = self.interface.get_relay_secondary_z_measures_of(\
                                                                polarizing_element)
                                    if self.output_detail >= OutputDetail.NORMAL:
                                        self.interface.print("       Measured Z = " + \
                                                             str(zline) + " sec.Ohm")

                                    z_max = zline[0] * mult1 \
                                            if zline[0] * mult1 > z_max else z_max
                            except Exception as e:
                                pass
                if self.output_detail >= OutputDetail.NORMAL:
                    self.interface.print("          Max Measured Z limit = " + str(round(z_max, 2)) + \
                                          " sec.Ohm")
                z1_z_reach = z_max
                z2_z_reach = z_max * mult2 / mult1
            else:
                z1_z_reach = branch_z * mult1
                z2_z_reach = branch_z * mult2
        else:
            z1_z_reach = branch_z * mult1
            z2_z_reach = branch_z * mult2
        return z1_z_reach, z2_z_reach

    def get_z3_impedance(self, relay, line):
        '''
        calculate the Z3 impedance for the given relay applying 3 phase faults
        at the 2nd step remote bus bars and selecting the greatest impedance not 
        greater than the smallest trafo impedance
        '''
        mult = 1.2
        if self.output_detail >= OutputDetail.NORMAL:
            self.interface.print("    Calculating setting for Zone 3")
        z3_zscc_value = namedtuple("z3_zscc_value", \
        "line relay observation_cubicle fault_type fault_bus zscc_values relay_bus open_breaker trafo_shc")
        # get the cubicle hosting the given relay
        cubicle = self.interface.get_relay_cubicle_of(relay)
        # get the relay polarizing element
        polarizing_element = self.interface.get_relay_polarizing_element(relay)
        # get the Ct ratio
        ct_list = self.interface.get_cubicle_CT_of(cubicle)
        if ct_list:
            ct_ratio = self.interface.get_ct_ratio(ct_list[0])
        else:
            ct_ratio = 1
        # set the Ct direction to the branch
        # temporary code!
        if ct_list:
            self.interface.set_attribute(ct_list[0], 'ciorient', 0)
        # get the Vt ratio
        vt_list = self.interface.get_cubicle_VT_of(cubicle)
        if vt_list:
            vt_ratio = self.interface.get_vt_ratio(vt_list[0])
        else:
            vt_ratio = 1
        # return value
        z3_max = 0
        # z limit due to the transformers
        trafo_z_limit = 100000
        if cubicle and polarizing_element:
            # set fault type and resistance
            self.interface.set_fault(self.fault_types['3PH'].type, 0)
            # get the relay busbar
            bus_bar = self.interface.get_cubicle_busbar_of(cubicle)
            # get the given line branch
            branch = self.grid.get_branch_of(line)
            # get the 1st step remote busbar
            z3_zscc_data = []
            if branch and self.interface.is_out_of_service(line) == False:
                # use a list for the remote busbar to take care of the taps
                remote_bus_bars = []
                remote_elements = []
                remote_bus_bar = branch.get_last_busbar_from(bus_bar)

                if remote_bus_bar and self.is_tap_bus(remote_bus_bar):
                    remote_bus_bars = self._get_remote_busses_of(remote_bus_bar, branch, \
                                                        remote_bus_bars, remote_elements)
                else:
                    remote_bus_bars.append(remote_bus_bar)
                # list of all busses/lines already investigated
                processed_lines = []
                processed_busses = []
                if len(remote_bus_bars):
                    for remote_bus_bar in remote_bus_bars:
                        connected_element_list = self.interface.get_bus_connections_of(\
                                                                        remote_bus_bar)
                        # list of all branches already investigated
                        processed_branches = []
                        for element in connected_element_list:
                            element_branches = self.grid.get_branches_of(element)
                            for element_branch in element_branches:
                                # skip the line from which we are coming from
                                if element_branch == branch or element_branch == None:
                                    continue
                                # skip branches with a single terminal
                                if element_branch.is_terminal_branch() == True:
                                    if self.output_detail >= OutputDetail.NORMAL:
                                        self.interface.print("\n\t     Skipping " + \
                                                                str(element_branch.name))
                                    continue
                                if self.output_detail >= OutputDetail.NORMAL:
                                    self.interface.print("\n\tChecking " + \
                                    self.interface.get_name_of(element) + " (branch: " + \
                                    element_branch.name + ")")
                                if element_branch in processed_branches:
                                    if self.output_detail >= OutputDetail.NORMAL:
                                        self.interface.print("\t\t Already processed!")
                                    continue
                                processed_branches.append(element_branch)
                                # if the branch is a transformer
                                trafo_branch = True if element_branch.is_transformer_branch()\
                                          else False
                                remote_bus_bar_2nd_step = element_branch.get_last_busbar_from(\
                                                                            remote_bus_bar)
                                remote_bus_bars_2nd_step = []
                                remote_elements_2nd_step = []
                                if remote_bus_bar_2nd_step and\
                                 self.is_tap_bus(remote_bus_bar_2nd_step):
                                    if self.output_detail >= OutputDetail.NORMAL:
                                                        self.interface.print("\t" +
                                                        self.interface.get_name_of(\
                                                        remote_bus_bar_2nd_step) + \
                                                         " is a tap busbar")
                                    remote_bus_bars_2nd_step = self._get_remote_busses_of(\
                                            remote_bus_bar_2nd_step, branch, \
                                            remote_bus_bars_2nd_step, remote_elements_2nd_step)
                                else:
                                    remote_bus_bars_2nd_step.append(remote_bus_bar_2nd_step)
                                    remote_elements_2nd_step.append(element)
                                if remote_bus_bars_2nd_step:
                                    initial_busbars_number = len(\
                                                        remote_bus_bars_2nd_step)
                                    for bus_index, (remote_bus_bar_2nd_step, remote_element_2nd_step)  \
                                            in enumerate(zip(remote_bus_bars_2nd_step,
                                                             remote_elements_2nd_step)):
                                        if trafo_branch == False and \
                                        bus_index < initial_busbars_number:
                                            # list of bus bars connected to the "remote_bus_bars_2nd_step"
                                            # throw a transformer
                                            connected_element_list_2nd = self.interface.\
                                            get_bus_connections_of(remote_bus_bar_2nd_step)
                                            # add in the list the trafo connection to the
                                            # bus 2nd step
                                            remote_bus_bars_3rd_step = []
                                            # and the elements
                                            remote_elements_3rd_step = []
                                            for element_2nd in connected_element_list_2nd:
                                                # skip the line from which we are coming from
#                                                 element_branch_2nd =\
#                                                  self.grid.get_branch_of(element_2nd)
                                                element_branches_2nd = \
                                                 self.grid.get_branches_of(element_2nd)
                                                for element_branch_2nd in element_branches_2nd:
                                                    if element_branch_2nd == element_branch or\
                                                     element_branch_2nd == None:
                                                        continue
                                                    # skip branches where I cannot apply a fault
        #                                             if element_branch_2nd.is_transformer_branch() == False and\
        #                                             element_branch_2nd.is_line_branch() == False:
        #                                                 if self.output_detail >= OutputDetail.NORMAL:
        #                                                     self.interface.print("\t\tSkipping " +\
        #                                                                 element_branch_2nd.name)
        #                                                 continue
                                                    if element_branch_2nd.is_transformer_branch():
                                                        remote_bus_bar_3rd_step = \
                                                        element_branch_2nd.get_last_busbar_from(\
                                                            remote_bus_bar_2nd_step)
                                                        if remote_bus_bar_3rd_step not in\
                                                        remote_bus_bars_3rd_step:
                                                            remote_bus_bars_3rd_step.append(\
                                                                remote_bus_bar_3rd_step)
                                                            remote_elements_3rd_step.append(\
                                                            element_branch_2nd.transformer)
                                                            if self.output_detail >= OutputDetail.NORMAL:
                                                                self.interface.print("\t\t adding " +
                                                                self.interface.get_name_of(\
                                                                remote_bus_bar_3rd_step) + \
                                                                 " as remote busbar trafo busbar")
                                            remote_bus_bars_2nd_step += remote_bus_bars_3rd_step
                                            remote_elements_2nd_step += remote_elements_3rd_step
                                        # avoid duplicates
                                        remote_bus_bars_2nd_step = list(set(remote_bus_bars_2nd_step))
                                        remote_elements_2nd_step = list(set(remote_elements_2nd_step))
                                        if self.output_detail >= OutputDetail.NORMAL:
                                            self.interface.print("\t 3PH SHC at " + \
                                            self.interface.get_name_of(remote_element_2nd_step) + " (bus " + \
                                            self.interface.get_name_of(remote_bus_bar_2nd_step) + \
                                            ")")
                                        # if remote_bus_bar_2nd_step in remote_bus_bars_3rd_step:
                                        if bus_index >= initial_busbars_number:
                                            trafo_branch = True
                                            if self.output_detail >= OutputDetail.NORMAL:
                                                self.interface.print(\
                                                '          Checking Remote busbar trafo busbar  ' + \
                                                self.interface.get_name_of(\
                                                    remote_bus_bar_2nd_step))
                                        remote_element_2nd_step_branch = \
                                        self.grid.get_branch_of(remote_element_2nd_step)
                                        if trafo_branch == False and remote_element_2nd_step_branch and\
                                        remote_element_2nd_step_branch.is_line_branch() == True:
                                            # fault at the 2nd step remote bus bar location
                                            # but along the line
                                            if self.output_detail >= OutputDetail.NORMAL:
                                                self.interface.print('          Fault at line ' + \
                                                self.interface.get_name_of(remote_element_2nd_step))
                                            if remote_element_2nd_step in processed_lines:
                                                if self.output_detail >= OutputDetail.NORMAL:
                                                    self.interface.print(\
                                                       "          skipped: Fault already executed ")
                                                continue
                                            processed_lines.append(remote_element_2nd_step)
                                            self.interface.set_fault_position(\
                                                    remote_element_2nd_step, \
                                                    position=0.01, \
                                                    single_shc=True, \
                                                    reference_busbar=remote_bus_bar_2nd_step)
                                        else:
                                            if self.output_detail >= OutputDetail.NORMAL:
                                                self.interface.print(\
                                                    "          Fault at bus ")
                                            if remote_bus_bar_2nd_step in processed_busses:
                                                if self.output_detail >= OutputDetail.NORMAL:
                                                    self.interface.print(\
                                                    "          skipped: Fault already executed ")
                                                continue
                                            processed_busses.append(remote_bus_bar_2nd_step)
                                            self.interface.set_fault_position(\
                                                        remote_bus_bar_2nd_step, \
                                                        single_shc=True)

                                        # find the breaker at the 2nd remote end
                                        breaker = None
                                        if line:
                                            breakers = self.interface.\
                                            get_branch_breakers_of(remote_element_2nd_step)
                                            bus_side1 = self.interface.\
                                            get_breaker_bus(breakers[0]) if breakers\
                                            else None
                                            # check that the right breaker (right side)
                                            # has been selected
                                            if bus_side1 and breakers:
                                                breaker = breakers[0] \
                                                if bus_side1 == remote_bus_bar_2nd_step or\
                                                len(breakers) == 1 else breakers[1]
                                            if breaker == None:
                                                if self.output_detail >= OutputDetail.NORMAL:
                                                        self.interface.print(\
                                                        "          Breaker not available")
                                        # double iteration with breaker open and closed
                                        operate_breaker_mode = [True, False]
                                        for breaker_active in operate_breaker_mode:
                                            if breaker and breaker_active == True and\
                                            trafo_branch == False:
                                                if self.output_detail >= OutputDetail.NORMAL:
                                                    self.interface.print(\
                                                    "          Open breaker " + \
                                                    self.interface.get_full_name_of(breaker))
                                                    self.interface.print("\t\t Running SHC")
                                                self.interface.set_echo_off()
                                                self.interface.open_(breaker)
                                                self.interface.set_echo_on()
                                            # run shc
                                            try:
                                                self.interface.set_echo_off()
                                                error = self.interface.run_shc()
                                                self.interface.set_echo_on()
                                                if error == 1:
                                                    if self.output_detail >= OutputDetail.NORMAL:
                                                        self.interface.print(\
                                                            "\t\tSHC Failed. Run now a IEC SHC")
                                                    self.interface.set_shc_basic_configuration(
                                                                        self.interface.SHC_Mode.IEC60909)
                                                    self.interface.set_echo_off()
                                                    error = self.interface.run_shc()
                                                    self.interface.set_shc_basic_configuration(
                                                                        self.interface.SHC_Mode.COMPLETE)
                                                    self.interface.set_echo_on()
                                                if error == 0:
                                                    z3 = self.interface.get_relay_secondary_z_measures_of(\
                                                                                polarizing_element)
                                                    cz3 = self.interface.\
                                                    get_relay_secondary_complex_z_measures_of(\
                                                                    polarizing_element)
                                                    z3_primary = [ z * vt_ratio / ct_ratio \
                                                        for z in z3]
                                                    cz3_primary = [ cz * vt_ratio / ct_ratio \
                                                        for cz in cz3]
                                                    z3_zscc_data.append(z3_zscc_value(\
                                                        line=line,
                                                        relay=relay,
                                                        observation_cubicle=cubicle,
                                                        fault_type='3PH',
                                                        fault_bus=remote_bus_bar_2nd_step,
                                                        zscc_values=cz3_primary,
                                                        relay_bus=bus_bar,
                                                        trafo_shc=trafo_branch,
                                                        open_breaker=breaker_active)\
                                                                        )
                                                    if self.output_detail >= OutputDetail.NORMAL:
                                                        if trafo_branch == True:
                                                            self.interface.print(\
                                                        "\t\tTrafo detected. Breaker left closed. SHC at the busbar")
                                                        self.interface.print("\t\tMeasured Z = " + \
                                                                             str(z3_primary) + " pri.Ohm")
                                                    if trafo_branch == True:
                                                        trafo_z_limit = z3_primary[0] * 0.80\
                                                        if z3_primary[0] * 0.80 < trafo_z_limit\
                                                        else trafo_z_limit
                                                    else:
                                                        z3_max = z3_primary[0] * mult \
                                                        if z3_primary[0] * mult > z3_max\
                                                        else z3_max
                                            except Exception as e:
                                                pass
                                            if breaker and breaker_active == True and\
                                            trafo_branch == False:
                                                if self.output_detail >= OutputDetail.NORMAL:
                                                    self.interface.print(\
                                                    "          Closed breaker ")
                                                    self.interface.print("\t\t Running SHC")
                                                self.interface.set_echo_off()
                                                self.interface.close_(breaker)
                                                self.interface.set_echo_on()
                                            # if we are evaluating one of the additional
                                            # bus bars we reset the trafo branch flag
                                            if bus_index >= initial_busbars_number\
                                            or trafo_branch:
                                                trafo_branch = False
                                                break
                    if z3_zscc_data:
                        if self.output_detail >= OutputDetail.NORMAL:
                            self.interface.print("\n\tstored " + str(len(z3_zscc_data)) + \
                            " set of Z values for " + self.interface.get_name_of(relay))
                        self.z3s_zscc_data.update({relay:z3_zscc_data})
        if self.output_detail >= OutputDetail.NORMAL:
            self.interface.print("\n\n    Max Measured Z limit = " + str(round(z3_max, 2)) + \
                                  " sec.Ohm")
            self.interface.print("    Min Measured trafo Z limit= " + \
                                 str(round(trafo_z_limit, 2)) + " sec.Ohm")
        return z3_max if z3_max < trafo_z_limit and z3_max > 0 else trafo_z_limit

    def get_relay_ratios_of(self, relay):
        '''
        function returning the ct and the vt ratio of the given relay 
        '''
        # default value for the return variables
        ct_ratio = 1
        vt_ratio = 1
        # get the cubicle hosting the given relay
        cubicle = self.interface.get_relay_cubicle_of(relay)
        if cubicle:
            # get the Ct ratio
            ct_list = self.interface.get_cubicle_CT_of(cubicle)
            if ct_list:
                ct_ratio = self.interface.get_ct_ratio(ct_list[0])
            # set the Ct direction to the branch
            # temporary code!
            if ct_list:
                self.interface.set_attribute(ct_list[0], 'ciorient', 0)
            # get the Vt ratio
            vt_list = self.interface.get_cubicle_VT_of(cubicle)
            if vt_list:
                vt_ratio = self.interface.get_vt_ratio(vt_list[0])
        return ct_ratio, vt_ratio

    def get_relay_measurements_in(self, meas_cubicle):
        '''
        get the phase + N currents from the relay in the given cubicle
        '''
        ikss = []
        ct_list = self.interface.get_cubicle_CT_of(meas_cubicle)
        if len(ct_list) > 0 :
            ct_ratio = self.interface.get_ct_ratio(ct_list[0])
            meas_relays = self.interface.get_cubicle_relay_of(meas_cubicle)
            if meas_relays:
                one_relay_active = False
                for meas_relay in meas_relays:
                    if self.interface.is_out_of_service(meas_relay) == False:
                        if self.output_detail >= OutputDetail.DEBUG:
                            self.interface.print("\nGetting measures from relay " + \
                                        self.interface.get_name_of(meas_relay))
                        measurement = self.interface.get_relay_measurement_element(meas_relay)
                        new_values = [sec_meas * ct_ratio for sec_meas \
                            in self.interface.get_relay_current_measures_of(measurement)]
                        ikss += new_values
                        one_relay_active = True
                        break
                    else:
                        if self.output_detail >= OutputDetail.DEBUG:
                            self.interface.print("\nMeasurement relay " + \
                            self.interface.get_name_of(meas_relay) + " is disabled ")
                if one_relay_active == False:
                    ikss += [0, 0, 0, 0]
                    if self.output_detail >= OutputDetail.DEBUG:
                            self.interface.print("\nNo Measurement relay in service")
            else:  # no relay add empty values
                ikss += [0, 0, 0, 0]
        else:  # no CT add empty values
            ikss += [0, 0, 0, 0]
        return ikss

    def get_relay_complex_measurements_in(self, meas_cubicle):
        '''
        get the phase + N currents from the relay in the given cubicle as real and 
        immaginary part
        '''
        ikss = []
        ct_list = self.interface.get_cubicle_CT_of(meas_cubicle)
        if len(ct_list) > 0 :
            ct_ratio = self.interface.get_ct_ratio(ct_list[0])
            meas_relays = self.interface.get_cubicle_relay_of(meas_cubicle)
            if len(meas_relays) > 0:
                measurement = self.interface.get_relay_measurement_element(meas_relays[0])
                new_values = [sec_meas * ct_ratio for sec_meas \
                        in self.interface.get_relay_complex_current_measures_of(measurement)]
                ikss += new_values
        else:  # no CT add empty values
            ikss += [0, 0, 0, 0]
        return ikss

    def calculate_shunt_SHCs(self, window, input_settings=None):
        '''
        function calculating a ph-grnd shc values at the busbar connected to the
        given shunt
        '''
        from math import pi
        from cmath import phase
        settings = self.initialize(window, input_settings)
        if self.is_dialog_setting_ok(window, settings) == False:
            return 1
        # Print some info in the PF output window
        if self.output_detail.value >= OutputDetail.NORMAL.value:
            self.interface.print("*********************************************************")
            self.interface.print("    Time Distance Creator Tool (Beta) 0.01")
            self.interface.print("          Calculate shunt SHC values      ")
            self.interface.print("*********************************************************n")

        if self.output_detail >= OutputDetail.NORMAL:
            self.interface.print("\nCalculating Shcs at shunts ")
        shunts = self.interface.get_available_shunts()
        self.workbook, worksheet = self.create_shunt_output_file_header(window)
        ended = False
        # data structure to save the shc values and the relevant grid elements
        Shunt_iscc_value = namedtuple("Shunt_iscc_value", \
                    "Shunt observation_cubicle fault_type fault_bus iscc_values remote_bus")
        self.shunts_iscc_data = {}
        for row_index, shunt in enumerate(shunts):
            shunt_name = self.interface.get_name_of(shunt)
            if self.output_detail >= OutputDetail.NORMAL:
                self.interface.print("\tCalcukating Shc at shunt " + shunt_name)
            bus = self.interface.get_branch_busses_of(shunt, self.interface.Side.Side_1)
            worksheet.write('A' + str(row_index + 2), shunt_name)
            worksheet.write('B' + str(row_index + 2), self.interface.get_name_of(bus))
            worksheet.write('C' + str(row_index + 2) , '1PH-Grnd')

            # fault at the busbar
            self.interface.set_fault_position(bus, position=0.01,
                                              single_shc=True)
            # set fault type and resistance
            fault_type = self.fault_types["1PH"].type
            self.interface.set_fault(fault_type, 0)
            # run shc
            try:
                self.interface.set_echo_off()
                error = self.interface.run_shc()
                self.interface.set_echo_on()
                if error == 0:
                    cubicles = self.interface.get_branch_cubicles_of(shunt)
                    if cubicles:
                        if self.output_detail >= OutputDetail.NORMAL:
                            self.interface.print("\t    " + str(fault_type) + \
                                        " SHC executed at " + shunt_name)
                        ikss = []
                        ikss += self.get_relay_complex_measurements_in(cubicles[0])
                        meas_number = 4
                        for i in range(meas_number):
                            worksheet.write(chr(ord('D') + i * 2) + str(row_index + 2), \
                                                    str(round(abs(ikss[i]), 1)))
                            worksheet.write(chr(ord('E') + i * 2) + str(row_index + 2), \
                                                str(round(phase(ikss[i]) * 180 / pi, 2)))
                else:
                    if self.output_detail >= OutputDetail.NORMAL:
                        self.interface.print("\t  ERROR " + str(error))
            except KeyboardInterrupt:
                if self.output_detail >= OutputDetail.NORMAL:
                    self.interface.print("\nInterrupted by the user!")
                ended = True
            except Exception as e:
                pass
            if ended == True:
                break
            # just for debug
#             if row_index == 20:
#                 break
        if self.output_detail >= OutputDetail.NORMAL:
            self.interface.print("\n\t Shunt Shcs calculation completed! ")

    def calculate_generator_SHCs(self, window, input_settings=None):
        '''
        function calculating a 3ph shc values at the busbar connected to the
        given generator
        '''
        from math import pi
        from cmath import phase
        settings = self.initialize(window, input_settings)
        if self.is_dialog_setting_ok(window, settings) == False:
            return 1
        if self.output_detail.value >= OutputDetail.NORMAL.value:
            self.interface.print("*********************************************************")
            self.interface.print("    Time Distance Creator Tool (Beta) 0.01")
            self.interface.print("          Calculate Gnerator SHC values      ")
            self.interface.print("*********************************************************\n")
        if self.output_detail >= OutputDetail.NORMAL:
            self.interface.print("\nCalculating SHCs at generators ")
        # create grid layout only
        generators = self.interface.get_generators(self.generator_selection_criteria)
        self.create_system_layout(window, input_settings, grid_only=True)
        worksheet = self.create_generator_output_file_header(window)
        ended = False
        # data structure to save the shc values and the relevant grid elements
        Generator_zscc_value = namedtuple("Generator_zscc_value", \
                    "generator observation_cubicles fault_type fault_bus zscc_values remote_bus open_breaker")
        self.generators_iscc_data = {}
        row_index = 2
        if self.output_detail >= OutputDetail.NORMAL:
                self.interface.print("       SHCs will be calculated for " + \
                                      str(len(generators)) + " generators")
        for gen_index, generator in enumerate(generators):
            if self.output_detail >= OutputDetail.NORMAL:
                if self.interface.is_out_of_service(generator):
                    self.interface.print("       " + self.interface.\
                    get_name_of(generator) + " is out of service")
                if self.interface.is_energized(generator) == False:
                    self.interface.print("       " + self.interface.\
                    get_name_of(generator) + " is not energized")
            if self.interface.is_out_of_service(generator) == True or\
            self.interface.is_energized(generator) == False:
                continue
            # list of the busses where the shc should be calculated
            remote_bus_bars = []
            # list of the lines where the shc should be calculated
            remote_lines = []
            generator_name = self.interface.get_name_of(generator)
            if self.output_detail >= OutputDetail.NORMAL:
                self.interface.print("    Calculating Shc for generator " + \
                generator_name + ' (' + str(gen_index) + '\\' + \
                                                    str(len(generators)) + ')')
            bus = self.interface.get_branch_busses_of(generator, self.interface.Side.Side_1)
            branch = self.grid.get_branch_of(generator)
            line = None
            # if the branch is a transformer branch get the other side
            if branch:
                if branch.is_transformer_branch():
                    if self.output_detail >= OutputDetail.NORMAL:
                        self.interface.print("      Step up trafo: " + \
                                self.interface.get_name_of(branch.transformer))
                    bus = branch.get_other_busbar_of(bus)
                else:  # otherwise multiple generators connected to the same busbar
                    branches = self.grid.get_branch_of(bus)
                    if branches:
                        for branch in branches:
                            if branch.is_transformer_branch():
#                                 if self.output_detail >= OutputDetail.NORMAL:
#                                     self.interface.print("      Step up trafo2: " +\
#                                     self.interface.get_name_of(branch.transformer))
                                bus = branch.get_other_busbar_of(bus)
                                line = branch.lines[0] if branch.lines else None
                                break
                if bus:
                    remote_branches = self.grid.busbar_list[bus]
                    # remote_bus_bars.append(bus)
                    # remote_lines.append(line)
                    remote_bus_bars += [branch.get_other_busbar_of(bus)\
                                       for branch in remote_branches \
                                       if branch.is_line_branch()]
                    new_lines = [branch.lines for branch in remote_branches
                                       if branch.is_line_branch()]
                    if new_lines:
                        new_lines_flat = [line for sublist in new_lines for line in sublist]
                        if new_lines_flat and len(new_lines_flat) > 0:
#                             if self.output_detail >= OutputDetail.VERBOSEDEBUG:
#                                 for line in new_lines_flat:
#                                     self.interface.print(self.interface.\
#                                                          get_name_of(line))
                            remote_lines += new_lines_flat

            worksheet.write('A' + str(row_index), generator_name)
            if bus:
                worksheet.write('B' + str(row_index), \
                                self.interface.get_name_of(bus))
            worksheet.write('C' + str(row_index) , '3 Phase')

            # set fault type and resistance
            fault_type = self.fault_types["3PH"].type
            self.interface.set_shc_basic_configuration(
                                            self.interface.SHC_Mode.COMPLETE)
            self.interface.set_fault(fault_type, 0)

            # get the relay at the measurement point
            cubicles = self.interface.get_branch_cubicles_of(generator)
            polarizing_element = None
            breakers = []
            if cubicles:
                relays = self.interface.get_cubicle_relay_of(cubicles[0])
                if relays:
                    # get the relay polarizing element
                    polarizing_element = self.interface.\
                                    get_relay_polarizing_element(relays[0])
                    ct_list = self.interface.get_cubicle_CT_of(cubicles[0])
                    if ct_list:
                        ct_ratio = self.interface.get_ct_ratio(ct_list[0])
                    else:
                        ct_ratio = 1
                    # set the Ct direction to the busbar
                    # temporary code!
                    if ct_list:
                        self.interface.set_attribute(ct_list[0], 'ciorient', 1)
                    vt_list = self.interface.get_cubicle_VT_of(cubicles[0])
                    if vt_list:
                        vt_ratio = self.interface.get_vt_ratio(vt_list[0])
                    else:
                        vt_ratio = 1
                else:
                    if self.output_detail >= OutputDetail.NORMAL:
                        self.interface.print("\n ERROR: no relay available!")
            generator_zscc_data = []
            # avoid duplicated elements in the busbars and lines lists
            remote_bus_bars = list(set(remote_bus_bars))
            remote_lines = list(set(remote_lines))
            for bus_bar, line in zip(remote_bus_bars, remote_lines):
                # chek if bus is a tap
                remote_paths = []
                if bus_bar and self.is_tap_bus(bus_bar):
                    if bus and self.output_detail >= OutputDetail.NORMAL:
                        self.interface.print("   Bus " + self.interface.\
                            get_name_of(bus) + " is a tap busbar")
                    remote_bus_bars += self._get_remote_busses_of(bus_bar, \
                                                        branch, remote_bus_bars, \
                                                         remote_paths)
                # get the remote line breaker
                breaker = None
                if line:
                    breakers = self.interface.get_branch_breakers_of(line)
                    bus_side1 = self.interface.get_branch_busses_of(line, \
                                                self.interface.Side.Side_1)
                    if bus_side1 and breakers:
                        breaker = breakers[0] if bus_side1 == bus_bar or\
                        len(breakers) == 1 else breakers[1]

                # fault at the busbar
                self.interface.set_fault_position(bus_bar if line == None \
                                                  else line, \
                        position=99.9 if line == None or \
                         self.interface.get_cubicle_busbar_of(\
                        self.interface.get_line_cubicle_j_of(line)) == bus_bar \
                         else 0.01,
                        single_shc=True)
                # run shc
                try:
                    operate_breaker_mode = [True, False]
                    for breaker_active in operate_breaker_mode:
                        zshc = []
                        if breaker and breaker_active == True:
                            if self.output_detail >= OutputDetail.NORMAL:
                                self.interface.print("          Operating breaker " + \
                                self.interface.get_full_name_of(breaker))
                            self.interface.open_(breaker)
                        self.interface.set_echo_off()
                        error = self.interface.run_shc()
                        self.interface.set_echo_on()
                        if error == 1:
                            self.interface.set_shc_basic_configuration(
                                                self.interface.SHC_Mode.IEC60909)
                            self.interface.set_echo_off()
                            error = self.interface.run_shc()
                            self.interface.set_shc_basic_configuration(
                                                self.interface.SHC_Mode.COMPLETE)
                            self.interface.set_echo_on()
                        if error == 0:
                            if self.output_detail >= OutputDetail.NORMAL:
                                if line == None:
                                    self.interface.print("          " + str(fault_type) + \
                                                " SHC executed at busbar " + self.interface.\
                                                get_name_of(bus_bar))
                                else:
                                    self.interface.print("          " + str(fault_type) + \
                                                " SHC executed at line " + self.interface.\
                                                get_name_of(line))
                            if polarizing_element:

                                new_zshc = self.interface.\
                                get_relay_secondary_complex_z_measures_of(\
                                                            polarizing_element)
                                new_primary_zshc = [ z * vt_ratio / ct_ratio \
                                                    for z in new_zshc]
                                zshc += new_primary_zshc
                                generator_zscc_data.append(Generator_zscc_value(\
                                                generator=generator,
                                                observation_cubicles=cubicles,
                                                fault_type=fault_type,
                                                fault_bus=bus_bar,
                                                zscc_values=zshc,
                                                remote_bus=bus,
                                                open_breaker=breaker_active))
                                if self.output_detail >= OutputDetail.NORMAL:
                                    self.interface.print("          " +
                                    str(fault_type) + " Z = " + str(zshc) + " ohm")
                                    self.interface.print("          " +
                                    str(fault_type) + " Measured in " + \
                                    self.interface.get_full_name_of(polarizing_element))
                        else:
                            if self.output_detail >= OutputDetail.NORMAL:
                                self.interface.print("\n ERROR = " + str(error))
                        if breaker and breaker_active == True:
                            if self.output_detail >= OutputDetail.NORMAL:
                                self.interface.print("          Closing breaker ")
                            self.interface.set_echo_off()
                            self.interface.close_(breaker)
                            self.interface.set_echo_on()
                except KeyboardInterrupt:
                    if self.output_detail >= OutputDetail.NORMAL:
                        self.interface.print("\nInterrupted by the user!")
                    ended = True
                except Exception as e:
                    pass
                if ended == True:
                    break

                self.generators_zscc_data.update({generator:generator_zscc_data})
            if self.output_detail >= OutputDetail.NORMAL:
                self.interface.print("     Saving data for generator " + \
                                    self.interface.get_name_of(generator))
            if generator_zscc_data:
                self.fill_generator_output_file(generator_zscc_data, \
                                  worksheet,
                                  row_index)

            row_index += 1
#           just for debug
#             if gen_index > 0:
#                 break
        if self.output_detail >= OutputDetail.NORMAL:
            self.interface.print("\n Generator Shcs calculation completed! ")

    def calculate_line_SHCs(self, window, input_settings=None):
        '''
        function calculating the shc values at 0% and at 100% of the line
        '''
        settings = self.initialize(window, input_settings)
        if self.is_dialog_setting_ok(window, settings) == False:
            return 1
        if self.output_detail.value >= OutputDetail.NORMAL.value:
            self.interface.print("*********************************************************")
            self.interface.print("    Time Distance Creator Tool (Beta) 0.01")
            self.interface.print("          Calculate Line SHC values      ")
            self.interface.print("*********************************************************\n")
        if self.output_detail >= OutputDetail.NORMAL:
            self.interface.print("\nCalculating Shcs at lines ")
        # create grid layout only
        self.create_system_layout(window, input_settings, grid_only=True)

        if len(settings['StudySelectedBus']) == 0:
            lines = self.interface.get_lines(self.line_selection_criteria)
        else:
            lines = self.grid.get_bus_lines_from(busbar_name=settings['StudySelectedBus'],
                                    number_of_steps=settings['StudySelectedBusExtent'])
        # row index in the output file
        start_index = 2
        line_sheet = self.create_line_output_file_header(window)

        # data structure to save the shc values and the relevant grid elements
        Line_iscc_value = namedtuple("Line_iscc_value", \
                    "line observation_cubicles fault_type fault_bus iscc_values remote_bus")
        self.lines_iscc_data = {}
        ended = False
        for line_index, line in enumerate(lines):
            if self.interface.is_energized(line) and\
                    self.interface.is_out_of_service(line) == False:
                if self.output_detail >= OutputDetail.NORMAL:
                    self.interface.print("  Processing line " + \
                    self.interface.get_name_of(line) + '  (' + str(line_index + 1) + \
                               '\\' + str(len(lines)) + ')')
                cubicles = [self.interface.get_line_cubicle_i_of(line), \
                        self.interface.get_line_cubicle_j_of(line)]
                # check if the busses delimiting the line are tap busses
                remote_bus_bars = []
                remote_paths = []
                first_remote_bus = self.interface.get_cubicle_busbar_of(\
                                    self.interface.get_line_cubicle_j_of(line))
                for cubicle in cubicles:
                    line_bus_bar = self.interface.get_cubicle_busbar_of(cubicle)
                    remote_bus_bars.append(line_bus_bar)
                    line_branch = self.grid.get_branch_of(line)
                    if line_bus_bar and self.is_tap_bus(line_bus_bar):
                        if self.output_detail >= OutputDetail.NORMAL:
                            self.interface.print("   Bus " + self.interface.\
                                get_name_of(line_bus_bar) + " is a tap busbar")
                        remote_bus_bars += self._get_remote_busses_of(line_bus_bar, \
                                                            line_branch, remote_bus_bars, \
                                                             remote_paths)
                line_iscc_data = []
                fault_types = [ftype_key for ftype_key, nfault_type \
                            in self.fault_types.items() if nfault_type.active == True]
                remote_bus_bars = list(set(remote_bus_bars))
                first_remote_bus_index = remote_bus_bars.index(first_remote_bus)
                if self.output_detail >= OutputDetail.NORMAL:
                                self.interface.print("   Calculating SHC for line " + \
                                            self.interface.get_name_of(line))
                for fault_type in fault_types:
                    for bus_index, busbar in enumerate(remote_bus_bars):
                        ikss = []
                        # fault at the busbar for the over the tap remote busbar
                        # otherwise at the line
                        self.interface.set_fault_position(line if bus_index == 0 or\
                                        busbar == first_remote_bus else busbar,
                                position=0.01 if bus_index < first_remote_bus_index\
                                else 99.9
                                                          , single_shc=True)
                        # set fault type and resistance
                        self.interface.set_fault(self.fault_types[fault_type].type, 0)
                        # run shc
                        try:
                            if self.output_detail >= OutputDetail.NORMAL:
                                self.interface.print("      " + str(fault_type) + \
                                                     " SHC executed at " + \
                                            self.interface.get_name_of(busbar))
                            self.interface.set_echo_off()
#                             self.interface.disable_pf_gui_update()
                            error = self.interface.run_shc()
                            self.interface.set_echo_on()
#                             self.interface.enable_pf_gui_update()
                            if error == 0:
                                cub_i = self.interface.get_line_cubicle_i_of(line)
                                ikss += self.get_relay_measurements_in(cub_i)
                                cub_j = self.interface.get_line_cubicle_j_of(line)
                                ikss += self.get_relay_measurements_in(cub_j)
                            else:
                                ikss = [-1, -1, -1, -1, -1, -1, -1, -1]
                                if self.output_detail >= OutputDetail.NORMAL:
                                    self.interface.print("ERROR running the SHC!")
                            line_iscc_data.append(Line_iscc_value(line=line, \
                                            observation_cubicles=[cub_i, cub_j],
                                            fault_type=fault_type,
                                            fault_bus=busbar,
                                            iscc_values=ikss,
                                            remote_bus=False \
                                            if bus_index < first_remote_bus_index\
                                             else True))
                        except KeyboardInterrupt:
                            if self.output_detail >= OutputDetail.NORMAL:
                                self.interface.print("\nInterrupted by the user!")
                            ended = True
                        except Exception as e:
                            if self.output_detail >= OutputDetail.NORMAL:
                                self.interface.print("Exception triggered!")

                self.lines_iscc_data.update({line:line_iscc_data})
                if self.output_detail >= OutputDetail.NORMAL:
                                self.interface.print("   Saving data for line " + \
                                            self.interface.get_name_of(line))
                self.fill_line_output_file(line_iscc_data, \
                                      line_sheet,
                                      start_index)

                start_index += len(fault_types) * 2
                if ended == True:
                    break
        if self.output_detail >= OutputDetail.NORMAL:
            self.interface.print("\n Line Shcs calculation completed! ")

    def create_TCD(self, window, input_settings=None):
        '''
        function perfoming the TCD creation process
        '''

        settings = self.initialize(window, input_settings)
        if self.is_dialog_setting_ok(window, settings) == False:
            return 1

        # Print some info in the PF output window
        if self.output_detail.value >= OutputDetail.NORMAL.value:
            self.interface.print("*********************************************************")
            self.interface.print("    Time Distance Creator Tool (Beta) 0.01")
            self.interface.print("                 Creating TCDs            ")
            self.interface.print("*********************************************************\n")

        # just for debugging
        # self.calculate_line_SHCs(settings, window)
        if self.output_detail >= OutputDetail.NORMAL:
            self.interface.print("\nCreating TC diagrams ")

        # to remove
        if len(settings['StudySelectedBus']) == 0:
            transformers = self.interface.get_transformers(self.trafo_selection_criteria)
        else:
            transformers = self.grid.get_bus_transformers_from(busbar_name=settings['StudySelectedBus'],
                                    number_of_steps=settings['StudySelectedBusExtent'])
        # end to remove
        # to restore
        # lines, transformers = self.create_system_layout(window, input_settings)
        trafo_sheet = self.create_trafo_output_file_header(window)

        # max number of transformer diagrams in a test case
        max_number_trafo_in_testcase = 1000
        # max_number_trafo_before_refresh
        max_number_trafo_refresh = 100
        # number of trafo for a new message in the output window
        max_number_trafo_message = 10
        # row index in the output file
        start_index = 2
        self.interface.disable_pf_gui_update()
        self.interface.set_echo_off()

        # data structure to save the shc values and the graphical objects where the
        # vertical lines are plotted
        Trafo_iscc_value = namedtuple("Trafo_iscc_value", "diagram label iscc_value iscc_value_at_v1")
        trafo_iscc_data = []

        for trafo_index, transformer in enumerate(transformers):
            if self.output_detail >= OutputDetail.NORMAL:
                    self.interface.print("  Processing trafo " + \
                    self.interface.get_name_of(transformer) + '  (' + str(trafo_index + 1) + \
                               '\\' + str(len(transformers)) + ')')
#             if self.interface.get_name_of(transformer) != 'tr3_101032_801090_801091_1' and\
#             self.interface.get_name_of(transformer) != 'trf_101028_801088_2' and\
#             self.interface.get_name_of(transformer) != 'trf_408001_908004_1':
#                 continue
            # output_file, trafo_sheet = self.get_workbook_and_worksheet(window)
            # if the elemnt is not a transformer ...skip it!
            if 'ElmTr' not in self.interface.get_class_name_of(transformer):
                continue
#             self.interface.enable_last_enabled_study_case()
#             if trafo_index > 0 and trafo_index % max_number_trafo_message == 0:
#                 self.interface.set_echo_on()
#                 self.interface.print(" Transformer # " + str(trafo_index) + "/" + \
#                                      str(len(transformers)))
#                 self.interface.set_echo_off()
            if trafo_index > 0 and trafo_index % max_number_trafo_refresh == 0:
                active_project = self.interface.deactivate_project()
#                 self.interface.refresh_pf()
                self.interface.activate_project(active_project)
#             if trafo_index > 0 and trafo_index % max_number_trafo_in_testcase == 0:
#                 self.interface.disable_current_study_case()
#                 new_study_case = self.interface.\
#                             create_study_case_from_last_enabled_study_case()
#                 self.interface.delete_TCC_pages(new_study_case)
            cubicles = self.interface.get_transformer_cubicle_of(transformer)
            items = [tuple([transformer])]
            # shc current list
            ikss = []
            ikss_v1 = []  # shc values at the primary voltage
            # toc diagram list
            toc_diagram = []

            # get the tarnsformer winding voltages
            trafo_voltages = self.interface.get_transformer_voltages_of(transformer)

            # iterate throw all cubicles
            for index, cubicle in enumerate(cubicles):
                found_relays = self.interface.get_cubicle_relay_of(cubicle)
                items.append(tuple(found_relays))

                # try to get the SHC current for each relay
                # set the shc main param
                self.interface.set_shc_basic_configuration(self.interface.SHC_Mode.COMPLETE)
                # skip the HV side
                # if index == 0:
                #    continue

                fault_types = [ftype_key for ftype_key, nfault_type \
                        in self.fault_types.items() if nfault_type.active == True]
                # set the fault position
                fault_bus_bar = self.interface.get_cubicle_busbar_of(cubicle)
                if self.output_detail >= OutputDetail.NORMAL:
                    if self.interface.is_out_of_service(transformer) == True:
                        self.interface.print("     Trafo " + \
                        self.interface.get_name_of(transformer) + " is disabled")
                if self.interface.is_energized(fault_bus_bar):
                    self.interface.print("    Applying SHC at " + \
                                         self.interface.get_name_of(fault_bus_bar))
                    self.interface.set_fault_position(\
                                fault_bus_bar, \
                                position=0, single_shc=True)

                    for fault_type in fault_types:
                        # set fault type and resistance
                        self.interface.set_fault(self.fault_types[fault_type].type, 0)
                        # run shc
                        try:
                            error = self.interface.run_shc()
                            self.interface.print("    " + str(fault_type) + \
                                                         " SHC executed")
                            if error == 0:
                                if self.output_detail >= OutputDetail.NORMAL:
                                    self.interface.print("    Getting SHC results")
                                # get the current at the first cubible (hopefully the HV cubicle)
                                for cub_index, meas_cubicle in enumerate(cubicles):
                                    ct_list = self.interface.\
                                            get_cubicle_CT_of(meas_cubicle)
                                    if len(ct_list) > 0 :
                                        ct_ratio = self.interface.get_ct_ratio(ct_list[0])
                                        meas_relays = self.interface.\
                                                    get_cubicle_relay_of(meas_cubicle)
                                        if len(meas_relays) > 0:
                                            measurement = self.interface.\
                                                    get_relay_measurement_element(meas_relays[0])
                                            new_values_v1 = [sec_meas * ct_ratio * \
                                            trafo_voltages[cub_index] / trafo_voltages[0] for \
                                                        sec_meas in self.interface.\
                                                    get_relay_current_measures_of\
                                                                    (measurement)]
                                            new_values = [sec_meas * ct_ratio  \
                                                for sec_meas in self.interface.\
                                                    get_relay_current_measures_of\
                                                                    (measurement)]
                                            ikss += new_values
                                            ikss_v1 += new_values_v1
                                        else:
                                            self.interface.print("    No measurement element has been found")
                                            ikss += [0, 0, 0, 0]
                                    else:  # no CT add empty values
                                        ikss += [0, 0, 0, 0]
                                        ikss_v1 += [0, 0, 0, 0]
                                if len(cubicles) == 2:
                                    ikss += [0, 0, 0, 0, 0, 0, 0, 0]
                                    ikss_v1 += [0, 0, 0, 0, 0, 0, 0, 0]
                                elif len(cubicles) == 3:
                                    ikss += [0, 0, 0, 0]
                                    ikss_v1 += [0, 0, 0, 0]
                        except Exception as e:
                            self.interface.print("    Exception Getting SHC results!")
                            pass
                else:
                    for i in range(len(fault_types)):
                        ikss += [0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0]
                        ikss_v1 += [0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0]
            # the number of values read
            num_of_values = 4  # the 3 phase currents + zero seq
            # draw the TCD diagrams
            sides = ['(HV)', ' (LV)', ' (MV)']
            for fault_winding in range(0, len(cubicles)):
                offset_fault_win = fault_winding * len(fault_types) * num_of_values * 4
                for fault_type in fault_types:
                    fault_index = fault_types.index(fault_type)
#                     fault_cubicle_offset = len(fault_types) * 16
                    toc_diagram.append(self.interface.create_time_current_diagram("TCC-TR " + \
                        fault_type + ' ' + self.interface.get_name_of(transformer), items))
                    # offset for the In
                    offset = 0
                    diagram_type = 'Phase Relays'
                    if 'G' in fault_type or '1PH' in fault_type:
                        diagram_type = 'Earth Relays'
                        offset = num_of_values - 1
                    self.interface.set_time_current_diagram_type(toc_diagram[len(toc_diagram) - 1], \
                                                                  diagram_type)

                    # store the values for the vertical lines
                    for side_index, side in enumerate(sides):
                        if len(ikss) > offset_fault_win + \
                        fault_index * num_of_values * 4 + \
                         side_index * num_of_values:
                            trafo_iscc_data.append(Trafo_iscc_value(label=fault_type + side,
                                            diagram=toc_diagram[len(toc_diagram) - 1],
                                            iscc_value_at_v1=ikss_v1[ offset_fault_win + \
                                            fault_index * num_of_values * 4 + \
                                            offset + side_index * num_of_values],
                                            iscc_value=ikss[ offset_fault_win + \
                                            fault_index * num_of_values * 4 + \
                                            offset + side_index * num_of_values]))

#                 self.interface.print("    Saving SHC results")
            # write the calculated I in the trafo_sheet
            if len(ikss) > 0:
                self.fill_trafo_output_file(ikss, \
                                  self.interface.get_name_of(transformer), \
                                  trafo_sheet,
                                  start_index, cubicles)
#             else:
#                 if self.output_detail >= OutputDetail.NORMAL:
#                     self.interface.print("    No SHC data")
            # len(ikssI) == 0, means no relay available, skip it!
            if self.interface.is_energized(fault_bus_bar) and\
                    self.interface.is_out_of_service(transformer) == False and\
                    len(ikss) > 0:
                start_index += len(fault_types) * len(cubicles)

#             if start_index > 5:
#                 break

        # draw the vertical lines
        self.interface.disable_current_study_case()
        for data in trafo_iscc_data:
            self.interface.create_TOC_diagram_verticalline(data.label, \
                                                data.diagram, data.iscc_value_at_v1)
        self.interface.enable_last_enabled_study_case()
        self.interface.deactivate_project()

        # output_file.close()
        self.interface.enable_pf_gui_update()
        self.interface.set_echo_on()
        if self.output_detail >= OutputDetail.NORMAL:
            self.interface.print("\nTC diagrams creation completed!")

    def create_trafo_output_file_header(self, window):
        '''
        create the output file where the shc values and the other elaboration results
        for the trafos are saved
        '''
        if self.workbook == None:
            self.create_copy_of(window.results_file_name.GetValue())
            self.workbook = Workbook(window.results_file_name.GetValue())
        cell_bold_format = self.workbook.add_format({'bold': True})
        # add the network summary sheet
        trafo_shc_worksheet = self.workbook.add_worksheet("Trafo SHC")

        # write the summary header
        trafo_shc_worksheet.set_column(0, 0, 35)
        trafo_shc_worksheet.set_column(1, 14, 20)
        trafo_shc_worksheet.write('A1' , "Trafo name", cell_bold_format)
        trafo_shc_worksheet.write('B1' , "Fault Type", cell_bold_format)
        trafo_shc_worksheet.write('C1' , "Fault Winding #", cell_bold_format)
        trafo_shc_worksheet.write('D1' , "IA1 (A)", cell_bold_format)
        trafo_shc_worksheet.write('E1' , "IB1 (A)", cell_bold_format)
        trafo_shc_worksheet.write('F1' , "IC1 (A)", cell_bold_format)
        trafo_shc_worksheet.write('G1' , "IN1 (A)", cell_bold_format)
        trafo_shc_worksheet.write('H1' , "IA2 (A)", cell_bold_format)
        trafo_shc_worksheet.write('I1' , "IB2 (A)", cell_bold_format)
        trafo_shc_worksheet.write('J1' , "IC2 (A)", cell_bold_format)
        trafo_shc_worksheet.write('K1' , "IN2 (A)", cell_bold_format)
        trafo_shc_worksheet.write('L1' , "IA3 (A)", cell_bold_format)
        trafo_shc_worksheet.write('M1' , "IB3 (A)", cell_bold_format)
        trafo_shc_worksheet.write('N1' , "IC3 (A)", cell_bold_format)
        trafo_shc_worksheet.write('O1' , "IN3 (A)", cell_bold_format)
        return trafo_shc_worksheet

    def create_shunt_output_file_header(self, window):
        '''
        create the output file where the shc values and the other elaboration results
        for the shunts are saved
        '''
        if self.workbook == None:
            self.create_copy_of(window.results_file_name.GetValue())
            self.workbook = Workbook(window.results_file_name.GetValue())
        cell_bold_format = self.workbook.add_format({'bold': True})
        # add the network summary sheet
        shunt_shc_worksheet = self.workbook.add_worksheet("Shunt SHC")

        # write the summary header
        shunt_shc_worksheet.set_column(0, 0, 35)
        shunt_shc_worksheet.set_column(1, 14, 20)
        shunt_shc_worksheet.write('A1' , "Shunt name", cell_bold_format)
        shunt_shc_worksheet.write('B1' , "Fault Bus", cell_bold_format)
        shunt_shc_worksheet.write('C1' , "Fault Type", cell_bold_format)
        shunt_shc_worksheet.write('D1' , "IA (A)", cell_bold_format)
        shunt_shc_worksheet.write('E1' , "Phi A (deg)", cell_bold_format)
        shunt_shc_worksheet.write('F1' , "IB (A)", cell_bold_format)
        shunt_shc_worksheet.write('G1' , "Phi B (deg)", cell_bold_format)
        shunt_shc_worksheet.write('H1' , "IC (A)", cell_bold_format)
        shunt_shc_worksheet.write('I1' , "Phi C (deg)", cell_bold_format)
        shunt_shc_worksheet.write('J1' , "IN (A)", cell_bold_format)
        shunt_shc_worksheet.write('K1' , "Phi N (deg)", cell_bold_format)
        return shunt_shc_worksheet

    def create_z3_output_file_header(self, window, branch_type):
        '''
        create the output file where the shc values and the other elaboration results
        for the z3 calculation are saved
        note; the 'branch_type' input parameter is a string which as added in 
        the sheet name
        '''
        if self.workbook == None:
            self.create_copy_of(window.results_file_name.GetValue())
            self.workbook = Workbook(window.results_file_name.GetValue())
        cell_bold_format = self.workbook.add_format({'bold': True})
        # add the network summary sheet
        z3_shc_worksheet = self.workbook.add_worksheet("Z3 SHC (" + \
                                                        branch_type + ")")

        # write the summary header
        z3_shc_worksheet.set_column(0, 0, 35)
        z3_shc_worksheet.set_column(1, 1000, 25)
        z3_shc_worksheet.write('A1' , "Relay name", cell_bold_format)
        z3_shc_worksheet.write('B1' , "Fault Bus", cell_bold_format)
        z3_shc_worksheet.write('C1' , "Fault Type", cell_bold_format)
        z3_shc_worksheet.write('D1' , "ZA", cell_bold_format)
        z3_shc_worksheet.write('E1' , "Phi ZA", cell_bold_format)
        z3_shc_worksheet.write('F1' , "ZB", cell_bold_format)
        z3_shc_worksheet.write('G1' , "Phi ZB", cell_bold_format)
        z3_shc_worksheet.write('H1' , "ZC", cell_bold_format)
        z3_shc_worksheet.write('I1' , "Phi ZC", cell_bold_format)

        max_number_of_taps = 100
        for i in range (0, max_number_of_taps):
            z3_shc_worksheet.write(0, self._calculate_tap_values_column(i, 0, \
                                start_column='J', number_of_values=7), \
                                      "Overtap Fault Bus " + str(i + 1), cell_bold_format)
            z3_shc_worksheet.write(0, self._calculate_tap_values_column(i, 1, \
                                start_column='J', number_of_values=7), \
                                       "ZA " + str(i + 1), cell_bold_format)
            z3_shc_worksheet.write(0, self._calculate_tap_values_column(i, 2, \
                                start_column='J', number_of_values=7), \
                                       "Phi ZA " + str(i + 1), cell_bold_format)
            z3_shc_worksheet.write(0, self._calculate_tap_values_column(i, 3, \
                                start_column='J', number_of_values=7), \
                                       "ZB " + str(i + 1), cell_bold_format)
            z3_shc_worksheet.write(0, self._calculate_tap_values_column(i, 4, \
                                start_column='J', number_of_values=7), \
                                       "Phi ZB " + str(i + 1), cell_bold_format)
            z3_shc_worksheet.write(0, self._calculate_tap_values_column(i, 5, \
                                start_column='J', number_of_values=7), \
                                       "ZC " + str(i + 1), cell_bold_format)
            z3_shc_worksheet.write(0, self._calculate_tap_values_column(i, 6, \
                                start_column='J', number_of_values=7), \
                                       "Phi ZC " + str(i + 1), cell_bold_format)
        return z3_shc_worksheet

    def create_generator_output_file_header(self, window):
        '''
        create the output file where the shc values and the other elaboration results
        for the generators are saved
        '''
        if self.workbook == None:
            self.create_copy_of(window.results_file_name.GetValue())
            self.workbook = Workbook(window.results_file_name.GetValue())
        cell_bold_format = self.workbook.add_format({'bold': True})
        # add the network summary sheet
        generator_shc_worksheet = self.workbook.add_worksheet("Generator SHC")

        # write the summary header
        generator_shc_worksheet.set_column(0, 0, 35)
        generator_shc_worksheet.set_column(1, 92, 25)
        generator_shc_worksheet.write('A1' , "Generator name", cell_bold_format)
        generator_shc_worksheet.write('B1' , "Fault Bus", cell_bold_format)
        generator_shc_worksheet.write('C1' , "Fault Type", cell_bold_format)
        generator_shc_worksheet.write('D1' , "ZA", cell_bold_format)
        generator_shc_worksheet.write('E1' , "Phi ZA", cell_bold_format)
        generator_shc_worksheet.write('F1' , "ZB", cell_bold_format)
        generator_shc_worksheet.write('G1' , "Phi ZB", cell_bold_format)
        generator_shc_worksheet.write('H1' , "ZC", cell_bold_format)
        generator_shc_worksheet.write('I1' , "Phi ZC", cell_bold_format)

        max_number_of_taps = 6
        for i in range (max_number_of_taps):
            generator_shc_worksheet.write(0, self._calculate_tap_values_column(i, 0, \
                                start_column='J', number_of_values=7), \
                                      "Overtap Bus " + str(i + 1), cell_bold_format)
            generator_shc_worksheet.write(0, self._calculate_tap_values_column(i, 1, \
                                start_column='J', number_of_values=7), \
                                       "ZA " + str(i + 1), cell_bold_format)
            generator_shc_worksheet.write(0, self._calculate_tap_values_column(i, 2, \
                                start_column='J', number_of_values=7), \
                                       "Phi ZA " + str(i + 1), cell_bold_format)
            generator_shc_worksheet.write(0, self._calculate_tap_values_column(i, 3, \
                                start_column='J', number_of_values=7), \
                                       "ZB " + str(i + 1), cell_bold_format)
            generator_shc_worksheet.write(0, self._calculate_tap_values_column(i, 4, \
                                start_column='J', number_of_values=7), \
                                       "Phi ZB " + str(i + 1), cell_bold_format)
            generator_shc_worksheet.write(0, self._calculate_tap_values_column(i, 5, \
                                start_column='J', number_of_values=7), \
                                       "ZC " + str(i + 1), cell_bold_format)
            generator_shc_worksheet.write(0, self._calculate_tap_values_column(i, 6, \
                                start_column='J', number_of_values=7), \
                                       "Phi ZC " + str(i + 1), cell_bold_format)
        return generator_shc_worksheet

    def create_line_output_file_header(self, window):
        '''
        create the output file where the shc values and the other elaboration results
        for the lines are saved
        '''
        if self.workbook == None:
            self.create_copy_of(window.results_file_name.GetValue())
            self.workbook = Workbook(window.results_file_name.GetValue())
        cell_bold_format = self.workbook.add_format({'bold': True})
        # add the network summary sheet
        trafo_shc_worksheet = self.workbook.add_worksheet("Line SHC")

        # write the summary header
        trafo_shc_worksheet.set_column(1, 90, 20)
        trafo_shc_worksheet.set_column(0, 0, 35)
        trafo_shc_worksheet.set_column(12, 12, 35)

        trafo_shc_worksheet.write('A1' , "Line name", cell_bold_format)
        trafo_shc_worksheet.write('B1' , "Fault Type", cell_bold_format)
        trafo_shc_worksheet.write('C1' , "Fault Position", cell_bold_format)
        trafo_shc_worksheet.write('D1' , "IA1", cell_bold_format)
        trafo_shc_worksheet.write('E1' , "IB1", cell_bold_format)
        trafo_shc_worksheet.write('F1' , "IC1", cell_bold_format)
        trafo_shc_worksheet.write('G1' , "I0x31", cell_bold_format)
        trafo_shc_worksheet.write('H1' , "IA2", cell_bold_format)
        trafo_shc_worksheet.write('I1' , "IB2", cell_bold_format)
        trafo_shc_worksheet.write('J1' , "IC2", cell_bold_format)
        trafo_shc_worksheet.write('K1' , "I0x32", cell_bold_format)
        max_number_of_taps = 6
        for i in range (max_number_of_taps):
            trafo_shc_worksheet.write(0, self._calculate_tap_values_column(i, 0), \
                                      "Overtap bus " + str(i + 1), cell_bold_format)
            trafo_shc_worksheet.write(0, self._calculate_tap_values_column(i, 1), \
                                       "Overtap IA1 " + str(i + 1), cell_bold_format)
            trafo_shc_worksheet.write(0, self._calculate_tap_values_column(i, 2), \
                                       "Overtap IB1 " + str(i + 1), cell_bold_format)
            trafo_shc_worksheet.write(0, self._calculate_tap_values_column(i, 3), \
                                       "Overtap IC1 " + str(i + 1), cell_bold_format)
            trafo_shc_worksheet.write(0, self._calculate_tap_values_column(i, 4), \
                                       "Overtap I0x31 " + str(i + 1), cell_bold_format)
            trafo_shc_worksheet.write(0, self._calculate_tap_values_column(i, 5), \
                                       "Overtap IA2 " + str(i + 1), cell_bold_format)
            trafo_shc_worksheet.write(0, self._calculate_tap_values_column(i, 6), \
                                       "Overtap IB2 " + str(i + 1), cell_bold_format)
            trafo_shc_worksheet.write(0, self._calculate_tap_values_column(i, 7), \
                                       "Overtap IC2 " + str(i + 1), cell_bold_format)
            trafo_shc_worksheet.write(0, self._calculate_tap_values_column(i, 8), \
                                       "Overtap I0x32 " + str(i + 1), cell_bold_format)

        return trafo_shc_worksheet

###############################################################################
###############################################################################
###############################################################################
    def activate_variants(self, variants):
        for variant in variants:
            variation_obj = self.interface.get_variation(variant) 
            if variation_obj == 1: 
                self.interface.enable_pf_gui_update()
                self.interface.set_echo_on()  
                self.interface.print("ERROR: '" + variant + 
                         "' variation doesn't exist !!!")
                continue
            self.interface.set_echo_on()
            self.interface.enable_pf_gui_update()
            self.interface.print("Activating variation '" + 
                                 variant + "'")
            self.interface.disable_pf_gui_update()
            self.interface.set_echo_off()
            self.interface.activate_variation(variation_obj)
            
    def deactivate_variants(self, variants):
        for variant in variants:
            variation_obj = self.interface.get_variation(variant) 
            if variation_obj == None: 
                self.interface.enable_pf_gui_update()
                self.interface.set_echo_on()  
                self.interface.print("ERROR: '" + variant + 
                         "' variation doesn't exist !!!")
                continue
            self.interface.set_echo_on()
            self.interface.enable_pf_gui_update()
            self.interface.print("Deactivating variation '" + 
                                 variant + "'")
            self.interface.disable_pf_gui_update()        
            self.interface.set_echo_off()
            self.interface.deactivate_variation(variation_obj)
            
    def create_results_header(self, file_name):    
        if self.workbook == None:
            self.workbook = Workbook(file_name)
        cell_bold_format = self.workbook.add_format({'bold': True})
        # add the network summary sheet
        results_worksheet = self.workbook.add_worksheet("Results")

        # write the summary header
        results_worksheet.set_column(0, 2, 12)
        results_worksheet.set_column(3, 3, 42)
        results_worksheet.set_column(4, 4, 15)
        results_worksheet.set_column(5, 5, 28)
        results_worksheet.set_column(6, 6, 40)
        results_worksheet.set_column(7, 7, 32)
        results_worksheet.set_column(8, 8, 45)
        results_worksheet.set_column(9, 9, 28)
        results_worksheet.set_column(10, 10, 45)
        results_worksheet.set_column(11, 11, 20)
        results_worksheet.write('A1' , "Configuration", cell_bold_format)
        results_worksheet.write('B1' , "Line ID", cell_bold_format)
        results_worksheet.write('C1' , "Line Name", cell_bold_format)
        results_worksheet.write('D1' , "Line Type", cell_bold_format)
        results_worksheet.write('E1' , "Line Length (km)", cell_bold_format)
        results_worksheet.write('F1' , self.projects[0] + " P (MW)", cell_bold_format)
        results_worksheet.write('G1' , self.projects[0] + " Capitalized cost(MUSD) ", cell_bold_format)
        results_worksheet.write('H1' , self.projects[1] + " P (MW)", cell_bold_format)
        results_worksheet.write('I1' , self.projects[1] + " Capitalized cost(MUSD) ", cell_bold_format)
        results_worksheet.write('J1' , self.projects[2] + " P (MW)", cell_bold_format)
        results_worksheet.write('K1' , self.projects[2] + " Capitalized cost(MUSD) ", cell_bold_format)
        results_worksheet.write('L1' , "Weighted Cost (MUSD)", cell_bold_format)
        self.workbook.close()
        
            

    def calculate_lines_optimal_types(self):
       
        #definitions...
        self.projects = ["Wet Peak Adjust 28-29", "Wet Off Peak Adjust 28-29", 
                        "DRY Peak Adjust 28-29"]
        self.project_weights = [0.0625, 0.4375, 0.5 ]
        configurations = ["1", "2A", "2B"]
        self.lines = ["TL1", "TL2", "TL3", "TL4"]
        self.line_limits = [1120, 1368, 1019, 2038, 1243, 2486, 4972]
        self.line_types = {"220-1": "220 kV Double Circuit 2*Bison Transmission Line",
                      "220-2": "220 kV Double Circuit 2*Moose Transmission Line",
                      "400-1": "400 kV s/c ACSR 2* Bison Transmission Line",
                      "400-2": "400 kV Double Circuit 2*Bison Transmission Line",
                      "400-3": "400 kV s/c ACSR 2*Moose Transmission Line",
                      "400-4": "400 kV Double Circuit 2*Moose transmission Line",
                      "400-5": "400 kV Double Circuit 4* Moose Transmission Line"}
        template_sheets = {"220-1": "Template 220 kV Bison D2",
                           "220-2": "Template 220 kV Mose D2",
                           "400-1": "Template 400 kV Bison 2",
                           "400-2": "Template 400 kV Bison D2 ",
                           "400-3": "Template 400 kV Moose 2 ",
                           "400-4": "Template 400 kV Moose D2 ",
                           "400-5": "Template 400 kV Moose D4"}
        result_row = {"220-1": "61",
                      "220-2": "61",
                      "400-1": "59",
                      "400-2": "59",
                      "400-3": "59",
                      "400-4": "59",
                      "400-5": "59"}
        self.trafo_keys = {"TL1-220": ["", ""],
                      "TL1-400": ["Arun-400-220", "Inaruwa-400-220"],
                      "TL2-1-220": ["Tingla-220-132", "Dudhkoshi-220-132"],
                      "TL2-2A-220": ["Tingla-220-132", "Dudhkoshi-220-132"],
                      "TL2-2B-220": ["Tingla-220-132", "Dudhkoshi2-220-132"],
                      "TL2-1-400": ["Tingla-400-132", "Dudhkoshi-400-132"],
                      "TL2-2A-400": ["Tingla-400-132", "Dudhkoshi-400-132"],
                      "TL2-2B-400": ["Tingla-400-132", "Dudhkoshi2-400-132"],                     
                      "TL3-1-220": ["", "Dudhkoshi-220-132"],
                      "TL3-2A-220": ["", "Tingla-220-132"],
                      "TL3-2B-220": ["", "Tingla-220-132"],
                      "TL3-1-400": ["Arun-400-220", "Dudhkoshi-400-132"],
                      "TL3-2A-400": ["Arun-400-220", "Tingla-400-132"],
                      "TL3-2B-400": ["Arun-400-220", "Tingla-400-132"],                     
                      "TL4-220": ["Anarmani-220-132", ""],
                      "TL4-400": ["Anarmani-400-132", "Inaruwa-400-220"]               
                      }
        self.all_line_keys = ["TL1-220",
                          "TL1-400",
                          "TL2-1-220",
                          "TL2-2A-220",
                          "TL2-2B-220",
                          "TL2-1-400",
                          "TL2-2A-400",
                          "TL2-2B-400",                     
                          "TL3-1-220",
                          "TL3-2A-220",
                          "TL3-2B-220",
                          "TL3-1-400",
                          "TL3-2A-400",
                          "TL3-2B-400",                     
                          "TL4-220",
                          "TL4-400"               
                      ]
        self.trafo_powers = {"220": 167,
                             "400": 315}
        self.energy_costs = {"DRY": 104.16,
                             "Wet": 59.64}
        self.cost_row = {"220-1": "18",
                      "220-2": "18",
                      "400-1": "16",
                      "400-2": "16",
                      "400-3": "16",
                      "400-4": "16",
                      "400-5": "16"}
        
        report_path = "C:\\Users\\Alberto's laptop\\Documents\\ELC - ElectroConsult\\Nepal eastern region evacuation study\\Phase 2\\Report"
        # path to the template prototypes
        source_template_path = "C:\\Users\\Alberto's laptop\\Documents\\ELC - ElectroConsult\\Nepal eastern region evacuation study\\Phase 2\\Templates"
        # add the actual date in the path
        now = datetime.datetime.now()
        date_str = now.strftime("%Y %m %d, %H-%M")
        report_path += " " + date_str
        # create the result report main dir
        if not os.path.exists(report_path):
            os.mkdir(report_path)
        
        base_template_path = report_path + "\\Templates"
        # create the result template dir
        if not os.path.exists(base_template_path):
            os.mkdir(base_template_path)
        #load excel
        self.excel = win32.gencache.EnsureDispatch('Excel.Application')       
        self.interface.deactivate_project()
        # creater the xls result file header
        self.result_file_name = report_path + "\\Results.xlsx";
        self.create_results_header(self.result_file_name)
        
        self.workbook = load_workbook(self.result_file_name)
        
        # create the file to store the temporary results
        tmp_results_file = open(report_path + "\\tmp_results.txt",'w')
        
        # iterate between projects
        for project_index, project in enumerate(self.projects):
            change_project = False
            self.output_row_index = 2
            # the 4 lines initial variants
            active_variants = ["", "", "", ""]
            
            self.interface.activate_project(project)
            self.interface.print("Activating " + project + " project\n")
            
            #create the line objects array
            self.line_types_objects = []
            for tl_type in self.line_types.keys():
                self.line_types_objects.append(self.
                                        interface.get_element_by_foreign_key(
                                            "T-" + tl_type))
            
            
            # activate the 4 lines initial variants
            #self.activate_variants(active_variants)   
            # make the project report directory 
            if not os.path.exists(report_path + "\\" + project):
                os.mkdir(report_path + "\\" + project)
            
            # iterate throw the configurations
            for config_index, configuration in enumerate(configurations):
                self.interface.enable_pf_gui_update()
                self.interface.set_echo_on()
                self.interface.print("   Using " + configuration +
                                  " configuration\n")
                # make the configuration report directory
                final_report_path = report_path + "\\" + project + \
                                    "\\Configuration " +configuration
                if not os.path.exists(final_report_path):
                    os.mkdir(final_report_path)
                
                #make a copy of the template file
                template_path = base_template_path + "\\Scenario " + configuration
                
                # iterate throw the lines
                # TL1
                self.interface.set_echo_off()
                # put out of service all lines
                self.disable_all_lines()
                try:
                    #debug
                    max_num_config = 4
                    config_number = 0
                    for line_key1, line_type1 in self.line_types.items():
                        self.interface.disable_pf_gui_update()
                        variation_name = "TL1-" + line_key1
                        if active_variants[0][:-2] != variation_name[:-2]:
                            if config_index > 0:
                                continue 
                            self.set_variant(variation_name)
                            self.disable_line(active_variants[0])
                            active_variants[0] = variation_name
                        else:
                            self.enable_line(variation_name)
                            active_variants[0] = variation_name
                        for line_key2, line_type2 in self.line_types.items():
                            variation_name = "TL2-" + configuration + "-" + line_key2
                            if active_variants[1][:-2] != variation_name[:-2]:
                                self.set_variant(variation_name)
                                self.disable_line(active_variants[1])
                                active_variants[1] = variation_name
                            else:
                                self.enable_line(variation_name)
                                active_variants[1] = variation_name
                            for line_key3, line_type3 in self.line_types.items():
                                variation_name = "TL3-" + configuration + "-" + line_key3
                                if active_variants[2][:-2] != variation_name[:-2]:
                                    self.set_variant(variation_name)
                                    self.disable_line(active_variants[2])
                                    active_variants[2] = variation_name
                                else:
                                    self.enable_line(variation_name)
                                    active_variants[2] = variation_name
                                for line_key4, line_type4 in self.line_types.items():
                                    variation_name = "TL4-" + line_key4
                                    if active_variants[3][:-2] != variation_name[:-2]:
                                        if config_index > 0:
                                            continue
                                        self.set_variant(variation_name)
                                        self.disable_line(active_variants[3])
                                        active_variants[3] = variation_name
                                    else:
                                        self.enable_line(variation_name)
                                        active_variants[3] = variation_name
                                    self.interface.enable_pf_gui_update()
                                    self.interface.set_echo_on()
                                    self.interface.print("Running lDF")
                                    self.interface.set_echo_off()
                                    self.interface.disable_pf_gui_update()
                                    self.interface.run_ldf()
                                    self.interface.enable_pf_gui_update()
                                    self.interface.set_echo_on()
                                    
                                    config_number += 1
                                    if self.interface.is_ldf_valid():
                                        self.interface.print("Variation: " + str(variation_name))
                                        self.interface.print("lDF calculation completed successfully")
                                        
                                        #"\nNumber of measurement items: {}\n".\
                                        #format(number_of_measurement_items)
                    
                                        # create the line types dir
                                        line_types_string = active_variants[0] +\
                                                    " " + active_variants[1] +\
                                                    " " + active_variants[2] +\
                                                    " " + active_variants[3]
                                        out_path = final_report_path + "\\" +\
                                        line_types_string
                                        if not os.path.exists(out_path):
                                            os.mkdir(out_path) 
                                        self.line_keys = [line_key1,
                                                      line_key2,
                                                      line_key3,
                                                      line_key4]
                                        for line_index, line in enumerate(self.lines):
                                            #self.interface.print("Debug 4")
                                            line_key = line + "-"
                                            if line != "TL1" and line != "TL4":
                                                line_key += configuration + "-"
                                                                                    
                                            #find the file
                                            template_list = list(Path(source_template_path
                                                + "\\Scenario " + configuration).
                                                                glob(line + '*')) 
                                            #create the copy
                                            calculation_file = copy2(template_list[0], 
                                                final_report_path + "\\" + line_types_string)
                                            # read the copied file
                                            workbook = load_workbook(calculation_file)
                                            summary_worksheet = workbook["Template summary"]
                                            #disable the not used sheets
                                            #summary_worksheet.activate()
                                            all_sheet_names  = workbook.get_sheet_names()
                                            for sheet_name in all_sheet_names:
                                                if  template_sheets[self.line_keys[line_index]]\
                                                            not in sheet_name and\
                                                     "Template summary" not in sheet_name:
                                                    workbook.remove_sheet(
                                                        workbook.get_sheet_by_name(sheet_name))     
                                                
                                            line_obj = self.interface.get_element_by_foreign_key(
                                                        active_variants[line_index][:-2])
                                            # get the line active power
                                            if line_obj is not None:
                                                self.interface.print("Getting '" + active_variants[line_index][:-2] + "'")
                                                line_power = self.interface.get_attribute(line_obj, "m:P:bus1")
                                                line_length = self.interface.get_attribute(line_obj, "b:dline") 
                                                # write the power value in the cell
                                                summary_worksheet.cell(row=4, column=3).\
                                                            value =\
                                                round(abs(line_power), 2)
                                                # write the number of transformers
                                                voltage = "220" if "220" in\
                                                 template_sheets[self.line_keys[line_index]] else "400"
                                                summary_worksheet.cell(row=1, column=6).\
                                                            value =\
                                                int(
                                                abs(line_power)/self.trafo_powers[voltage]) + 1
                                                # write the cost of energy
                                                calculation_sheet = workbook[\
                                                template_sheets[self.line_keys[line_index]]]
                                                cost = self.energy_costs["DRY"]\
                                                if "DRY" in project else\
                                                    self.energy_costs["Wet"] 
                                                cost_row = self.cost_row[self.line_keys[line_index]]
                                                calculation_sheet.cell(row=\
                                                    int(cost_row), column=6).\
                                                            value = cost
                                                workbook.save(calculation_file)
                                                if line_power is None:
                                                    self.interface.print(\
                                                    "Cannot read line power value")
                                                    sys.exit()
                                            else:
                                                self.interface.enable_pf_gui_update()
                                                self.interface.set_echo_on()
                                                self.interface.print("ERROR: " + 
                                                            active_variants[line_index] +
                                                            " foreign key no found")
                                            # get the result from the sheet
                                            workbook = self.excel.Workbooks.Open(calculation_file)
                                            workbook.Save()
                                            workbook.Close()
                        
                                            workbook = load_workbook(calculation_file, data_only=True)
                                            sheet = workbook[template_sheets
                                                    [self.line_keys[line_index]]]
                                            calculated_value = sheet["G" + 
                                                result_row[self.line_keys[line_index]]].value
                                            self.interface.print
                                            ("Calculated value" + str(calculated_value))    
                                            # here the part to save the data
                                            self.save_line_data(project_index,
                                                            configuration,
                                                            active_variants[line_index],
                                                            line_index,
                                                            line_length,
                                                            round(abs(line_power), 2),
                                                            round(calculated_value,2))
                                            tmp_results_file.write\
                                                (str(self.output_row_index));
                                        if config_number > max_num_config:
                                            change_project = True
                                            break
                                    if change_project:
                                        break
                                if change_project:
                                        break
                            if change_project:
                                        break
                        if change_project:
                                        break
                    if change_project:
                                        break                                  
                except KeyboardInterrupt:
                    self.interface.set_echo_on()
                    if self.output_detail >= OutputDetail.NORMAL:
                        self.interface.print("\nInterrupted by the user!")
                        
            self.interface.deactivate_project()
        self.calculate_total_costs()
        self.save_best_configuration()
        return 0
    
    def disable_all_lines(self):
        '''
        put out of service all lines listed in line_keys
        '''
        for line in self.all_line_keys:
            self.disable_line(line)
        
    def disable_line(self, line_key):
        '''
        put out of service the given line
        '''
        if len(line_key) > 0:
            if line_key[-2] == '-':
                line_key = line_key[:-2]
            line_obj = self.interface.get_element_by_foreign_key(line_key)
            if line_obj is not None:
                # deactivate the line
                self.interface.set_attribute(line_obj, "outserv", 1) 
                #self.interface.print("Disabling '" + str(line_key) + "'")
            else:
                self.interface.enable_pf_gui_update()
                self.interface.set_echo_on()
                self.interface.print("'" + line_key + "' key doesn't exits!")
                #sys.exit()
    def enable_line(self, line_key):
        '''
        put in service the given line
        '''
        if len(line_key) > 0:
            if line_key[-2] == '-':
                line_key = line_key[:-2]
            line_obj = self.interface.get_element_by_foreign_key(line_key)
            if line_obj is not None:
                # deactivate the line
                self.interface.set_attribute(line_obj, "outserv", 0)
                #self.interface.print("Enabling '" + str(line_key) + "'") 
            else:
                self.interface.enable_pf_gui_update()
                self.interface.set_echo_on()
                self.interface.print("'" + line_key + "' key doesn't exits!")
                #sys.exit()
        
    def set_variant(self, variation_name):
                            
        self.config_line_type(variation_name)
        self.config_transformer_number(variation_name)
    
    def config_line_type(self, variation_name):
        '''
        function setting the line type
        '''
        # get the line object using the foreign key
        index = int(variation_name[-1])
        # the 4000 line types start from the third element
        if "400" in variation_name:
            index += 2
        variation_name = variation_name[:-2]        
        line_obj = self.interface.get_element_by_foreign_key(variation_name)
        if line_obj is not None:
            # activate the line
            self.interface.set_attribute(line_obj, "outserv", 0) 
            # set the line type
            self.interface.set_attribute(line_obj, "typ_id", 
                                         self.line_types_objects[index-1])
        else:
            self.interface.enable_pf_gui_update()
            self.interface.set_echo_on()
            self.interface.print("'" + variation_name + "' key doesn't exits!")
            #sys.exit()
    
    def config_transformer_number(self, variation_name):
        '''
        function setting the number of parallel transformers
        '''
        index = int(variation_name[-1])
        for key in self.trafo_keys[variation_name[:-2]]:  
            if len(key) > 0:         
                trafo_obj = self.interface.get_element_by_foreign_key(key)
                if trafo_obj is not None:
                    # activate the transformer
                    self.interface.set_attribute(trafo_obj, "outserv", 0)
                    # set the number of parallel trafos
                    trafo_rated_power = self.trafo_powers["400"]\
                                        if "400" in variation_name\
                                        else self.trafo_powers["220"]
                    number_of_trafos =\
                    int(self.line_limits[index]/trafo_rated_power) + 1
                    self.interface.set_attribute(trafo_obj, "ntnum", number_of_trafos)
                else:
                    self.interface.enable_pf_gui_update()
                    self.interface.set_echo_on()
                    self.interface.print("'" + variation_name + "' key doesn't exits!")
                    #sys.exit() 
    
    def calculate_total_costs(self):
        '''
        function calculating the total costs in the result file
        '''  
        results_worksheet = self.workbook["Results"]
        row_index = 2
        number_of_projects = 3
        number_of_lines = 4
        # varibles for the final report
        self.min_total_cost = 9999999;
        self.min_cost_row = 0
        while row_index < self.output_row_index:
            total_cost = 0
            for line_number in range(number_of_lines):
                line_total_cost = 0
                for project_number in range(number_of_projects): 
                    project_weight = float(self.project_weights[project_number])
                    line_cost = float(results_worksheet.cell(\
                                    row_index, 7 + project_number * 2).value)
                    line_total_cost += line_cost * project_weight
                # write the total cost for the line
                results_worksheet.\
                    cell(row_index, 6 + number_of_projects * 2).value =\
                                                                line_total_cost
                total_cost += line_total_cost
                row_index += 1
            # write the total cost 
            results_worksheet.\
                cell(row_index, 6 + number_of_projects * 2).value =\
                                                                total_cost
            if total_cost < self.min_total_cost:
                self.min_total_cost = total_cost
                self.min_cost_row = row_index
            row_index += 1
        self.workbook.save(self.result_file_name)
    
    def save_best_configuration(self):
        '''
        save in a xls file the best configuration
        '''
        results_worksheet = self.workbook["Results"]
        
        best_result_file_name = self.result_file_name.replace(\
                                            "\\Results.xlsx", "\\Best_result.xlsx");
        # create the best results file
        self.best_results_workbook = Workbook(best_result_file_name)
        # access the results sheet
        cell_bold_format = self.best_results_workbook.add_format({'bold': True})
        best_results_worksheet = self.best_results_workbook.add_worksheet("Best_config")
        
        best_results_worksheet.set_column(0, 2, 12)
        best_results_worksheet.set_column(3, 3, 42)
        best_results_worksheet.set_column(4, 4, 15)
        best_results_worksheet.set_column(5, 5, 28)
        best_results_worksheet.set_column(6, 6, 40)
        best_results_worksheet.set_column(7, 7, 32)
        best_results_worksheet.set_column(8, 8, 45)
        best_results_worksheet.set_column(9, 9, 28)
        best_results_worksheet.set_column(10, 10, 45)
        best_results_worksheet.set_column(11, 11, 20)
        number_of_columns = 12
        #copy the first line
        for column in range(number_of_columns):
            value = results_worksheet.cell(1, column + 1).value
            best_results_worksheet.write(chr(ord('A') + column) + '1' , 
                value, cell_bold_format)
        #copy the data
        for row in range(5):
            for column in range(number_of_columns):
                best_results_worksheet.write(chr(ord('A') + column) + str(row+2), 
                results_worksheet.cell(
                self.min_cost_row - 4 + row, column + 1).value)
        
        self.best_results_workbook.close()
        
    def save_line_data(self, project_index, configuration, line_name, 
                       line_index, line_length, line_power, line_cost):
        '''
        save the line data in the result file
        '''    
        # add the network summary sheet
        results_worksheet = self.workbook["Results"]
 
        # write the summary header
        results_worksheet.cell(self.output_row_index, 1).value = configuration
        results_worksheet.cell(self.output_row_index, 2).value = self.lines[line_index]
        results_worksheet.cell(self.output_row_index, 3).value = line_name
        results_worksheet.cell(self.output_row_index, 4).value = self.line_types[
                                                    self.line_keys[line_index]]
        results_worksheet.cell(self.output_row_index, 5).value = line_length
        results_worksheet.cell(self.output_row_index, 6 + project_index*2).value = line_power
        results_worksheet.cell(self.output_row_index, 7 + project_index*2).value = line_cost
         
        self.output_row_index += 1
        # if it's the last of the 4 lines add a void row
        if line_index == 3:
            results_worksheet.cell(self.output_row_index, 11).value = "TOT"
            self.output_row_index += 1
        self.workbook.save(self.result_file_name)
###############################################################################
# auxiliary function
###############################################################################

    def _calculate_tap_values_column(self, tap_number, value_position,
                                    start_column='L', number_of_values=9):
        '''
        function calcaulating the excel number identifying a column in the shc
        value output for the busses from a tap assuming that the given number of
         values are displayed for each bus
        '''
        first_set_dim = ord(start_column) - ord('A')
        index = tap_number * number_of_values + value_position + first_set_dim
        return index

###############################################################################

    def get_workbook_and_worksheet(self, window):
        '''
        function getting the shc result files workbook and worksheet
        '''
        self.workbook = load_workbook(window.results_file_name.GetValue())
        trafo_shc_worksheet = self.workbook["Trafo SHC"]
        return trafo_shc_worksheet

    def create_copy_of(self, file_name, extension='.xls'):
        '''
        create a copy of the given faile name (if existing) adding the date/time in
        the file name
        '''
        import os.path
        from shutil import copyfile
        from datetime import datetime
        import time
        if os.path.isfile(file_name):
            # now = datetime.now()
            file_date_str = time.ctime(os.path.getctime(file_name))
            file_date_str = file_date_str.replace(':', '-')
            copyfile(file_name, file_name.replace(extension, file_date_str + extension))

    def save_diagrams_as_wmf(self, window, diagram_name='', path=None, \
                                    create_copy=True):
        '''
        save as a wmf file in the given path directory all graphical diagram
         with name containing or equal to diagram_name
        if no path is provided the path of the script output file is used
        if the create_copy parameter is True a copy of the file with the date/hour
        is created 
        '''
        dir_path = path if path != None else\
                os.path.split(window.results_file_name.GetValue())[0] + '\\Pictures'
        if not os.path.exists(dir_path):
            os.mkdir(dir_path)

        diagram_pages = self.interface.get_diagram_pages(diagram_name)
        for diagram_page in diagram_pages:
            file_name = (dir_path + '\\' + self.interface.\
                         get_name_of(diagram_page))
            if create_copy == True:
                self.create_copy_of(file_name + '.wmf', extension='.wmf')
            try:
                os.remove(file_name + '.wmf')
            except:
                pass
            if self.output_detail >= OutputDetail.NORMAL:
                            self.interface.print("    Saving " + file_name)
            self.interface.save_page_in_wmf(diagram_page, file_name)

    def create_word_with_wmf(self, window, diagram_name, file_name, file_path=None,
                             wmf_path=None):
        '''
        create the the given file_name word file containing the wmf with name  
        containing the or equal to diagram_name. the wmf file are searched in the 
        wmf_path is provided otherwise in the "wmf" directory in the result file
        path
        '''
        file_path = file_path if file_path != None else\
                os.path.split(window.results_file_name.GetValue())[0]
        wmf_path = wmf_path if wmf_path != None else\
                os.path.split(window.results_file_name.GetValue())[0] + '\\wmf'
        # open word
        try:
            # word = client.Dispatch("Word.Application")
            word = client.Dispatch("kwps.Application")
        except Exception as e:
            try:
                word = client.Dispatch("Word.Application")
            except Exception as e:
                self.interface.print("ERROR:    Exception Running Word!")
                return
        Doc = word.Documents.Open(file_path + '\\' + file_name)
        word.Visible = True

        diagram_pages = self.interface.get_diagram_pages(diagram_name)
        for index, diagram_page in enumerate(diagram_pages):
            pic = Doc.Paragraphs(index + 1).Range.Words(2).InlineShapes.AddPicture\
            (wmf_path + '\\' + self.interface.get_name_of(diagram_page) + '.wmf')

    def add_wmfs_in_word(self, word_file_name, wmf_path):
        '''
        replace inside the given word file all the tags with the relevant wmf 
        pictures found in the given wmf_path
        '''
        # open word
        try:
            word = client.Dispatch("kwps.Application")
        except Exception as e:
            try:
                word = client.Dispatch("Word.Application")
            except Exception as e:
                self.interface.print("ERROR:    Exception Running Word!")
                return

        picture_tag = "<PIC>"
        picture_end_tag = "<\PIC>"

        doc = word.Documents.Open(word_file_name)
        word.Visible = True
        for i in range(doc.Paragraphs.Count - 1):
            try:
                paragraph = doc.Paragraphs(i + 1).Range.Text
                if picture_tag in paragraph:
                    start_index = paragraph.find(picture_tag) + len(picture_tag)
                    end_index = paragraph.find(picture_end_tag)
                    wmf_info = paragraph[start_index:end_index].split(':')
                    try:
                        if self.output_detail >= OutputDetail.NORMAL:
                            self.interface.print("Transfering " + wmf_info[0] + \
                                                 " -" + wmf_info[1])
                        doc.Paragraphs(i + 1).Range.Text = ""
                        inlineshapes = doc.Paragraphs(i + 1).Range.Words(1).InlineShapes
                        new_picture = inlineshapes.AddPicture\
                        (wmf_path + '\\' + wmf_info[0] + '\\' + wmf_info[1] + '.wmf', \
                        doc.Paragraphs(i + 2).Range)
                        shape = inlineshapes.Item(1).ConvertToShape()
                        # magic numbers to crop the picture from the original PF wmf
                        new_picture.PictureFormat.CropBottom = 210
                        new_picture.PictureFormat.CropRight = 290
                        new_picture.ScaleWidth = wmf_info[2] if len(wmf_info) > 2\
                                                             else 51.5
                        new_picture.ScaleHeight = wmf_info[3] if len(wmf_info) > 3\
                                                             else 40                       
                        shape.WrapFormat.Type = 4  # wdWrapFront
                        shape.WrapFormat.AllowOverlap = False
                        shape.Left = word.CentimetersToPoints(wmf_info[4] \
                                                    if len(wmf_info) > 4 else 0.01)
                        shape.Top = word.CentimetersToPoints(wmf_info[5]\
                                                    if len(wmf_info) > 5 else 0.1)
                    except Exception as e:
                        self.interface.print("ERROR: " + wmf_path + '\\' + \
                        wmf_info[0] + '\\' + wmf_info[1] + '.wmf' + " not found!")
            except Exception as e:
                self.interface.print("ERROR: at doc line " + str(i) + "(of" + \
                                     str(doc.Paragraphs.Count) + ")")

        doc.SaveAs(word_file_name.replace(".docx", " with pictures.docx"))
#         document = Document()
#
#         document.add_heading('EEP Network relay coordination results', 0)
#
#         diagram_pages = self.interface.get_diagram_pages(diagram_name)
#         for diagram_page in diagram_pages:
#             picture_full_path = '"' + wmf_path + '\\' + \
#                             self.interface.get_name_of(diagram_page) + '.wmf' + '"'
#             Image.open("C:\\pippo.wmf").save("C:\\pippo.wmf")
#             document.add_picture("C:\\pippo.wmf", \
#                                     width=Inches(1.25))
#
#             document.add_picture(picture_full_path, \
#                                     width=Inches(1.25))
#             document.add_page_break()
#
#         document.save(file_path + '\\' + file_name + '.docx')

    def fill_trafo_output_file(self, i_list, trafo_name, trafo_shc_worksheet, \
                                    start_index, cubicles):
        '''
        put the current trafo current values in the given worksheet as difference 
        with the given starting index
        '''
        if self.output_detail >= OutputDetail.NORMAL:
            self.interface.print(" Saving " + trafo_name + " data")
        num_of_meas = 4  # 3 phases + N
        fault_strings = [ftype_key for ftype_key, nfault_type \
                        in self.fault_types.items() if nfault_type.active == True]
        num_of_faults = len(fault_strings)
        windings_number = len(cubicles)
        for fault_winding_index in range(windings_number):
            for fault_index in range(0, num_of_faults):
                row_index = start_index + fault_index + fault_winding_index * num_of_faults
                bus_name = self.interface.get_name_of(\
                            self.interface.get_cubicle_busbar_of(cubicles[fault_winding_index]))
                trafo_shc_worksheet.write('A' + str(row_index), trafo_name)
                trafo_shc_worksheet.write('B' + str(row_index) , \
                                              fault_strings[fault_index])
                trafo_shc_worksheet.write('C' + str(row_index) , \
                                        str(fault_winding_index) + ' - ' + bus_name)
                for winding in range(0, windings_number):
                    for i in range(0, num_of_meas):
                        trafo_shc_worksheet.write(chr(ord('D') + i + winding * num_of_meas) \
                             +str(row_index) , \
                        str(round(i_list[fault_winding_index * 32 + winding * num_of_meas + \
                        fault_index * num_of_meas * 4 + i], 5)).\
                        replace(".", ","))

    def fill_generator_output_file(self, generator_zscc_data, gen_sheet, start_index):
        '''
        put the current trafo current values in the given worksheet as difference 
        with the given starting index
        '''
        from math import pi
        from cmath import phase
        num_of_meas = 3  # zA, phi_A, zB, phi_B, zC, phi_C   as complex
#         fault_strings = [ftype_key for ftype_key, nfault_type \
#                     in self.fault_types.items() if nfault_type.active == True]
        fault_strings = ['3psc']
        num_of_faults = len(fault_strings)
        for fault_index in range(0, num_of_faults):
                data_index = 0
                row_index = start_index + fault_index
                gen_sheet.write('A' + str(row_index) , \
                    self.interface.get_name_of(generator_zscc_data[data_index].line))
                gen_sheet.write('B' + str(row_index) , \
                                           fault_strings[fault_index])
#                 bus_name = self.interface.get_name_of(\
#                          self.interface.get_cubicle_busbar_of(\
#                     generator_zscc_data[data_index].observation_cubicles[side_index]))
#                 gen_sheet.write('D' + str(row_index) ,\
#                 self.interface.get_name_of(generator_zscc_data[data_index].fault_bus))

                # created a list of zscc data ordered using the fault_bus
                # the first item is not ordered
                if generator_zscc_data:
                    sorted_generator_zscc_data = generator_zscc_data.copy()
                    # remove the first element
                    sorted_generator_zscc_data.pop(0)
                    if sorted_generator_zscc_data:
                        try:
                            sorted_generator_zscc_data.sort(key=lambda x: x.fault_bus, \
                                                                    reverse=False)
                        except Exception as e:
                            pass
                    # add again the first element
                    sorted_generator_zscc_data.insert(0, generator_zscc_data[0])
                for data in sorted_generator_zscc_data:
                    if fault_strings[fault_index] != data.fault_type.value:
                        continue
                    breaker_status = "breaker open" if data.open_breaker == True\
                                else "breaker closed"
                    column = self._calculate_tap_values_column(\
                                            data_index, 0, \
                                 start_column='B' if data_index == 0 else 'C', \
                                 number_of_values=7)
                    gen_sheet.write(row_index, column, \
                      self.interface.get_name_of(data.fault_bus) + "-" + \
                      breaker_status)

                    for i in range(0, num_of_meas):
                        column = self._calculate_tap_values_column(\
                                        data_index, i * 2, \
                                start_column='D', number_of_values=7)
                        gen_sheet.write(row_index, column, \
                        str(round(abs(data.zscc_values[i]), 4)).\
                        replace(".", ","))
                        column = self._calculate_tap_values_column(\
                                        data_index, i * 2 + 1, \
                                start_column='D', number_of_values=7)
                        gen_sheet.write(row_index, column, \
                        str(round(phase(data.zscc_values[i]) * 180 / pi, 2)).\
                        replace(".", ","))
                    data_index += 1

    def fill_z3_output_file(self, z3_line_workbook, z3_trafo_workbook, \
                        initial_output_row=0, trafo_initial_output_row=0):
        '''
        function using the Z3s_zsc_data to send in the results ouput files the
        Z values used by the Z3 calculation
        '''
        from math import pi
        from cmath import phase
        num_of_meas = 3  # zA, phi_A, zB, phi_B, zC, phi_C   as complex
        # iterator for the trafo items
        trafo_row_index = trafo_initial_output_row
        # iterate throw the lines
        row_index = 0
        for row_index, z3_zscc_data in enumerate(self.z3s_zscc_data.values()):
            # created a list of zscc data ordered using the fault_bus
            # the first item is not ordered
            sorted_z3_zscc_data = None
            if z3_zscc_data:
                sorted_z3_zscc_data = z3_zscc_data.copy()
#                 # remove the first element
#                 sorted_z3_zscc_data.pop(0)
#                 if sorted_z3_zscc_data:
#                     try:
#                         sorted_z3_zscc_data.sort(key = lambda x: x.fault_bus,\
#                                                                 reverse = False)
#                     except Exception as e:
#                         pass
#                 # add again the first element
#                 sorted_z3_zscc_data.insert(0, z3_zscc_data[0])
            # iterator for the trafo items
            trafo_data_index = 0
            data_index = 0
            for data in sorted_z3_zscc_data:
                sheet = z3_trafo_workbook if data.trafo_shc else\
                    z3_line_workbook
                # special index management for trafo data
                row_index = trafo_row_index + trafo_initial_output_row\
                         if data.trafo_shc else row_index + initial_output_row
                used_data_index = trafo_data_index if data.trafo_shc == True \
                                                         else data_index
                # add the relay name only for the first data item
                if used_data_index == 0:
                    sheet.write(row_index + 1, 0 , \
                    self.interface.get_name_of(data.relay))
                sheet.write(row_index + 1, 2 , '3PH')
                breaker_status = "breaker open" if data.open_breaker == True and\
                            data.trafo_shc == False else "breaker closed"
                column = self._calculate_tap_values_column(\
                                        used_data_index, 0, \
                             start_column='B' if used_data_index == 0 else 'C', \
                             number_of_values=7)
                sheet.write(row_index + 1, column, \
                  self.interface.get_name_of(data.fault_bus) + "-" + \
                  breaker_status)
                # iterate throw the line Z values for a shc in a bus bar
                for i in range(0, num_of_meas):
                    column = self._calculate_tap_values_column(\
                                    used_data_index, i * 2, \
                            start_column='D', number_of_values=7)
                    sheet.write(row_index + 1, column, \
                    str(round(abs(data.zscc_values[i]), 4)).\
                    replace(".", ","))
                    column = self._calculate_tap_values_column(\
                                    used_data_index, i * 2, \
                            start_column='E', number_of_values=7)
                    angle_value = phase(data.zscc_values[i]) * 180 / pi \
                                if data.zscc_values[i] != 0 else 0
                    sheet.write(row_index + 1, column, \
                    str(round(angle_value, 2)).\
                    replace(".", ","))
                if data.trafo_shc == True:
                    trafo_data_index += 1
                else:
                    data_index += 1
            # if at least on transformer data set has been recorded increase the
            # transformer row count
            if  trafo_data_index > 0:
                trafo_row_index += 1
        return row_index, trafo_row_index

    def fill_line_output_file(self, line_iscc_data, line_sheet, start_index):
        '''
        put the current trafo current values in the given worksheet as difference 
        with the given starting index
        '''
        num_of_meas = 8  # A,B,C,N localy and remotely
        fault_strings = [ftype_key for ftype_key, nfault_type \
                    in self.fault_types.items() if nfault_type.active == True]
        num_of_faults = len(fault_strings)
        for fault_index in range(0, num_of_faults):
            for side_index, side in enumerate(['Local', 'Remot']):
                data_index = 0
                row_index = start_index + fault_index * 2 + side_index
                line_sheet.write(row_index, 0 , \
                    self.interface.get_name_of(line_iscc_data[data_index].line))
                line_sheet.write(row_index, 1 , \
                                           fault_strings[fault_index])
#                 bus_name = self.interface.get_name_of(\
#                          self.interface.get_cubicle_busbar_of(\
#                     line_iscc_data[data_index].observation_cubicles[side_index]))
                line_sheet.write(row_index, 2 , \
                side + ' - ' + \
                self.interface.get_name_of(line_iscc_data[data_index].fault_bus))
                for data in line_iscc_data:
                    if (side == 'Local' and data.remote_bus == True) or\
                    (side == 'Remot' and  data.remote_bus == False):
                        continue
                    if fault_strings[fault_index] != data.fault_type:
                        continue
                    column = self._calculate_tap_values_column(\
                                            data_index, 0, start_column='C')
                    line_sheet.write(row_index, column, \
                     side + ' - ' + self.interface.get_name_of(data.fault_bus))
                    for i in range(0, num_of_meas):
                        column = self._calculate_tap_values_column(\
                                        data_index, i + 1, start_column='C')
                        line_sheet.write(row_index, column, \
                        str(round(data.iscc_values[i], 1)).\
                        replace(".", ","))
                    data_index += 1

    def get_path_branches(self, path, grid):
        '''
        function returning in a list all branches which are part of the given path in
        the given grid
        '''
        return_branch_list = []
        # get the lines contained in the given path
        lines_references = self.interface.get_content(path, '*.IntRef')
        lines = []
        # from  the refrences list  create the lines list
        for reference in lines_references:
            line_object = self.interface.get_referenced_object_of(reference)
            if self.interface.get_class_name_of(line_object) == 'ElmLne' or\
                'ElmTr' in self.interface.get_class_name_of(line_object):
                lines.append(line_object)
        for branch in grid.branch_list:
            next_branch = False
            branch_lines = branch.get_lines()
            for branch_line in branch_lines:
                for line in lines:
                    if self.interface.get_name_of(branch_line) == \
                    self.interface.get_name_of(line):
                        return_branch_list.append(branch)
                        next_branch = True
                        break
                if next_branch == True:
                    break
        return return_branch_list

    def refresh_tdds(self):
        '''
        entry point refreshing all Time distance diagrams
        '''
        self.interface.refresh_TD_diagram(None)
        # the code here below requires PF2018
#         self.interface.SHC_sweep_update('spgf')
#         self.interface.SHC_sweep_update('spgf25')
#         self.interface.SHC_sweep_update('spgf50')
#         self.interface.SHC_sweep_update('3psc')

    def create_TDD(self, window, input_settings=None):
        '''
        function performing the TDD creation process
        '''
#         import pydevd
#         pydevd.settrace()
        lines, trafos = self.create_system_layout(window, input_settings)
        settings = input_settings if input_settings != None \
        else window.GetSettings()
        path_list = [window.path_list[pathindex] for pathindex in settings["PathList"]]
        if len(path_list):
            self.interface.print("\nUsing selected paths")
        else:
            if self.grid != None:
                try:
                    if self.output_detail >= OutputDetail.NORMAL:
                        self.interface.print("\nCreating paths")
                    path_list = self.grid.create_paths(lines, self.relay_list)
                except KeyboardInterrupt:
                    if self.output_detail >= OutputDetail.NORMAL:
                        self.interface.print("\nInterrupted by the user!")

        # Print some info in the PF output window
        if self.output_detail.value >= OutputDetail.NORMAL.value:
            self.interface.print("*********************************************************")
            self.interface.print("    Time Distance Creator Tool (Beta) 0.01")
            self.interface.print("                 Creating TDDs            ")
            self.interface.print("*********************************************************\n")

        # iterate throw the path list...
        if path_list:
            try:  # main loop
                self.interface.set_ldf_basic_configuration()

                for path in path_list:
                    branches = self.get_path_branches(path, self.grid)
                    relays_in_path = []
                    trafo_in_branches = False
                    for branch in branches:
                        relays_in_path.append(tuple(branch.get_relay_list(0)))
                        relays_in_path.append(tuple(branch.get_relay_list(1)))
                        if branch.is_transformer_branch() == True:
                            trafo_in_branches = True
                            break
                    # don't draw the paths which contain a trafo
                    if trafo_in_branches == True:
                        continue
                    # TDD with kilometric calculation
                    diagram_name = 'KM ' + self.interface.get_name_of(path)
                    time_distance_diagram = self.interface.create_time_distance_diagram(\
                                                    diagram_name, path, relays_in_path)
                    self.interface.set_path_method(time_distance_diagram, 'kilometrical')

                    shc_already_set = False

                    # association between a relay type and the line formatting
                    line_format = namedtuple("line_format", "style tickness")
                    relay_type_line_format = {"F21 Distance Mho" :\
                                                line_format(style=13,
                                                            tickness=1),
                                    "F67N_F50N_F51N Neutral directional overc":\
                                                line_format(style=12,
                                                            tickness=1)}

                    # TDD with SHC sweep for the different  fault types
                    for ftype_key, nfault_type in self.fault_types.items():
                        # Loop through the fault types
                        if nfault_type.active == False:  # Skip this fault type
                            continue
                        # set the shc object only once
                        if shc_already_set == False:
                            shc_already_set = True
                            SC_object = self.interface.get_shc_sweep_object()
                            # set the calculation method as "Complete"
                            if SC_object != None:
                                SC_object.iopt_mde = 3
                        # set the shc object fault type
                        self.interface.set_fault(faultype=nfault_type.type, \
                                   resistance=nfault_type.R, \
                                   single_shc=True, \
                                    shc_object=SC_object)

                        diagram_name = 'SW-' + ftype_key + ' ' + \
                        str(nfault_type.R if nfault_type.R > 0 else '') + ' ' + \
                                            self.interface.get_name_of(path)
                        time_distance_diagram = self.interface.create_time_distance_diagram(\
                                            diagram_name, path, relays_in_path, \
                                            relay_type_line_format)
                        self.interface.set_path_method(time_distance_diagram, 'short_circuit_sweep')
                        # refresh the page to recalculate the shc sweep (only for PF V17)
                        self.interface.set_echo_on()
                        self.interface.refresh_TD_diagram(time_distance_diagram)
                self.interface.set_echo_on()
    #           self.save_diagrams_as_wmf(window, 'TD-')
    #             self.create_word_with_wmf(window,
    #                                    diagram_name = 'TD-',
    #                                    file_name =  "report.doc")
            except KeyboardInterrupt:
                if self.output_detail >= OutputDetail.NORMAL:
                    self.interface.print("\nInterrupted by the user!")

        self.interface.set_echo_on()
        return 0

    def is_dialog_setting_ok(self, window, settings):
        ''' 
        function checking that the dialog settings are correct
        '''
        # Check if a PF project is active
        if self.interface.is_project_active() == False:
            dlg = wx.MessageDialog(
                window, "Please activate a project.", "No project is active", \
                wx.OK | wx.ICON_WARNING)
            dlg.ShowModal()
            dlg.Destroy()
            return False
        # check if the specified bus bar is available in the power system
        if  len(settings['StudySelectedBus']) > 0 and \
        len(self.interface.get_element_by_name(settings['StudySelectedBus'])) == 0:
            dlg = wx.MessageDialog(
                window, "'" + settings['StudySelectedBus'] + "' busbar not available", \
                "Error", wx.OK | wx.ICON_WARNING)
            dlg.ShowModal()
            dlg.Destroy()
            return False
        # check if at least one SHC type has been selected.
        if settings['Fslg'] == False and \
           settings['Fdlg'] == False and \
           settings['Fltl'] == False and \
           settings['Ftph'] == False and \
           settings['Fslgr'] == False and \
           settings['Fltlr'] == False and \
           settings['Fdlgr'] == False:
            dlg = wx.MessageDialog(
                window, "Please activate at least one short circut type", \
                "No short circuit type is active", \
                wx.OK | wx.ICON_WARNING)
            dlg.ShowModal()
            dlg.Destroy()
            return False
        return True

    def create_protecton_layout_logic(self, grid, lines, number_of_levels):
        '''
         function finding how the relays are linked together to protect each 
         line in zone 1, zone 2 etc
         '''
        self.relay_matrix = grid.create_relay_matrix_for(lines, \
                            number_of_levels if number_of_levels > 0 else 1)

    #=========================================================================
    # Trafo/Line selection Logic functions
    #=========================================================================

    def get_element(self, transformer, dummy):
        '''
        function returning the trafo passed as parameter
        '''

        return transformer

    #=========================================================================
    # Logic functions
    #=========================================================================

    def get_line_fault_removal_time_of(self, line):
        '''
        return the fault removal time of the given line
        '''

        fault_clearance_time = self.interface.get_line_fault_removal_time_of(line)
        # in case of a "real" trip time return it
        if fault_clearance_time < self.interface.get_relay_NO_TRIP_constant() - 0.1:
            return fault_clearance_time
        ikss_1 = self.interface.get_attribute(line, 'm:Ikss:bus1')
        ikss_2 = self.interface.get_attribute(line, 'm:Ikss:bus2')
        # the trip time i 9999.999 but there is not current at the busses
        # so I get the total shc_trace trace time which is when the fault has been removed
        if ikss_1 == 0 and ikss_2 == 0:
            return self.interface.get_shc_trace_actual_time_step()
        else:
            return fault_clearance_time

    def get_regional_lines_of(self, line):
        '''
        function returning as regional lines the lines belonging to the same 
        grid of the given linbe 
        '''
        lines = self.interface.get_grid_lines_of(
            self.interface.get_line_grid_of(line))
        lines.remove(line)
        return lines

    #=========================================================================
    #   Dummy functions for Dictionary + function logics
    #=========================================================================

    def void_function(self, line, distance):
        '''
        generator function doing nothing 
        '''
        yield ""

    def void_function_param(self, pf_object):
        '''
        null function with one input parameter for the search criteria
        '''
        return object

    def void_function_2_params(self, pf_object, key):
        '''
        null function with two input parameters for the search criteria
        '''
        return object

    #=========================================================================
    #   Contingency functions
    #=========================================================================

    def n_1_Outage_of_Largest_Local_Infeed(self, line, fault_location):
        '''
        generator function disabling the  largest local infeed line
        '''
        if self.output_detail >= OutputDetail.DEBUG:
            self.interface.print("\t\t\t'Outage_of 1st largest local infeed.'")
        busbar = self.get_shc_busbar(
            line, fault_location.d)  # get the bus where to apply the SHC
        if self.interface.is_energized(busbar) == True:
            # fault position on the busbar
            self.interface.set_fault_position(busbar, single_shc=True)
            error = self.interface.run_shc(single_shc=True)  # run the SHC
            if error == 0 and self.ldf_status.ldf_failed == False:
                # get the lines connected to such a busbar
                connections = [connection for connection in \
                               self.interface.get_bus_connections_of(busbar) \
                               if connection != line]
                # create a list of the SHC value for any line
                shc_currents = [self.interface.get_shc_I_bus1_in(
                    nconnection) for nconnection in connections]
                # remove the line only if at least an active line is available
                if len([i for i in shc_currents if i > 0]) > 0:
                    # get the index of the max value
                    max_index = max(self.xrange(len(shc_currents)),
                                    key=shc_currents.__getitem__)
                    maxline = connections[max_index]
                    self.interface.switch_off(maxline)  # switch off the max current line
                    yield self.interface.get_name_of(maxline)
                    self.interface.switch_on(maxline)  # switch on the max current line
                else:
                    if self.output_detail >= OutputDetail.NORMAL:
                        self.interface.print(
                            "\t\t\t'Outage of Largest Local Infeed': '" +
                            self.interface.get_name_of(line) +
                            "' is the unique energized line. Skipping  network configuration")
            else:
                if self.output_detail >= OutputDetail.NORMAL:
                    self.interface.print(
                        "\t\t\t'Outage of Largest Local Infeed': " +
                        self.interface.get_name_of(busbar) +
                        "' shc_trace not possible. Skipping  network configuration")
        else:
            if self.output_detail >= OutputDetail.NORMAL:
                    self.interface.print(
                        "\t\t\t'Outage of Largest Local Infeed': " +
                        self.interface.get_name_of(busbar) +
                        "' is not energized. Skipping  network configuration")

    def n_1_Outage_of_2nd_Largest_Local_Infeed(self, line, fault_location):
        '''
        generator function disabling the 2nd largest local infeed line
        '''
        if self.output_detail >= OutputDetail.DEBUG:
            self.interface.print("\t\t\t'Outage_of 2nd largest local infeed.'")
        busbar = self.get_shc_busbar(
            line, fault_location.d)  # get the bus where to apply the SHC
        if self.interface.is_energized(busbar) == True:
            # fault position on the busbar
            self.interface.set_fault_position(busbar, single_shc=True)
            error = self.interface.run_shc(single_shc=True)  # run the SHC
            if error == 0 and self.ldf_status.ldf_failed == False:
                # get the lines connected to such a busbar
                connections = [connection for connection in \
                               self.interface.get_bus_connections_of(busbar) \
                               if connection != line]
                # create a list of the SHC value for any line
                shc_currents = [self.interface.get_shc_I_bus1_in(
                    nconnection) for nconnection in connections]
                # remove the line only if at least an active lines is available
                if len([i for i in shc_currents if i > 0]) > 0:
                    # get the index of the max value
                    max_index = max(self.xrange(len(shc_currents)),
                                    key=shc_currents.__getitem__)
                    shc_currents[max_index] = 0  # remove this line
                    # get the index of the max value again so it's the 2nd max value
                    max_index = max(self.xrange(len(shc_currents)),
                                    key=shc_currents.__getitem__)
                    maxline = connections[max_index]
                    self.interface.switch_off(maxline)  # switch off the max current line
                    yield self.interface.get_name_of(maxline)
                    self.interface.switch_on(maxline)  # switch on the max current line
                else:
                    if self.output_detail >= OutputDetail.NORMAL:
                        self.interface.print(
                            "\t\t\t'Outage of 2nd Largest Local Infeed': two \
                            energized lines are not available. \
                            Skipping network configuration")
            else:
                if self.output_detail >= OutputDetail.NORMAL:
                    self.interface.print(
                        "\t\t\t'Outage of 2nd Largest Local Infeed': '" +
                        self.interface.get_name_of(busbar) +
                        "' shc_trace not possible. Skipping network configuration")
        else:
            if self.output_detail >= OutputDetail.NORMAL:
                    self.interface.print(
                        "\t\t\t'Outage of 2nd Largest Local Infeed': '" +
                        self.interface.get_name_of(busbar) +
                        "' is not energized. Skipping network configuration")

    def n_1_Outage_of_Largest_Regional_Infeed(self, line, fault_location):
        '''
        generator function disabling the  largest regional line
        '''
        if self.output_detail >= OutputDetail.DEBUG:
            self.interface.print("\t\t\t'Outage_of largest regional infeed.'")
        busbar = self.get_shc_busbar(
            line, fault_location.d)  # get the bus where to apply the SHC
        if self.interface.is_energized(busbar) == True:
            # fault position on the busbar
            self.interface.set_fault_position(busbar, single_shc=True)
            error = self.interface.run_shc(single_shc=True)  # run the SHC
            if error == 0 and self.ldf_status.ldf_failed == False:
                connections = self.get_regional_lines_of(line)
                # create a list of the SHC value for any line
                shc_currents = [self.interface.get_shc_I_bus1_in(nconnection) \
                            for nconnection in connections]
                if self.output_detail >= OutputDetail.VERBOSEDEBUG:
                    self.interface.print("\t\t\t'Outage of Largest Regional Infeed': ")
                    self.interface.print("\t\t\t\tLine SHC currents: ")
                    for index, i in enumerate(shc_currents):
                        self.interface.print("\t\t\t\t{}: {} kA".\
                            format(self.interface.get_name_of \
                                   (connections[index]), i))
                # remove the line only if at least 1 active line is available
                if len([i for i in shc_currents if i > 0]) > 0:
                    # get the index of the max value
                    max_index = max(self.xrange(len(shc_currents)),
                                    key=shc_currents.__getitem__)
                    maxline = connections[max_index]
                    self.interface.switch_off(maxline)  # switch off the max current line
                    yield self.interface.get_name_of(maxline)
                    self.interface.switch_on(maxline)  # switch on the max current line
                else:
                    if self.output_detail >= OutputDetail.NORMAL:
                        self.interface.print(
                            "\t\t\t'Outage of Largest Regional Infeed': a regional \
                            energized line is not available. Skipping network configuration")
            else:
                if self.output_detail >= OutputDetail.NORMAL:
                    self.interface.print(
                        "\t\t\t'Outage of Largest RegionalInfeed': '" +
                        self.interface.get_name_of(busbar) +
                        "' shc_trace not possible. Skipping network configuration")
        else:
            if self.output_detail >= OutputDetail.NORMAL:
                    self.interface.print(
                        "\t\t\t'Outage of Largest RegionalInfeed': '" +
                        self.interface.get_name_of(busbar) +
                        "' is not energized. Skipping network configuration")

    def n_1_Outage_of_Each_Local_Infeed_in_Turn(self, line, fault_location):
        '''
        generator function returning in turn all the lines connected to the 
        selected busbar
        the busbar is selected given the line element and the fault distance 
        along the line
        only the given line is ignored and not returned
        '''
        if self.output_detail >= OutputDetail.DEBUG:
            self.interface.print("\t\t\t'Outage_of each local infeed in turn.'")
        busbar = self.get_shc_busbar(
            line, fault_location.d)  # get the bus where to apply the SHC
        # fault position on the busbar
        self.interface.set_fault_position(busbar, single_shc=True)
        error = self.interface.run_shc(single_shc=True)  # run the SHC
        if error == 0 and self.ldf_status.ldf_failed == False:
            # get the lines connected to such a busbar
            connections = [connection for connection in \
                           self.interface.get_bus_connections_of(busbar)  \
                           if connection != line]
            if self.output_detail >= OutputDetail.VERBOSEDEBUG:
                self.interface.print("\t\t\t\t'Connected items:'")
                for nconnection in connections:
                    self.interface.print("\t\t\t\t\t" + \
                                         self.interface.get_name_of(nconnection))
            for nconnection in connections:
                if self.interface.get_shc_I_bus1_in(nconnection) > 0:
                    self.interface.switch_off(nconnection)  # switch off the line
                    yield self.interface.get_name_of(nconnection)
                    self.interface.switch_on(nconnection)  # switch on the line

    def n_2_Outage_of_Two_Largest_Infeeds(self, line, fault_location):
        '''
        generator function disabling the 2 largest local infeed line
        '''
        if self.output_detail >= OutputDetail.DEBUG:
            self.interface.print("\t\t\t'Outage_of Two Largest Infeeds'")
        busbar = self.get_shc_busbar(
            line, fault_location.d)  # get the bus where to apply the SHC
        if self.interface.is_energized(busbar) == True:
            # fault position on the busbar
            self.interface.set_fault_position(busbar, single_shc=True)
            error = self.interface.run_shc(single_shc=True)  # run the SHC
            if error == 0 and self.ldf_status.ldf_failed == False:
                # get the lines connected to such a busbar
                connections = [connection for connection in \
                               self.interface.get_bus_connections_of(busbar) \
                               if connection != line]
                # create a list of the SHC value for any line
                shc_currents = [self.interface.get_shc_I_bus1_in(
                    nconnection) for nconnection in connections]
                # remove the line only if at least 2 active lines are available
                if len([i for i in shc_currents if i > 0]) > 1:
                    # get the index of the max value
                    max_index1 = max(self.xrange(len(shc_currents)),
                                     key=shc_currents.__getitem__)
                    shc_currents[max_index1] = 0  # remove this line
                    maxline1 = connections[max_index1]
                    # get the index of the max value again so it's the 2nd max value
                    max_index2 = max(self.xrange(len(shc_currents)),
                                     key=shc_currents.__getitem__)
                    maxline2 = connections[max_index2]
                    self.interface.switch_off(maxline1)  # switch off the max current line 1
                    self.interface.switch_off(maxline2)  # switch off the max current line 2
                    yield self.interface.get_name_of(maxline1) + ", " + \
                                    self.interface.get_name_of(maxline2)
                    self.interface.switch_on(maxline1)  # switch on the max current line 1
                    self.interface.switch_on(maxline2)  # switch on the max current line 2
                else:
                    if self.output_detail >= OutputDetail.NORMAL:
                        self.interface.print(
                            "\t\t\t'Outage_of Two Largest Infeeds': three energized \
                        lines are not available. Skipping network configuration")
            else:
                if self.output_detail >= OutputDetail.NORMAL:
                    self.interface.print(
                        "\t\t\t'Outage_of Two Largest Infeeds': '" +
                        self.interface.get_name_of(busbar) +
                        "' shc_trace not possible. Skipping  network configuration")
        else:
            if self.output_detail >= OutputDetail.NORMAL:
                    self.interface.print(
                        "\t\t\t'Outage_of Two Largest Infeeds': '" +
                        self.interface.get_name_of(busbar) +
                        "' is not energized. Skipping  network configuration")

    def stuck_Breaker_Circuit_Breaker_Fail(self, line, fault_location):
        '''
        function disabling all relays at that side of the line to simulate the 
        circuit breaker failure
        '''
        if self.output_detail >= OutputDetail.DEBUG:
            self.interface.print("\t\t\t'Breaker_Circuit_Breaker_Fail'")
        disabled_relays = []
        relay_list = self.interface.get_branch_relays_of(line,
                                            fault_location.reference_breaker)
        for relay in relay_list:
            if self.interface.is_out_of_service(relay) == False:
                disabled_relays.append(relay)
                self.interface.disable(relay)
        if len(disabled_relays) > 0:
            disabled_relays_names = ""
            for nrelay in disabled_relays:
                disabled_relays_names += self.interface.get_name_of(nrelay) + \
                str(", ")
            yield disabled_relays_names
        else:
            if self.output_detail >= OutputDetail.NORMAL:
                    self.interface.print("\t\t\t'Breaker_Circuit_Breaker_Fail':\
                    No relay available at '" + self.interface.get_name_of(line) +
                    "' (side {}) . Skipping  network configuration".
                    format(fault_location.reference_breaker))
        for relay in disabled_relays:
            self.interface.enable(relay)

    #=======================================================================
    #   Auxiliary functions
    #=======================================================================

    def get_shc_busbar(self, line, distance):
        '''
        function getting the right busbar where to apply a fault given the line 
        element and the position the fault along
        the line. The selected busbar is the closest one.
        '''
        return (self.interface.get_branch_bus1_of(line)
                if distance < 50 else self.interface.get_branch_bus2_of(line))

    def fill_relay_list(self, grid):
        '''
        fill self.relaylist with all relay data 
        '''
        pf_relay_list = self.interface.get_relays(
            [self.interface.RelayType.OVERCURRENT,
                            self.interface.RelayType.DISTANCE])

        for relay in pf_relay_list:
            # filter the relays not considering out of service relays and subrelays
            if self.interface.is_out_of_service(relay) or \
            self.interface.is_subrelay(relay):
                continue
            bus_bar = self.interface.get_relay_busbar_of(relay)
            breaker_list = self.interface.get_relay_breaker_of(relay)
            if len(breaker_list) > 0:
                breaker = breaker_list[0]
            else:
                breaker = None
            substation = self.interface.get_busbar_substation_of(bus_bar) \
                                    if bus_bar != None else None

            relay_branch = grid.get_branch_of(relay)
            from_station = self.create_station_label(bus_bar)
            if relay_branch != None:
                to_station = self.create_station_label(relay_branch.terminal_bus_list[1])\
                                if relay_branch.terminal_bus_list[1] != None and \
                                   relay_branch.terminal_bus_list[1] != bus_bar \
                                else self.create_station_label(\
                                            relay_branch.terminal_bus_list[0]) \
                                if relay_branch.terminal_bus_list[0] != bus_bar\
                                    else ""
            else:
                to_station = ""
            phase_threshold_list = \
            self.interface.get_relay_time_defined_overcurrent_elements_thresholds_of\
            (relay, self.interface.ElementType['phase'])
            phase_minimum_instantaneous_threshold = min(phase_threshold_list)\
                                        if len(phase_threshold_list) > 0  else 0
            ground_threshold_list = \
            self.interface.get_relay_time_defined_overcurrent_elements_thresholds_of\
            (relay, self.interface.ElementType['ground'])
            ground_minimum_instantaneous_threshold = min(ground_threshold_list)\
                                        if len(ground_threshold_list) > 0  else 0

            new_relay = self.Relay(name=self.interface.get_name_of(relay),
                                   manufacturer=self.interface.get_relay_manufacturer_of(relay),
                                   model=self.interface.get_relay_model_name_of(relay),
                                   substation=self.interface.get_name_of(substation) if substation != None else "",
                                   busbar=bus_bar,
                                   cubicle=self.interface.get_relay_cubicle_of(relay),
                                   protected_item=self.interface.get_relay_protected_item_of(relay),
                                   voltage=self.interface.get_busbar_rated_voltage_in(bus_bar) if bus_bar != None else 0,
                                   cbr_optime=self.interface.get_breaker_operating_time(breaker) if breaker != None  else 0,
                                   pf_relay=relay,
                                   is_backup_relay=not self.interface.is_main_protection(relay),
                                   phase_minimum_threshold=phase_minimum_instantaneous_threshold,
                                   ground_minimum_threshold=ground_minimum_instantaneous_threshold,
                                   measurement=self.interface.get_relay_measurement_element(relay),
                                   relay_branch=relay_branch,
                                   from_station=from_station,
                                   to_station=to_station)
            self.relay_list.append(new_relay)

    def fill_tripping_time_list(self, grid, network_configuration, use_shc_trace=False):
        '''
        function returning a list with the tripping time of all relays listed 
        inside self.relaylist
        '''
        if use_shc_trace == True:
            # get the list of the tripped relays
            # and fill self.tripping_data_list with the relays tripping times
            for tripped_device in self.interface.get_shc_trace_tripped_devices():
                for index, relay in enumerate(self.relay_list):
                    if tripped_device == relay.pf_relay:
                        self.tripping_data_list[index] = self.Tripping_data(\
                        trip_time=\
                        self.interface.get_shc_trace_device_trip_time_of(tripped_device),
                        tripping_element_string=self.tripping_data_list[index].tripping_element_string,
                        currents=self.tripping_data_list[index].currents, \
                        at_load_bus=grid.is_load_bus(\
                                    self.interface.get_relay_busbar_of(relay.pf_relay), \
                                    relay.protected_item, network_configuration), \
                        breaker_failure=False)
                        break

            return self.tripping_data_list
        else:
            return [self.Tripping_data(trip_time=self.interface.get_relay_tripping_time_of(relay.pf_relay),
                tripping_element_string=self.get_names_string_of(\
                    self.interface.get_relay_tripping_element_of(relay.pf_relay)),
                currents=self.interface.get_relay_current_measures_of(\
                relay.measurement) if relay.measurement != None else (0, 0, 0, 0), \
                at_load_bus=grid.is_load_bus(\
                                    self.interface.get_relay_busbar_of(relay.pf_relay), \
                                    relay.protected_item, network_configuration), \
                breaker_failure=self.interface.is_out_of_service(relay.pf_relay))
                for index, relay in enumerate(self.relay_list)]

    def create_notripping_time_list(self):
        ''' 
        function creating a list of NOTRIP triiping times for all available
        relays. It's used to fill the object when the LDF failed  
        '''
        return [self.Tripping_data(\
                trip_time=self.interface.get_relay_NO_TRIP_constant(),
                tripping_element_string='',
                currents=(0, 0, 0, 0), \
                at_load_bus=False, \
                breaker_failure=False)
                for _ in self.relay_list]

    def init_tripping_data(self, grid, network_configuration):
        '''
        function initializing the tripping data list (only for the SHC trace)
        '''
        # init the tripping data object with no trip data
        # creating a list with the tripping data "notrip" for all relays
        self.tripping_data_list = [(self.Tripping_data(\
                                trip_time=self.interface.get_relay_NO_TRIP_constant(),
                                tripping_element_string='', \
                                currents=(0, 0, 0, 0), \
                                at_load_bus=grid.is_load_bus(\
                                self.interface.get_relay_busbar_of(relay.pf_relay), \
                                relay.protected_item, network_configuration), \
                                breaker_failure=self.interface.is_out_of_service(relay.pf_relay))) \
                                for relay in self.relay_list]

    def collect_relays_trip_data(self, grid, network_configuration):
        '''
        collect the relay tripped elements and measured currents
        '''
        # set the current values and the tripped element names
        # in the self.tripping_data_list list (the list used by the shc_trace trace)

        tripped_devices_list = self.interface.get_shc_trace_tripped_devices()
        started_devices_list = self.interface.get_shc_trace_started_devices()
        for index, relay in enumerate(self.relay_list):
            if relay.pf_relay in started_devices_list:
                relay_currents = self.interface.get_relay_current_measures_of(relay.measurement) \
                if relay.measurement != None else (0, 0, 0, 0)
                self.tripping_data_list[index] = self.Tripping_data(\
                    trip_time=\
                    self.interface.get_shc_trace_device_trip_time_of(relay.pf_relay),
                    tripping_element_string=self.get_names_string_of(\
                        self.interface.get_relay_tripping_element_of(relay.pf_relay)),
                    currents=relay_currents if max(relay_currents) > 0 else \
                        self.tripping_data_list[index].currents, \
                        at_load_bus=grid.is_load_bus(\
                                    self.interface.get_relay_busbar_of(relay.pf_relay), \
                                    relay.protected_item, network_configuration), \
                                    breaker_failure=self.interface.is_out_of_service(relay.pf_relay))
            elif relay.pf_relay in tripped_devices_list:
                self.tripping_data_list[index] = self.Tripping_data(\
                    trip_time=\
                    self.interface.get_shc_trace_device_trip_time_of(relay.pf_relay),
                    tripping_element_string=self.tripping_data_list[index].tripping_element_string,
                    currents=self.tripping_data_list[index].currents, \
                    at_load_bus=grid.is_load_bus(\
                                self.interface.get_relay_busbar_of(relay.pf_relay), \
                                relay.protected_item, network_configuration), \
                    breaker_failure=self.interface.is_out_of_service(relay.pf_relay))
            else:
                relay_currents = self.interface.get_relay_current_measures_of(relay.measurement) \
                if relay.measurement != None else (0, 0, 0, 0)
                self.tripping_data_list[index] = self.Tripping_data(\
                trip_time=self.tripping_data_list[index].trip_time,
                tripping_element_string=self.tripping_data_list[index].tripping_element_string,
                currents=relay_currents if max(relay_currents) > 0 else \
                        self.tripping_data_list[index].currents, \
                        at_load_bus=grid.is_load_bus(\
                                    self.interface.get_relay_busbar_of(relay.pf_relay), \
                                            relay.protected_item, network_configuration), \
                        breaker_failure=self.interface.is_out_of_service(relay.pf_relay))

    def show_results_in_IE(self, file_name):
        '''
        Function running Internet Explorer and showing the result file
        '''
        ie = webbrowser.get(webbrowser.iexplore)
        ie.open("file://" + file_name)

    def create_station_label(self, busbar):
        '''
        function creating the busbar lable which includes busbar name and
        substation name
        '''
        busbar_name = self.interface.get_name_of(busbar)
        busbar_substation = self.interface.get_busbar_substation_of(busbar)
        substation_name = self.interface.get_name_of(busbar_substation)\
                            if busbar_substation != None else ""
        return substation_name + '/' + busbar_name if len(substation_name) > 0 \
          else busbar_name

    def get_names_string_of(self, input_list):
        '''
        function returning a string containing the names of the objects present 
        in the given list
        '''
        return_string = ""
        for item in input_list:
            return_string += self.interface.get_name_of(item) + " "
        return return_string

    def xrange(self, x):
        return iter(range(x))

    #===========================================================================
    # Object serialization
    #===========================================================================

    def serialize(self, object_to_serialize, file_name):
        '''
#         generic function to serialize the object passed as parameter
#         the object can be up a 2 dimension matrix
#         '''
#         file_name = file_name.replace('.xml', '.json')
#         file_name = file_name.replace('.XML', '.json')
#         file_name = file_name.replace('.', type(object_to_serialize).__name__+'.') # the file where I save the obejct contains the object class name
#
#
#         with open(file_name, mode='w') as output_file:
#             json.dump(object_to_serialize, output_file)
#             if type(object_to_serialize) is list:
#                 for item in object_to_serialize:
#                     if type(item) is list:
#                         for subitem in item:
#                             json.dump(len(item), output_file)
#                             json.dump(subitem, output_file)
#                     else:
#                         json.dump(len(object_to_serialize), output_file)
#                         json.dump(item, output_file)
#             else:
#                 json.dump(object_to_serialize, output_file)
#             output_file.close()

    #=========================================================================
    #   Input/ouput XML functions
    #=========================================================================

    def writeXSL(self, XMLfilename):
        XSLfilename = XMLfilename.replace('.xml', '.xsl')
        XSLfilename = XSLfilename.replace('.XML', '.XSL')

        XSDfilename = XMLfilename.rsplit("\\", 1)[-1]
        XSDfilename = XSDfilename.replace('.xml', '.xsd')
        XSDfilename = XSDfilename.replace('.XML', '.XSD')
        with open(XSLfilename, 'w') as XSLOutputFile:
            XSLOutputFile.write("<?xml version=\"1.0\" ?>\n")
            XSLOutputFile.write(
                "<xsl:stylesheet xmlns:xsl=\"http://www.w3.org/1999/XSL/Transform\" version=\"1.0\" xmlns:schemaLocation=\"" + XSDfilename + "\">\n")
            XSLOutputFile.write(
                "<xsl:output method=\"html\" version=\"1.0\" encoding=\"UTF-8\" indent=\"yes\" />\n")
            XSLOutputFile.write(
                "<!-- File automatically created by PSET for Digsilent PowerFactory-->\n")
            XSLOutputFile.write("<xsl:template match=\"/CESIPSETRESULTS\">\n")
            XSLOutputFile.write(
                "<html><head><title>CESI Time Distance Diagram Creator 2019 (BETA) Results v0.01</title>\n")
            XSLOutputFile.write("<style media=\"screen\" type=\"text/css\">\n")
            XSLOutputFile.write(
                "table{border-collapse: collapse; border-spacing: 0;}\n")
            XSLOutputFile.write(
                ".CESItableformat {margin:0px;padding:0px;width:100%;border:1px solid #000000;}\n")
            XSLOutputFile.write(".CESItableformat table{\n")
            XSLOutputFile.write(
                "    border-collapse: collapse; border-spacing: 0; width:100%; height:100%; margin:0px;padding:0px;}\n")
            XSLOutputFile.write(".CESItableformat tr:hover {\n")
            XSLOutputFile.write("background-color:#ffffff;}\n")
            XSLOutputFile.write(".CESItableformat td{\n")
            XSLOutputFile.write(
                "vertical-align:middle;background-color:#6d7175;border:1px solid #000000;border-width:0px 1px 1px 0px;text-align:left;padding:5px;\n")
            XSLOutputFile.write(
                "font-size:12px;font-family:verdana;font-weight:normal;color:#ffffff;}\n")
            XSLOutputFile.write(
                ".CESItableformat tr:hover td{background-color:#edeeef;font-size:12px;font-family:verdana;font-weight:bold;color:#6d7175;repeat-x 0 0;}\n")
            XSLOutputFile.write(".CESItableformat td:hover span{\n")
            XSLOutputFile.write(
                "display:inline; position:absolute;     border:2px solid #FFF;  \n")
            XSLOutputFile.write(
                "font-size:12px;font-family:verdana;font-weight:bold;color:#000000;    background:#edeeef repeat-x 0 0;}\n")
            XSLOutputFile.write(".CESItableformat th:hover span{\n")
            XSLOutputFile.write(
                "display:inline; position:absolute; border:2px solid #FFF;  \n")
            XSLOutputFile.write(
                "font-size:12px;font-family:verdana;font-weight:bold;color:#000000;    background:#edeeef repeat-x 0 0;}\n")
            XSLOutputFile.write(".CESItableformat th{\n")
            XSLOutputFile.write(
                "background-color:#003f7f;border:0px solid #000000;text-align:left;border-width:0px 0px 1px 1px;\n")
            XSLOutputFile.write(
                "font-size:12px;font-family:verdana;font-weight:bold;color:#ffffff; padding:7px;}\n")
            XSLOutputFile.write(".CESItableformat span {\n")
            XSLOutputFile.write("z-index:10;display:none; padding:3px 3px;\n")
            XSLOutputFile.write(
                "    margin-top:40px; margin-left:20px; width:1500px; line-height:16px;}\n")
            XSLOutputFile.write("div {padding: 300px 0px 0px 0px;}\n")
            XSLOutputFile.write("</style>\n")
            XSLOutputFile.write("</head>\n")
            XSLOutputFile.write(
                "<table border=\"1\" class=\"CESItableformat\">\n")
            XSLOutputFile.write("<thead>\n")
            XSLOutputFile.write(
                "<tr><th colspan=\"11\">CESI Time Distance Diagram Creator Tool 2019 (BETA) V0.01 Results</th></tr>\n")
            XSLOutputFile.write(
                "<tr><th colspan=\"2\">Results Created:</th><td colspan=\"9\"><xsl:value-of select=\"SimulationStartTime\"/><xsl:text> </xsl:text></td></tr>\n")
            XSLOutputFile.write("<tr><th colspan=\"2\"><span><table>\n")
            XSLOutputFile.write("<tr><th>Parameter</th><th>Value</th></tr>\n")
            XSLOutputFile.write(
                "<tr><th>Minimum CTI</th><td><xsl:value-of select=\"MinCTI\"/><xsl:text> </xsl:text> <xsl:value-of select=\"..//TimeUnit\"/><xsl:text> </xsl:text></td></tr>\n")
            XSLOutputFile.write(
                "<tr><th>Max Permissible Clearance Time for close-in Faults</th><td><xsl:value-of select=\"MaxClearanceTimeNearEnd\"/><xsl:text> </xsl:text> <xsl:value-of select=\"..//TimeUnit\"/><xsl:text> </xsl:text></td></tr>\n")
            XSLOutputFile.write(
                "<tr><th>Max Permissible Clearance Time for remote-end Faults</th><td><xsl:value-of select=\"MaxClearanceTimeFarEnd\"/><xsl:text> </xsl:text> <xsl:value-of select=\"..//TimeUnit\"/><xsl:text> </xsl:text></td></tr>\n")
            XSLOutputFile.write(
                "<tr><th>Definition of Fast Trip Time Faults</th><td><xsl:value-of select=\"MinClearanceTimeFarEnd\"/><xsl:text> </xsl:text> <xsl:value-of select=\"..//TimeUnit\"/><xsl:text> </xsl:text></td></tr>\n")
            XSLOutputFile.write(
                "<tr><th>Maximum Reach for Fast Trippping</th><td><xsl:value-of select=\"MinClearanceDistFarEnd\"/><xsl:text> </xsl:text> <xsl:text> % </xsl:text></td></tr>\n")
            XSLOutputFile.write(
                "<tr><th>Max Overall Permissible Fault Clearance Time</th><td><xsl:value-of select=\"MaxCT\"/><xsl:text> </xsl:text> <xsl:value-of select=\"..//TimeUnit\"/><xsl:text> </xsl:text></td></tr>\n")
            XSLOutputFile.write(
                "<tr><th>Total Short Circuits Applied</th><td><xsl:value-of select=\"NumberShortCircuits\"/><xsl:text> </xsl:text> </td></tr>\n")
            XSLOutputFile.write(
                "<tr><th>Total Violations Found</th><td><xsl:value-of select=\"NumberViolations\"/><xsl:text> </xsl:text> </td></tr>\n")
            XSLOutputFile.write(
                "</table></span>Database:</th><td colspan=\"9\"><xsl:value-of select=\"CAPEDatabase\"/><xsl:text> </xsl:text></td></tr>\n")
            XSLOutputFile.write(
                "<tr><th colspan=\"2\">Network Study Date:</th><td colspan=\"9\"><xsl:value-of select=\"StudyDate\"/><xsl:text> </xsl:text></td></tr>\n")
            XSLOutputFile.write(
                "<tr><th colspan=\"2\">Simulation Voltage:</th><td colspan=\"9\"><xsl:value-of select=\"StudyVoltage\"/><xsl:text> </xsl:text></td></tr>\n")
            XSLOutputFile.write(
                "<tr><th colspan=\"2\">Simulation Area:</th><td colspan=\"9\"><xsl:value-of select=\"StudyArea\"/><xsl:text> </xsl:text></td></tr>\n")
            XSLOutputFile.write(
                "<tr><th colspan=\"2\">Simulation Zone:</th><td colspan=\"9\"><xsl:value-of select=\"StudyZone\"/><xsl:text> </xsl:text></td></tr>\n")
            XSLOutputFile.write(
                "<tr><th colspan=\"2\">Simulation Grid:</th><td colspan=\"9\"><xsl:value-of select=\"StudyGrid\"/><xsl:text> </xsl:text></td></tr>\n")
            XSLOutputFile.write(
                "<tr><th colspan=\"2\">Simulation Path:</th><td colspan=\"9\"><xsl:value-of select=\"StudyPath\"/><xsl:text> </xsl:text></td></tr>\n")
            XSLOutputFile.write(
                "<tr><th colspan=\"2\">Simulation Bus:</th><td colspan=\"9\"><xsl:value-of select=\"StudyBus\"/><xsl:text> </xsl:text></td></tr>\n")
            XSLOutputFile.write(
                "<tr><th>Fault Number</th><th>From Station</th><th>To Station</th><th>Voltage (kV)</th><th>Circuit ID</th>\n")
            XSLOutputFile.write(
                "<th>Distance To Fault(%)</th><th>Fault Type</th><th>Contingency</th><th>Outage(s)</th><th>Fault Clearance Time (s)</th><th>Result</th></tr>\n")
            XSLOutputFile.write("</thead>\n")
            XSLOutputFile.write("<tbody>\n")
            XSLOutputFile.write("<xsl:apply-templates select=\".//Fault\"/>\n")
            XSLOutputFile.write("</tbody>\n")
            XSLOutputFile.write("</table>\n")
            XSLOutputFile.write(
                "<div>Copyright 2018 Electric Power Research Institute, Inc. All rights reserved.</div>\n")
            XSLOutputFile.write("</html>\n")
            XSLOutputFile.write("</xsl:template> \n")
            XSLOutputFile.write("<xsl:template match=\"Fault\">\n")
            XSLOutputFile.write("<tr>\n")
            XSLOutputFile.write("<td>\n")
            XSLOutputFile.write(
                "<span><table><tr><th>Station</th><th>Circuit Breaker</th><th>Voltage (kV)</th><th>Ckt ID</th><th>Tripping Relay</th><th>Tripping Element</th><th>Trip time (SECONDS)</th><th>psetdigsilent_test Result</th>\n")
            XSLOutputFile.write(
                "<th>IA (pu)</th><th>IB (pu)</th><th>IC (pu)</th><th>IN (pu)</th><th>Relay Setting Phase(Ground)(pu)</th></tr>\n")
            XSLOutputFile.write("<xsl:apply-templates select=\".//Relay\"/>\n")
            XSLOutputFile.write("</table></span>\n")
            XSLOutputFile.write("<xsl:value-of select=\"FaultNumber\"/>\n")
            XSLOutputFile.write("</td>\n")
            XSLOutputFile.write(
                "<td><xsl:value-of select=\"FromStation\"/></td>\n")
            XSLOutputFile.write(
                "<td><xsl:value-of select=\"RemoteStation\"/></td>\n")
            XSLOutputFile.write(
                "<td><xsl:value-of select=\"Voltage\"/></td>\n")
            XSLOutputFile.write(
                "<td><xsl:value-of select=\"CircuitID\"/></td>\n")
            XSLOutputFile.write(
                "<td><xsl:value-of select=\"DistanceToFault\"/></td>\n")
            XSLOutputFile.write(
                "<td><xsl:value-of select=\"FaultType\"/></td>\n")
            XSLOutputFile.write(
                "<td><xsl:value-of select=\"Contingency\"/></td>\n")
            XSLOutputFile.write(
                "<td><xsl:value-of select=\"OutagedElement\"/></td>\n")
            XSLOutputFile.write(
                "<td><xsl:value-of select=\"FaultClearanceTime\"/></td>\n")
            XSLOutputFile.write(
                "<td><xsl:value-of select=\"ProtectionPerformanceAssessment\"/></td>\n")
            XSLOutputFile.write("</tr>\n")
            XSLOutputFile.write("</xsl:template>\n")
            XSLOutputFile.write("<xsl:template match=\"Relay\">\n")
            XSLOutputFile.write("<tr>\n")
            XSLOutputFile.write(
                "<td><xsl:value-of select=\"FromStation\"/></td>\n")
            XSLOutputFile.write(
                "<td><xsl:value-of select=\"ToStation\"/></td>\n")
            XSLOutputFile.write(
                "<td><xsl:value-of select=\"Voltage\"/></td>\n")
            XSLOutputFile.write(
                "<td><xsl:value-of select=\"CircuitID\"/></td>\n")
            XSLOutputFile.write(
                "<td><xsl:value-of select=\"LZOPTag\"/></td>\n")
            XSLOutputFile.write(
                "<td><xsl:value-of select=\"TrippingElement\"/></td>\n")
            XSLOutputFile.write(
                "<td><xsl:value-of select=\"TripTime\"/></td>\n")
            XSLOutputFile.write(
                "<td><xsl:value-of select=\"RelayPerformanceAssessment\"/></td>\n")
            XSLOutputFile.write("<td><xsl:value-of select=\"IFA\"/></td>\n")
            XSLOutputFile.write("<td><xsl:value-of select=\"IFB\"/></td>\n")
            XSLOutputFile.write("<td><xsl:value-of select=\"IFC\"/></td>\n")
            XSLOutputFile.write("<td><xsl:value-of select=\"IFN\"/></td>\n")
            XSLOutputFile.write("<td><xsl:value-of select=\"Irelay\"/></td>\n")
            XSLOutputFile.write("</tr>\n")
            XSLOutputFile.write("</xsl:template>\n")
            XSLOutputFile.write("</xsl:stylesheet>\n")
            XSLOutputFile.close()

    def writeXSD(self, XSDfilename):
        # print ("Writing XSD file")
        # XMLfilename = self.results_file_name.GetValue()
        XSDfilename = XSDfilename.replace('.xml', '.xsd')
        XSDfilename = XSDfilename.replace('.XML', '.XSD')
        with open(XSDfilename, 'w') as XSDOutputFile:
            XSDOutputFile.write(
                "<?xml version=\"1.0\" encoding=\"UTF-8\" ?>\n")
            XSDOutputFile.write(
                "<xs:schema xmlns:xs=\"http://www.w3.org/2001/XMLSchema\">\n")
            XSDOutputFile.write(
                "<!--File automatically created by PSET for Digsilent PowerFactory-->\n")
            XSDOutputFile.write("<xs:element name=\"CESIPSETRESULTS\">\n")
            XSDOutputFile.write("  <xs:complexType>\n")
            XSDOutputFile.write("    <xs:sequence>   \n")
            XSDOutputFile.write(
                "<xs:element name=\"StudyDate\" type=\"xs:string\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"DatabaseFile\" type=\"xs:string\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"SimulationStartTime\" type=\"xs:string\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"SimulationID\" type=\"xs:string\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"StudyVoltage\" type=\"xs:string\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"StudyArea\" type=\"xs:string\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"StudyZone\" type=\"xs:string\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"StudyGrid\" type=\"xs:string\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"StudyPath\" type=\"xs:string\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"StudyBus\" type=\"xs:string\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"TimeUnit\" type=\"xs:string\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"MinCTI\" type=\"xs:decimal\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"MaxCT\" type=\"xs:decimal\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"MaxClearanceTimeNearEnd\" type=\"xs:decimal\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"MinClearanceDistFarEnd\" type=\"xs:decimal\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"MinClearanceTimeFarEnd\" type=\"xs:decimal\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"MaxClearanceTimeFarEnd\" type=\"xs:decimal\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"OvercurrentMargin\" type=\"xs:decimal\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"ImpedanceMargin\" type=\"xs:decimal\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"SimulationDepth\" type=\"xs:positiveInteger\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"MutualDepth\" type=\"xs:positiveInteger\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"StudyDate\" type=\"xs:string\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"SimulationDate\" type=\"xs:string\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"Fault\"  maxOccurs=\"unbounded\">\n")
            XSDOutputFile.write("<xs:complexType>\n")
            XSDOutputFile.write("<xs:sequence>\n")
            XSDOutputFile.write(
                "<xs:element name=\"SimulationID\" type=\"xs:string\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"NetworkCaseID\" type=\"xs:string\"  minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"NetworkStateID\" type=\"xs:string\" minOccurs=\"0\"/>  \n")
            XSDOutputFile.write(
                "<xs:element name=\"FaultNumber\" type=\"xs:string\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"FromStation\"  type=\"xs:string\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"FromStationID\"  type=\"xs:string\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"ToStation\"  type=\"xs:string\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"ToStationID\"  type=\"xs:string\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"RemoteStation\"  type=\"xs:string\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"RemoteStationID\"  type=\"xs:string\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"Voltage\"  type=\"xs:decimal\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"CircuitID\"  type=\"xs:positiveInteger\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"OutagedElement\"  type=\"xs:string\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"Contingency\"  type=\"xs:string\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"DistanceToFault\"  type=\"xs:decimal\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"FaultType\"  type=\"xs:string\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"FaultClearanceTime\"  type=\"xs:decimal\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"ProtectionPerformanceAssessment\"  type=\"xs:string\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"Relay\"  maxOccurs=\"unbounded\" minOccurs=\"0\">\n")
            XSDOutputFile.write("  <xs:complexType>\n")
            XSDOutputFile.write("<xs:sequence>  \n")
            XSDOutputFile.write(
                "<xs:element name=\"SimulationID\" type=\"xs:string\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"NetworkCaseID\" type=\"xs:string\"  minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"NetworkStateID\" type=\"xs:string\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"FaultNumber\" type=\"xs:string\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"FromStation\"  type=\"xs:string\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"ToStation\"  type=\"xs:string\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"RemoteStation\"  type=\"xs:string\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"CircuitID\"  type=\"xs:string\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"Voltage\"  type=\"xs:string\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"LZOPTag\"  type=\"xs:string\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"RelayTag\"  type=\"xs:string\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"RelayName\"  type=\"xs:string\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"RelayModel\"  type=\"xs:string\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"TrippingElement\"  type=\"xs:string\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"TrippingCharacteristicValue\"  type=\"xs:string\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"SimulationMeasuredValue\"  type=\"xs:string\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"TripTime\"  type=\"xs:string\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"CBOpenTime\"  type=\"xs:string\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"RelayPerformanceAssessment\"  type=\"xs:string\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"IFA\"  type=\"xs:decimal\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"IFB\"  type=\"xs:decimal\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"IFC\"  type=\"xs:decimal\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"IFN\"  type=\"xs:decimal\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"Irelay\"  type=\"xs:decimal\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write("</xs:sequence>\n")
            XSDOutputFile.write("  </xs:complexType>\n")
            XSDOutputFile.write("</xs:element> \n")
            XSDOutputFile.write("</xs:sequence>\n")
            XSDOutputFile.write("  </xs:complexType>\n")
            XSDOutputFile.write("</xs:element> \n")
            XSDOutputFile.write(
                "<xs:element name=\"BusesStudied\"  type=\"xs:integer\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"LinesStudied\"  type=\"xs:integer\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"TransformersStudied\"  type=\"xs:integer\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"NumberShortCircuits\"  type=\"xs:integer\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write(
                "<xs:element name=\"NumberViolations\"  type=\"xs:integer\" minOccurs=\"0\"/>\n")
            XSDOutputFile.write("    </xs:sequence>\n")
            XSDOutputFile.write("</xs:complexType>\n")
            XSDOutputFile.write("</xs:element>\n")
            XSDOutputFile.write("</xs:schema> \n")
            XSDOutputFile.close()

