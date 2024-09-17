'''
Created on August 1, 24

@author: MB
'''

#!/usr/bin/env python
# encoding: utf-8

import sys
import os
import copy

import wx

#import win32com 
import win32com.client as client

from sympy import *
from lxml import etree

# Excel access tool
import xlrd
import openpyxl
# Word access tool
from docx import Document
from pydoc import doc
# file copy utility
import shutil
# chart export
import PIL
from PIL import ImageGrab, Image
from PIL import WmfImagePlugin

from csv import excel
from collections import namedtuple

import psutil

from setuptools.command.easy_install import sys_executable

import time

from enum import IntEnum

class OutputDetail(IntEnum):
    DISABLED = 0
    NORMAL = 1
    DEBUG = 2
    VERBOSEDEBUG = 3


# delimiters for the excel cell tag
excel_cell_tag = "<E>"
excel_cell_end_tag = "<.E>" 
# delimiters for the pictures tags
picture_tag = "<PIC>"
picture_end_tag = "<.PIC>"
# delimiters for the formulas/equations tags
formula_tag = "<MATH>"
formula_end_tag = "<.MATH>"
#delimiters for the text templates
template_tag = "<T>"
template_end_tag = "<.T>"
end_template_tag = "<ENDT>"
end_template_end_tag = "<.ENDT>"
#delimeters for the ancors
ancor_tag = "<A>"
ancor_end_tag = "<.A>"






def save_excel_charts(excel_file_name):
    '''
    save all chart of the given excell files as wmf images in the pictures dir
    '''
    # open excel
    try:
        excel = client.Dispatch('Excel.Application')
    except Exception as e:      
        print("ERROR:    Exception Running Excel!")
        return
    try:
        #open file
        workbook = excel.Workbooks.Open(excel_file_name)
        #iterate through the sheets
        for sheet in workbook.Worksheets:
            #iterate through the charts
            for chart in sheet.ChartObjects():
                title = chart.chart.ChartTitle.Text
                chart.Copy()
                image = ImageGrab.grabclipboard()
                new_name = os.path.dirname(excel_file_name) + "\\Charts"
                if not os.path.exists(new_name):
                    os.mkdir(new_name)
                file_name_no_ext = os.path.basename(excel_file_name).replace(".xlsx", "")
                file_name_no_ext = file_name_no_ext.replace(".xlsm", "")
                new_name += "\\" + file_name_no_ext
                if not os.path.exists(new_name):
                    os.mkdir(new_name)
                new_name += "\\" + sheet.Name + " " + title + ".png"
                # delete existing pictures
                if os.path.exists(new_name):
                    os.remove(new_name)
                image.save(new_name, 'png') 
        workbook.Close(SaveChanges=False)
    except Exception as e:
        print("ERROR: '" + str(e))
  
  
def find_file(name, path):
    for root, dirs, files in os.walk(path):
        for file in files:
            if name in file:
                return os.path.join(root, file)        
          
          
def add_excel_chart_in_word(word_file_name, chart_path):
        '''
        replace inside the given word file all the tags with the relevant  
        pictures found in the given chart_path
        '''        
        # open word
        try:
            word = client.Dispatch("Word.Application")
        except Exception as e:
            try:
                word = client.Dispatch("kwps.Application")
            except Exception as e:
                print("ERROR:    Exception Running Word!")
                return
        #word.Visible = True
        doc = word.Documents.Open(word_file_name)
        found_pictures = 0
        for i in range(doc.Paragraphs.Count):
            if i == doc.Paragraphs.Count - found_pictures:
                break
            try:
                paragraph = doc.Paragraphs(i + 1).Range.Text
                if picture_tag in paragraph:
                    start_index = paragraph.find(picture_tag) + len(picture_tag)
                    end_index = paragraph.find(picture_end_tag)
                    png_info = paragraph[start_index:end_index].split(':')
                    try:
                        print("     Transfering '" + png_info[0] + \
                                                 " " + png_info[1] + "' chart")
                        doc.Paragraphs(i + 1).Range.Text = ""
                        inlineshapes = doc.Paragraphs(i + 1).Range.Words(1).InlineShapes
                        
                        chart_name = find_file(png_info[0] + ' ' + png_info[1],\
                                               chart_path)
                        new_picture = inlineshapes.AddPicture\
                                      (chart_name) # doc.Paragraphs(i + 2).Range
                        shape = inlineshapes.Item(1).ConvertToShape()
                        # magic numbers to crop the picture from the original PF wmf
                        # at home big screen 210/290
                        new_picture.PictureFormat.CropBottom = 0
                        new_picture.PictureFormat.CropRight = 0
                        new_picture.ScaleWidth = png_info[2] if len(png_info) > 2\
                                                             else 60
                        new_picture.ScaleHeight = png_info[3] if len(png_info) > 3\
                                                             else 75                       
                        shape.WrapFormat.Type = 4  # wdWrapFront
                        shape.WrapFormat.AllowOverlap = False
#                         shape.Left = word.CentimetersToPoints(png_info[4] \
#                                                     if len(png_info) > 4 else 0.01)
#                         shape.Top = word.CentimetersToPoints(png_info[5]\
#                                                     if len(png_info) > 5 else 0.1)
                        found_pictures += 1
                    except Exception as e:
                        print("ERROR: " + chart_path + '\\' + \
                        png_info[0] + ' ' + png_info[1] + '.png')
            except Exception as e:
                print("END OF FILE: at doc line " + str(i) + "(of " + \
                                     str(doc.Paragraphs.Count) + ")")
        if len(doc.TablesOfContents)>0:
            doc.TablesOfContents(1).Update()
        return doc
                
        
                
def add_excel_cell_value_in_string(str_object, cell_data_str, workbook):
    '''
    replace the str_object .text with the sheet content defined in the given
    cell_data_str 2nd parameter
    '''
    start_index = cell_data_str.find(excel_cell_tag) + len(excel_cell_tag)
    end_index = cell_data_str.find(excel_cell_end_tag)
    cell_info = cell_data_str[start_index:end_index].split(':')
    try:
        if len(cell_info) == 2:
            print("      Transfering " + cell_info[0] + ":" + cell_info[1])                                 
            # get the sheet by name      
            data_sheet = workbook[cell_info[0]]
                   
            #data_sheet = workbook.sheet_by_name(cell_info[0])
            # replace the old value with the new value
            # check/manage the case column with two litterals
            #remove spaces
            cell_info[1] = cell_info[1].replace(" ", "")
            if cell_info[1][1].isalpha():
                column = (ord(cell_info[1][0]) - ord('A') + 1) * 26 +\
                          ord(cell_info[1][1]) - ord('A')   
                row = int(cell_info[1][2:]) - 1
            else:  
                column = ord(cell_info[1][0]) - ord('A')
                row = int(cell_info[1][1:]) - 1   
            new_cell = data_sheet.cell(row + 1, column + 1) 
            new_value = data_sheet.cell(row + 1, column + 1).value
            start_index = str_object.text.find(excel_cell_tag)
            if start_index >= 0 and new_value != None:
                end_index = str_object.text.find(excel_cell_end_tag)
                if end_index < 0:
                    end_index = len(str_object.text) 
                else:
                    end_index += len(excel_cell_end_tag)
                full_tag = str_object.text[start_index: end_index]
                str_object.text = str_object.text.replace(full_tag, 
                                                new_value 
                                                if type(new_value) == str\
                                                else str(round(new_value,3)))
                print("          value = " +  (new_value if type(new_value) == str\
                                                else str(round(new_value,3))))  
        else:
            print("ERROR: '" + cell_data_str + "' wrong error format")
    except Exception as e:
        if end_index == -1:
            print("ERROR: '" + str(e))
  
  
def add_equation(object, workbook):  
    '''
    add the equation defined in the equation tags
    '''    
    inline = object.runs
    for i in range(len(inline)):
        if formula_tag in inline[i].text or formula_end_tag in inline[i].text:
            kV, MVA = symbols('kV MVA')
            try:
                # get the equation definition from the paragraph
                start_index = object.text.find(formula_tag) + len(formula_tag)
                end_index = object.text.find(formula_end_tag)
                whole_content = object.text[start_index:end_index]
                eq_data = whole_content.split(':')
                equation_content = eq_data[0]
                try:
                    if formula_tag in object.text and len(equation_content) > 0:               
                        expr1 = parse_expr(equation_content, evaluate = False)          
                        # create MathML structure
                        MVA, kV = symbols('MVA kV')
                        expr1xml = mathml(expr1, printer = 'presentation')
                        tree = etree.fromstring('<math xmlns="http://www.w3.org/1998/Math/MathML">'+expr1xml+'</math>')            
                        # convert to MS Office structure
                        xslt = etree.parse('c:\\Program Files (x86)\\Microsoft Office\\root\\Office16\\MML2OMML.XSL')
                        transform = etree.XSLT(xslt)
                        new_dom = transform(tree)
                        # append the equation in the paragraph
                        if len(eq_data) > 1:
                            object._element.insert(int(eq_data[1]), new_dom.getroot())
                        else:
                            object._element.append(new_dom.getroot())
                except Exception as e:
                    print("Parse equation ERROR: " + str(e) ) 
                # remove the formula info and tags
                inline[i].text = inline[i].text.replace(formula_tag, "")
                while i < len(inline):
                    if len(inline[i].text) > 3 and inline[i].text in equation_content:
                        inline[i].text = ""
                    if len(eq_data) > 1:
                        second_part = ":" + eq_data[1]
                        if second_part in inline[i].text:
                            inline[i].text = inline[i].text.replace(second_part, "")
                    if formula_end_tag in inline[i].text:
                        inline[i].text = inline[i].text.replace(formula_end_tag, "")
                        if len(inline[i].text) > 3 and inline[i].text in equation_content:
                            inline[i].text = ""
                        inline[i].text = inline[i].text.replace(whole_content, "")
                        inline[i].text = inline[i].text.replace(equation_content, "")
                        break
                    else:
                        inline[i].text = inline[i].text.replace(whole_content, "")
                        inline[i].text = inline[i].text.replace(equation_content, "")                
                    i += 1
            except Exception as e:
                print("Add equation ERROR: " + str(e) ) 
 
 
def duplicate_excel_sheets(excel_file_name, excel_output_file_name):
    '''
    create a copy of the Excel file in the Results dir
    '''

    #run Excel
    try:
        excel = client.Dispatch('Excel.Application')
        excel.Application.DisplayAlerts = False 
    except Exception as e:      
        print("ERROR:" + str(e) + " - Exception Running Excel!")
        return
    try:
        #open the Excel file for the api 
        api_workbook = excel.Workbooks.Open(Filename=excel_file_name)
        
        # save the changes
        api_workbook.SaveAs(excel_output_file_name)
        api_workbook.Close(SaveChanges=False)        
        return        
    except Exception as e:
        print("ERROR: '" + str(e) + " - " + excel_file_name)  
  
 
                       
################################################################################

def duplicate_paragraph_in(position_paragraph, paragraph, 
                           text_to_search, text_to_replace):
    """
    add before the given paragraph the given paragraph
    """

    output_para = position_paragraph.insert_paragraph_before()
    for run in paragraph.runs:
        output_run = output_para.add_run("")
        # Run's bold data
        output_run.bold = run.bold
        # Run's italic data
        output_run.italic = run.italic
        # Run's underline data
        output_run.underline = run.underline
        # Run's color data
        output_run.font.color.rgb = run.font.color.rgb
        # other front params
        output_run.font.size = run.font.size
        output_run.font.subscript = run.font.subscript
        output_run.font.superscript = run.font.superscript
        # Run's font data
        output_run.style.name = run.style.name
    # Paragraph's alignment data
    output_para.paragraph_format.alignment = paragraph.paragraph_format.alignment
    output_para.paragraph_format.first_line_indent = paragraph.paragraph_format.first_line_indent
    # add all elments
    for element in paragraph._element:
        output_para._element.append(copy.deepcopy(element))
    # replace the text
    for run in output_para.runs:
        if text_to_search in run.text:
            run.text = run.text.replace(text_to_search, text_to_replace)

def search_paragraph_containing(searched_text, document):
    '''
    function returning the paragraph containing the given text
    '''
    for paragraph in document.paragraphs:
        # check if I've found the template for the sheet
        if searched_text in paragraph.text:
            return paragraph

def delete_paragraph(paragraph, document):
    '''
    delete the given paragraph
    '''
    # clear all text
    # replace the text
    for run in paragraph.runs:
        run.text = ""
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None

################################################################################

def remove_tags(inline, start_index, end_index, paragraph, workbook,  cell_info_str):
    '''
    replace the tags with the actual values from excel and remove the text 
    between the tags
    '''
    char_counter = 0
    start_char =  start_index - len(excel_cell_tag)
    end_char = end_index + len(excel_cell_end_tag)
    for i in range(len(inline)):
        # check that the tag terminator is before the tag start
        # in that case I remove it
        first_terminator = paragraph.text.find(excel_cell_end_tag)
        first_start = paragraph.text.find(excel_cell_tag)
        if (first_terminator < first_start or first_start < 0) and\
            first_terminator > 0:
            first_part = inline[i].text[0:first_terminator+len(excel_cell_end_tag)]
            # check if removing the first part up to the terminator 
            if first_part in cell_info_str + excel_cell_end_tag:\
                inline[i].text = inline[i].text.replace(first_part, "", 1)
            else: # otherwise remove only the terminator
                inline[i].text = inline[i].text.replace(excel_cell_end_tag, "", 1)
        run_length = len(inline[i].text)
        if excel_cell_tag in inline[i].text:
            add_excel_cell_value_in_string(inline[i],
                                        paragraph.text,
                                        workbook)
            end_char -= run_length - len(inline[i].text)                           
            run_length = len(inline[i].text)
        if char_counter >= start_char and\
           (char_counter + run_length) <= end_char and\
           inline[i].text in cell_info_str+excel_cell_end_tag:
            inline[i].text = ""
        elif char_counter < start_char and\
            (char_counter + run_length) < end_char and\
            (char_counter + run_length) > start_char and\
             inline[i].text[start_char-char_counter:] in cell_info_str+excel_cell_end_tag:
                inline[i].text = inline[i].text.replace(\
                        inline[i].text[start_char-char_counter:], "")
        elif char_counter >= start_char and\
             char_counter < end_char and\
             (char_counter + run_length) > end_char and\
             inline[i].text[0:end_char-char_counter] in cell_info_str+excel_cell_end_tag:
                inline[i].text = inline[i].text.replace(\
                        inline[i].text[0:end_char-char_counter], "")           
        char_counter += run_length
    
def add_excel_cell_value_in_table(table, workbook, formula_found = False):
    '''
    replace the table .text with the sheet content
    '''
    for row in table.rows:
        for cell in row.cells:
            try:
                if formula_found == False and formula_tag in cell.text:  
                    add_excel_cell_value_in_table(table, workbook, True)
                    for paragraph in cell.paragraphs:
                        add_equation(paragraph, workbook)
                else:                                  
                    while excel_cell_tag in cell.text or excel_cell_end_tag in cell.text:
                        start_index = cell.text.find(excel_cell_tag) + len(excel_cell_tag)
                        end_index = cell.text.find(excel_cell_end_tag)
                        if end_index == -1:
                            end_index = len(cell.text) 
                        if excel_cell_tag in cell.text: 
                            cell_info_str = cell.text[start_index:end_index]
                        else:
                            cell_info_str = ""
                        for paragraph in cell.paragraphs:
                            inline = paragraph.runs
                            
                            remove_tags(inline, start_index, end_index, 
                                        paragraph, workbook, cell_info_str)                            
            except Exception as e:
                print("ERROR: " + str(e) )    


def add_excel_cell_value_in_paragraph(paragraph, workbook, formula_found = False):
    '''
    replace the paragraph .text with the sheet content
    '''
    try:              
        if formula_found == False and formula_tag in paragraph.text:  
            add_excel_cell_value_in_paragraph(paragraph, workbook, True)
            add_equation(paragraph, workbook)
        else: 
            while excel_cell_tag in paragraph.text or excel_cell_end_tag in paragraph.text:
                start_index = paragraph.text.find(excel_cell_tag) + len(excel_cell_tag)
                end_index = paragraph.text.find(excel_cell_end_tag)
                if end_index == -1:
                    end_index = len(paragraph.text) 
                if excel_cell_tag in paragraph.text: 
                    cell_info_str = paragraph.text[start_index:end_index]
                else:
                    cell_info_str = ""
                inline = paragraph.runs
                
                remove_tags(inline, start_index, end_index, paragraph, workbook, cell_info_str)
                
            print("                    Modified paragraph: '" + paragraph.text + "'")                                                        
    except Exception as e:
        print("ERROR:  " + str(e)) 

def store_word_text_templates(document):
    '''
    return a dictionary of paragraph lists containing the text templates 
    '''
    # iterate through the word paragraphs
    return_templates = {}
    storing_paragraph = False
    sheet_type = ""
    template = []
    for paragraph in document.paragraphs:
        if end_template_tag in paragraph.text:
            storing_paragraph = False
            if len(sheet_type) > 0:
                return_templates[sheet_type] = template
            sheet_type = ""
            template = []
        if storing_paragraph:
            template.append(paragraph) 
        # check if I've found the template for the sheet
        if template_tag in paragraph.text:
            template = []
            storing_paragraph = True
            start_index = paragraph.text.find(template_tag) + len(template_tag)
            end_index = paragraph.text.find(template_end_tag)
            sheet_type = paragraph.text[start_index:end_index]          
    return return_templates        
             

def instanciate_text_templates_in_word(word_file_name, new_sheets_data):
    '''
    create instances of the word text template
    '''
    #open Word
    doc = Document(word_file_name) 
    # get the templates
    templates = store_word_text_templates(doc)
    if new_sheets_data != None:
        position_paragraph = None
        previous_position_paragraph = None
        for new_sheet_data in new_sheets_data:
            searched_text = ancor_tag + new_sheet_data.template_name + ancor_end_tag
            position_paragraph = search_paragraph_containing(searched_text, doc)
            if position_paragraph != None:             
                if previous_position_paragraph != None  and \
                    position_paragraph.text != previous_position_paragraph.text:
                    delete_paragraph(previous_position_paragraph, doc)
                for new_paragraph in templates[new_sheet_data.template_name]:
                    duplicate_paragraph_in(position_paragraph, new_paragraph,
                                           text_to_search = new_sheet_data.template_name,
                                                   text_to_replace = new_sheet_data.sheet_name)
                previous_position_paragraph = position_paragraph  
        if  position_paragraph != None:    
            delete_paragraph(position_paragraph, doc)    
    return doc        


def add_excel_cell_values_in_word(word_file_name, excel_file_name):
    '''
    replace inside the given Word file all the tags with the relevant Excel 
    cell content of the given Excel file 
    '''
    try:
        #open the Excel file
        workbook = openpyxl.load_workbook(excel_file_name, data_only=True)
        #workbook = xlrd.open_workbook(excel_file_name)
        print("\n Trasferring " + excel_file_name)
        try:  
            #open Word
            doc = Document(word_file_name)        
            
            # iterate through the word headers
            for section in doc.sections:
                # header paragraphs
                for paragraph in section.header.paragraphs:
                    # leave if the template start has been found
                    if template_tag in paragraph.text and len(paragraph.text) > 0:
                        break
                    add_excel_cell_value_in_paragraph(paragraph, workbook) 
                # header tables
                for table in section.header.tables:
                    add_excel_cell_value_in_table(table, workbook)     
            # iterate through the word paragraphs
            for paragraph in doc.paragraphs:
                # leave if the template start has been found
                if template_tag in paragraph.text and len(paragraph.text) > 0:
                    break 
                add_excel_cell_value_in_paragraph(paragraph, workbook)
            # iterate through the word tables
            for table in doc.tables:
                add_excel_cell_value_in_table(table, workbook)   
            return doc
        except Exception as e:
            print("ERROR: '" + str(e) + " - " + word_file_name)
    except Exception as e:
        print("ERROR: '" + str(e) + " - " + excel_file_name)
 
 
def process_files(word_file_name, excel_file_name):    
    '''
    perform the whole replacement operation  for the given files
    '''
    print("\nExcel to Word Script running\n")
    print("Processing: '" + excel_file_name + "'\n using '" + word_file_name +\
          "' as template.")
    # compose the final output full path
    output_file_path = os.path.dirname(word_file_name) + "\\Results\\"
    output_word_file_name = output_file_path + os.path.basename(word_file_name) 
    output_excel_file_name = output_file_path + os.path.basename(excel_file_name)   
    # if the results file doesn't exist create it
    if not os.path.exists(output_file_path):
                    os.mkdir(output_file_path)            
    duplicate_excel_sheets(excel_file_name, output_excel_file_name)
    # save on the disk the chart present in the xls file
    print("  Saving the Excel charts")
    save_excel_charts(output_excel_file_name)
    # fill the values in word 
    doc = add_excel_cell_values_in_word(word_file_name, output_excel_file_name)
    doc.save(output_word_file_name)
    # add the chart pictures 
    file_name_no_ext = os.path.basename(excel_file_name).replace(".xlsx", "")
    file_name_no_ext = file_name_no_ext.replace(".xlsm", "")
    chart_path = os.path.dirname(excel_file_name) + "\\Results\\Charts\\" + file_name_no_ext
                   
    print("  Adding Excel charts in the Word report")                            
    doc = add_excel_chart_in_word(output_word_file_name, chart_path)
    doc.SaveAs(output_word_file_name) 
    doc .Close()
    
#     for proc in psutil.process_iter():
#         if proc.name() == "EXCEL.EXE":
#             proc.kill()
    
    

# for index in range(len(word_file_names)):  
#     process_files(word_file_names[index], excel_file_names[index])
# sys.exit()   
       
# code to uncomment to use the code from the command line
#
# number_of_passed_arguments = len(sys.argv)
# if number_of_passed_arguments < 3:
#     print("Excel to Word Script requires 2 Arguments")
#     print("Total arguments passed:", number_of_passed_arguments - 1)
#     print("Format: ExwoTransfer <full path of the word file> <full path of the excell file>")
# else:
#     word_file_name = sys.argv[1]
#     excel_file_name = sys.argv[2]
#             
#     process_files(word_file_name, excel_file_name)
# print("\nExcel to Word ScriptExecution completed!")   
# time.sleep(3)   
  
    
def transfer_excel_values(window, input_settings=None):
        '''
        function transferring in to the excel tags present in the template the
        relevant cell values present in the Excel file
        '''
        # just for test
        # self.interface.rebuild_pf()
        # return
        # just check that the settings are ok
        word_file_name = window.results_file_name.GetValue()  
        excel_file_name = window.source_file_name.GetValue() 
        if word_file_name == "":
            dlg = wx.MessageDialog(window, "Please specify a valid results filename.",
                                           "No result filename", \
                                           wx.OK | wx.ICON_WARNING)
            dlg.ShowModal()
            dlg.Destroy()
            return
        if excel_file_name == "":
            dlg = wx.MessageDialog(window, "Please specify a valid xls filename.",
                                           "No xls filename", \
                                           wx.OK | wx.ICON_WARNING)
            dlg.ShowModal()
            dlg.Destroy()
            return
        
        if window.output_detail >= OutputDetail.NORMAL:
                print("Trasfering Excel values  ...")
   
        process_files(word_file_name, excel_file_name)    
   

