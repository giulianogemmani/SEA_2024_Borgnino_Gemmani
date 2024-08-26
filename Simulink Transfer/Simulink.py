'''
Created on 13 lug 2024

@author: borgn
'''
import os
import matlab.engine
print("Current Working Directory:", os.getcwd())

from docx import Document
from docx.shared import Inches
# Start MATLAB engine
eng = matlab.engine.start_matlab()
# Alternatively, specify the full path if the script is not in the MATLAB path
eng.run('C:/Users/borgn/OneDrive/Desktop/Tesi Simulink/progetto_software_main.m', nargout=0)
# Create a new Document
doc = Document()

# Add a Title
doc.add_heading('MATLAB Figure Report', level=1)

# Add the figure
doc.add_picture('figure.png', width=Inches(5))  # Adjust the width as needed

# Add a caption
doc.add_paragraph('Figure 1: Example plot generated in MATLAB.')

# Save the document
doc.save('C:\\Users\\borgn\\OneDrive\\Desktop\\progetto software\\report.docx')
